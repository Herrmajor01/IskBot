"""
Модуль для расчета процентов по ст. 395 ГК РФ.
"""

import json
import logging
import os
import re
from datetime import datetime, timedelta
from typing import Any, Dict, List, Optional, Tuple

logging.basicConfig(
    level=logging.INFO,
    filename='bot.log',
    format='%(asctime)s - %(levelname)s - %(message)s'
)


def _is_leap_year(year: int) -> bool:
    """Проверяет, является ли год високосным."""
    return year % 4 == 0 and (year % 100 != 0 or year % 400 == 0)


def _days_in_year(year: int) -> int:
    """Возвращает количество дней в году."""
    return 366 if _is_leap_year(year) else 365


def _split_period_by_years(
    start: datetime,
    end: datetime
) -> List[Tuple[datetime, datetime, int]]:
    """
    Разбивает период на подпериоды по границам годов.

    Returns:
        Список кортежей (начало, конец, дней_в_году)
    """
    if start > end:
        return []

    result = []
    current = start

    while current <= end:
        year_end = datetime(current.year, 12, 31)
        period_end = min(year_end, end)
        days_in_yr = _days_in_year(current.year)
        result.append((current, period_end, days_in_yr))
        current = period_end + timedelta(days=1)

    return result


def _calculate_interest_for_period(
    amount: float,
    start: datetime,
    end: datetime,
    rate: float
) -> float:
    """
    Рассчитывает проценты по ст. 395 ГК РФ с учётом границ годов.
    Корректно обрабатывает периоды, пересекающие границу года.
    """
    total_interest = 0.0

    for period_start, period_end, year_days in _split_period_by_years(start, end):
        days = (period_end - period_start).days + 1
        interest = amount * days * rate / 100 / year_days
        total_interest += interest

    return round(total_interest, 2)


def _parse_numeric_value(text: str) -> Optional[float]:
    cleaned = re.sub(r'[^\d.,+\-]', '', text or '')
    cleaned = cleaned.strip().rstrip('.').rstrip(',')
    if not cleaned or cleaned in ('+', '-'):
        return None
    if not re.match(r'^[\+\-]?\d+([.,]\d+)?$', cleaned):
        return None
    return float(cleaned.replace(',', '.'))


def _parse_int_value(text: str) -> Optional[int]:
    cleaned = re.sub(r'[^\d]', '', text or '')
    if not cleaned:
        return None
    try:
        return int(cleaned)
    except ValueError:
        return None


def _parse_periods_from_rows(
    rows: List[List[str]]
) -> Tuple[List[Dict[str, Any]], float]:
    date_pattern = re.compile(r'\d{2}\.\d{2}\.\d{4}')
    periods: List[Dict[str, Any]] = []
    current_sum = 0.0

    for row in rows[1:]:
        if not row:
            continue

        amount_text = (row[0] or '').strip()
        if not amount_text:
            continue
        if re.search(r'[А-Яа-яA-Za-z]', amount_text):
            skip_words = ["сумма", "задолж", "процент", "итого"]
            if any(word in amount_text.lower() for word in skip_words):
                continue

        amount = _parse_numeric_value(amount_text)
        if amount is None:
            continue
        if amount_text.lstrip().startswith('+'):
            period_sum = current_sum + amount
        else:
            period_sum = amount

        date_cells = []
        for idx, cell in enumerate(row):
            match = date_pattern.search(cell or '')
            if match:
                date_cells.append((idx, match.group(0)))
        if len(date_cells) < 2:
            continue

        date_from_idx, date_from_text = date_cells[0]
        date_to_idx, date_to_text = date_cells[1]
        try:
            date_from = datetime.strptime(date_from_text, '%d.%m.%Y')
            date_to = datetime.strptime(date_to_text, '%d.%m.%Y')
        except Exception:
            continue

        days = None
        if date_to_idx + 1 < len(row):
            days = _parse_int_value(row[date_to_idx + 1])
        if not days or days <= 0:
            continue

        rate = None
        for cell in row:
            if '%' in (cell or ''):
                rate = _parse_numeric_value(cell)
                break
        if rate is None:
            for cell in row[date_to_idx + 2:]:
                val = _parse_numeric_value(cell)
                if val is None:
                    continue
                if 0 < val <= 100:
                    rate = val
                    break
        if rate is None or rate <= 0:
            continue

        year_days = None
        for cell in row:
            val = _parse_int_value(cell)
            if val in (365, 366):
                year_days = val
                break

        current_sum = period_sum
        periods.append({
            'sum': period_sum,
            'date_from': date_from,
            'date_to': date_to,
            'days': days,
            'rate': rate,
            'year_days': year_days,
        })

    if not periods:
        raise ValueError("Не удалось распарсить ни одной строки из таблицы")
    return periods, current_sum


def extract_interest_table_rows(docx_path: str) -> List[List[str]]:
    from docx import Document

    doc = Document(docx_path)
    if not doc.tables:
        raise ValueError("В файле .docx отсутствует таблица")
    table = doc.tables[0]
    rows = []
    for row in table.rows:
        rows.append([cell.text.strip() for cell in row.cells])
    if not rows:
        raise ValueError("Таблица пуста")
    return rows


def _parse_date_value(value: str) -> Optional[datetime]:
    try:
        return datetime.strptime(value.strip(), "%d.%m.%Y")
    except Exception:
        return None


def _format_decimal_ru(value: float, decimals: int = 2) -> str:
    fmt = f"{value:,.{decimals}f}"
    return fmt.replace(",", " ").replace(".", ",")


def _format_integer_ru(value: float) -> str:
    fmt = f"{value:,.0f}"
    return fmt.replace(",", " ")


def _update_interest_table_rows(
    rows: List[List[str]],
    new_end: datetime,
    new_interest: float,
    total_interest: float,
    total_days: int,
    average_rate: float,
    new_days: int,
    new_end_str: str
) -> List[List[str]]:
    updated = [row[:] for row in rows]

    # Update last data row (the last row with date_from/date_to)
    last_data_idx = None
    for idx in range(len(updated) - 1, -1, -1):
        row = updated[idx]
        if not row:
            continue
        if row[0].strip().lower().startswith("итого"):
            continue
        if len(row) < 3:
            continue
        if _parse_date_value(row[1]) and _parse_date_value(row[2]):
            last_data_idx = idx
            break

    if last_data_idx is not None:
        row = updated[last_data_idx]
        if len(row) >= 4:
            row[2] = new_end_str
            row[3] = str(new_days)
        if row:
            row[-1] = _format_decimal_ru(new_interest, 2)
        updated[last_data_idx] = row

    # Update total row
    total_idx = None
    for idx, row in enumerate(updated):
        if row and row[0].strip().lower().startswith("итого"):
            total_idx = idx
            break

    if total_idx is not None:
        row = updated[total_idx]
        if len(row) > 3:
            row[3] = str(total_days)
        if len(row) > 8:
            row[8] = _format_decimal_ru(average_rate, 2) + "%"
        if row:
            row[-1] = _format_decimal_ru(total_interest, 2)
        updated[total_idx] = row

    return updated


def parse_periods_from_docx(
    docx_path: str
) -> Tuple[List[Dict[str, Any]], float]:
    """
    Парсит таблицу из .docx и возвращает периоды для расчета процентов.
    Восстановленная версия из старого parsing.py.
    """
    from docx import Document

    doc = Document(docx_path)
    if not doc.tables:
        raise ValueError("В файле .docx отсутствует таблица")
    table = doc.tables[0]
    if len(table.rows) == 0:
        raise ValueError("Таблица пуста")

    rows = []
    for row in table.rows:
        rows.append([cell.text.strip() for cell in row.cells])

    max_cells = max(len(row) for row in rows)
    if max_cells < 6:
        raise ValueError(
            "Таблица должна содержать минимум 6 столбцов в строках данных"
        )

    return _parse_periods_from_rows(rows)


def _load_cached_rates(
    cache_path: str,
    ttl_hours: int
) -> Optional[List[Tuple[datetime, float]]]:
    if not cache_path or ttl_hours <= 0:
        return None
    try:
        if not os.path.exists(cache_path):
            return None
        cache_age = datetime.now() - datetime.fromtimestamp(
            os.path.getmtime(cache_path)
        )
        if cache_age > timedelta(hours=ttl_hours):
            return None
        with open(cache_path, 'r', encoding='utf-8') as handle:
            payload = json.load(handle)
        if isinstance(payload, dict):
            rates_payload = payload.get('rates', payload)
        elif isinstance(payload, list):
            rates_payload = payload
        else:
            return None
        if not isinstance(rates_payload, list):
            return None
        parsed = []
        for item in rates_payload:
            date_str = None
            rate_value = None
            if isinstance(item, dict):
                date_str = item.get('date') or item.get('date_from')
                rate_value = item.get('rate')
            elif isinstance(item, (list, tuple)) and len(item) >= 2:
                date_str = item[0]
                rate_value = item[1]
            if not date_str or rate_value is None:
                continue
            try:
                date_from = datetime.strptime(str(date_str), "%d.%m.%Y")
                rate = float(str(rate_value).replace(',', '.'))
            except Exception:
                continue
            parsed.append((date_from, rate))
        return parsed or None
    except Exception as exc:
        logging.warning("Ошибка чтения кэша ставок: %s", exc)
        return None


def _save_cached_rates(
    cache_path: str,
    rates: List[Tuple[datetime, float]]
) -> None:
    if not cache_path:
        return
    try:
        payload = {
            'rates': [
                {'date': date_from.strftime("%d.%m.%Y"), 'rate': rate}
                for date_from, rate in rates
            ]
        }
        with open(cache_path, 'w', encoding='utf-8') as handle:
            json.dump(payload, handle, ensure_ascii=False, indent=2)
    except Exception as exc:
        logging.warning("Ошибка сохранения кэша ставок: %s", exc)


def _fetch_rates_from_395gk(
    url: str,
    timeout: int
) -> List[Tuple[datetime, float]]:
    try:
        import requests
        from bs4 import BeautifulSoup
    except Exception as exc:
        raise RuntimeError(
            f"Не удалось импортировать зависимости для загрузки ставок: {exc}"
        ) from exc

    response = requests.get(url, timeout=timeout)
    response.raise_for_status()
    soup = BeautifulSoup(response.text, 'html.parser')
    header_cell = soup.find(
        string=re.compile(r'Дата начала применения', re.IGNORECASE)
    )
    table = header_cell.find_parent('table') if header_cell else None
    if not table:
        raise ValueError("Не найдена таблица ставок на странице")

    rates = []
    for row in table.find_all('tr'):
        cells = [c.get_text(strip=True) for c in row.find_all('td')]
        if len(cells) < 2:
            continue
        date_match = re.search(r'\d{2}\.\d{2}\.\d{4}', cells[0])
        rate_match = re.search(r'\d+(?:[.,]\d+)?', cells[1])
        if not date_match or not rate_match:
            continue
        date_from = datetime.strptime(date_match.group(0), "%d.%m.%Y")
        rate = float(rate_match.group(0).replace(',', '.'))
        rates.append((date_from, rate))

    if not rates:
        raise ValueError("Не удалось извлечь ставки из таблицы")

    unique = {}
    for date_from, rate in rates:
        unique[date_from] = rate
    return sorted(unique.items(), key=lambda x: x[0])


def _build_rate_periods(
    rates: List[Tuple[datetime, float]]
) -> List[Tuple[datetime, datetime, float]]:
    if not rates:
        return []
    ordered = sorted(rates, key=lambda x: x[0])
    result = []
    for i, (date_from, rate) in enumerate(ordered):
        date_to = (
            ordered[i + 1][0] - timedelta(days=1)
            if i + 1 < len(ordered) else datetime.max
        )
        result.append((date_from, date_to, rate))
    return result


def get_key_rates_from_395gk() -> List[Tuple[datetime, datetime, float]]:
    """
    Возвращает список ключевых ставок ЦБ РФ с датами действия.
    """
    rates_data = [
        ("01.08.2016", "10.50"), ("19.09.2016", "10.00"),
        ("27.03.2017", "9.75"), ("02.05.2017", "9.25"),
        ("19.06.2017", "9.00"), ("18.09.2017", "8.50"),
        ("30.10.2017", "8.25"), ("18.12.2017", "7.75"),
        ("12.02.2018", "7.50"), ("26.03.2018", "7.25"),
        ("17.09.2018", "7.50"), ("17.12.2018", "7.75"),
        ("17.06.2019", "7.50"), ("29.07.2019", "7.25"),
        ("09.09.2019", "7.00"), ("28.10.2019", "6.50"),
        ("16.12.2019", "6.25"), ("10.02.2020", "6.00"),
        ("27.04.2020", "5.50"), ("22.06.2020", "4.50"),
        ("27.07.2020", "4.25"), ("22.03.2021", "4.50"),
        ("26.04.2021", "5.00"), ("15.06.2021", "5.50"),
        ("26.07.2021", "6.50"), ("13.09.2021", "6.75"),
        ("25.10.2021", "7.50"), ("20.12.2021", "8.50"),
        ("14.02.2022", "9.50"), ("28.02.2022", "20.00"),
        ("11.04.2022", "17.00"), ("04.05.2022", "14.00"),
        ("27.05.2022", "11.00"), ("14.06.2022", "9.50"),
        ("25.07.2022", "8.00"), ("19.09.2022", "7.50"),
        ("24.07.2023", "8.50"), ("15.08.2023", "12.00"),
        ("18.09.2023", "13.00"), ("30.10.2023", "15.00"),
        ("18.12.2023", "16.00"), ("29.07.2024", "18.00"),
        ("16.09.2024", "19.00"), ("28.10.2024", "21.00"),
        ("09.06.2025", "20.00"), ("28.07.2025", "18.00"),
        ("15.09.2025", "17.00"), ("27.10.2025", "16.50"),
        ("22.12.2025", "16.00"),
    ]
    fetched_rates = None
    try:
        from config import CALCULATION_CONFIG
        cache_path = CALCULATION_CONFIG.get('cache_file')
        ttl_hours = int(CALCULATION_CONFIG.get('cache_ttl_hours', 24))
        rates_url = CALCULATION_CONFIG.get(
            'rates_url',
            'https://395gk.ru/svedcb.htm'
        )
        timeout = int(CALCULATION_CONFIG.get('request_timeout', 10))
        cached = _load_cached_rates(cache_path, ttl_hours)
        if cached:
            fetched_rates = cached
        else:
            fetched_rates = _fetch_rates_from_395gk(rates_url, timeout)
            _save_cached_rates(cache_path, fetched_rates)
    except Exception as exc:
        logging.warning("Не удалось загрузить ставки с 395gk: %s", exc)

    key_rates = []
    if fetched_rates:
        key_rates = fetched_rates
    else:
        for date_str, rate_str in rates_data:
            try:
                date_from = datetime.strptime(date_str, "%d.%m.%Y")
                rate = float(rate_str.replace(",", "."))
                key_rates.append((date_from, rate))
            except Exception as e:
                logging.warning(
                    "Ошибка парсинга ставки: %s, %s, %s",
                    date_str, rate_str, e
                )
                continue

    result = _build_rate_periods(key_rates)

    if not result:
        logging.warning(
            "Не удалось получить ключевые ставки, "
            "используется резервная ставка 21%"
        )
        return [(datetime(2025, 1, 1), datetime.max, 21.0)]
    return result


def split_period_by_key_rate(
    start: datetime,
    end: datetime,
    key_rates: List[Tuple[datetime, datetime, float]]
) -> List[Tuple[datetime, datetime, float]]:
    """
    Делит период на подпериоды по ключевым ставкам.
    """
    periods = []
    for k_start, k_end, rate in key_rates:
        actual_start = max(start, k_start)
        actual_end = min(end, k_end)
        if actual_start <= actual_end:
            periods.append((actual_start, actual_end, rate))
    return sorted(periods, key=lambda x: x[0])


def calc_395_on_periods(
    base_sum: float, periods: List[Tuple[datetime, datetime, float]]
) -> Tuple[float, List[Dict[str, Any]]]:
    """
    Рассчитывает проценты по ст. 395 ГК РФ для списка периодов.
    Корректно обрабатывает периоды, пересекающие границу года.
    """
    total_interest = 0.0
    detailed_calc = []

    for start, end, rate in periods:
        # Разбиваем период по годам для корректного расчёта
        year_subperiods = _split_period_by_years(start, end)

        period_interest = 0.0
        formula_parts = []

        for sub_start, sub_end, year_days in year_subperiods:
            sub_days = (sub_end - sub_start).days + 1
            sub_interest = base_sum * sub_days * rate / 100 / year_days
            period_interest += sub_interest
            formula_parts.append(
                f"{base_sum:,.2f} × {sub_days} × {rate}% / {year_days}"
                .replace(',', ' ')
            )

        period_interest = round(period_interest, 2)
        total_interest += period_interest
        total_days = (end - start).days + 1

        detailed_calc.append({
            'period': (
                f"{start.strftime('%d.%m.%Y')} - "
                f"{end.strftime('%d.%m.%Y')}"
            ),
            'date_from': start.strftime('%d.%m.%Y'),
            'date_to': end.strftime('%d.%m.%Y'),
            'sum': base_sum,
            'days': total_days,
            'rate': rate,
            'interest': period_interest,
            'formula': ' + '.join(formula_parts)
        })

    return round(total_interest, 2), detailed_calc


def calculate_full_395(
    docx_path: str,
    today: Optional[datetime] = None,
    key_rates: Optional[List[Tuple[datetime, datetime, float]]] = None
) -> Dict[str, Any]:
    """
    Полный расчет процентов по ст. 395 ГК РФ с парсингом из .docx.
    """
    if today is None:
        today = datetime.now()

    if key_rates is None:
        key_rates = get_key_rates_from_395gk()

    table_rows = None
    try:
        periods, base_sum = parse_periods_from_docx(docx_path)
    except Exception as exc:
        logging.warning(
            "Ошибка парсинга таблицы процентов из %s: %s",
            docx_path,
            exc
        )
        return {
            'total_interest': 0.0,
            'detailed_calc': [],
            'base_sum': 0.0,
            'periods_count': 0,
            'error': str(exc)
        }
    try:
        table_rows = extract_interest_table_rows(docx_path)
    except Exception as exc:
        logging.warning(
            "Ошибка чтения таблицы процентов из %s: %s",
            docx_path,
            exc
        )

    if not periods:
        return {
            'total_interest': 0.0,
            'detailed_calc': [],
            'base_sum': base_sum,
            'error': 'Не найдены периоды для расчета'
        }

    original_last = periods[-1]
    replacement_periods: List[Dict[str, Any]] = []
    if original_last['date_to'] < today:
        split_periods = split_period_by_key_rate(
            original_last['date_from'],
            today,
            key_rates
        )
        if not split_periods:
            split_periods = [
                (original_last['date_from'], today, original_last['rate'])
            ]
        for start, end, rate in split_periods:
            days = (end - start).days + 1
            replacement_periods.append({
                'sum': original_last['sum'],
                'date_from': start,
                'date_to': end,
                'days': days,
                'rate': rate,
            })
        periods = periods[:-1] + replacement_periods

    # Расчет для каждого периода из таблицы
    total_interest = 0.0
    detailed_calc = []

    for p in periods:
        start = p['date_from']
        end = p['date_to']
        rate = p['rate']
        amount = p['sum']

        # Разбиваем период по годам для корректного расчёта
        year_subperiods = _split_period_by_years(start, end)

        period_interest = 0.0
        formula_parts = []

        for sub_start, sub_end, year_days in year_subperiods:
            sub_days = (sub_end - sub_start).days + 1
            sub_interest = amount * sub_days * rate / 100 / year_days
            period_interest += sub_interest
            formula_parts.append(
                f"{amount:,.2f} × {sub_days} × {rate}% / {year_days}"
                .replace(',', ' ')
            )

        period_interest = round(period_interest, 2)
        total_interest += period_interest

        detailed_calc.append({
            'period': (
                f"{start.strftime('%d.%m.%Y')} - "
                f"{end.strftime('%d.%m.%Y')}"
            ),
            'date_from': start.strftime('%d.%m.%Y'),
            'date_to': end.strftime('%d.%m.%Y'),
            'sum': amount,
            'days': p['days'],
            'rate': rate,
            'interest': period_interest,
            'formula': ' + '.join(formula_parts)
        })

    updated_table_rows = None
    if (
        original_last['date_to'] < today
        and len(replacement_periods) == 1
        and table_rows
    ):
        rp = replacement_periods[0]
        new_end = rp['date_to']
        new_days = rp['days']
        new_rate = rp['rate']
        # Используем корректный расчёт с учётом границ года
        new_interest = _calculate_interest_for_period(
            original_last['sum'],
            rp['date_from'],
            new_end,
            new_rate
        )
        total_days = sum(int(item.get('days', 0) or 0) for item in periods)
        if total_days > 0:
            avg_rate = sum(
                float(item.get('rate', 0) or 0)
                * float(item.get('days', 0) or 0)
                for item in periods
            ) / total_days
        else:
            avg_rate = 0.0
        updated_table_rows = _update_interest_table_rows(
            table_rows,
            new_end,
            new_interest,
            round(total_interest, 2),
            total_days,
            avg_rate,
            new_days,
            new_end.strftime('%d.%m.%Y'),
        )

    return {
        'total_interest': round(total_interest, 2),
        'detailed_calc': detailed_calc,
        'base_sum': base_sum,
        'periods_count': len(periods),
        'table_rows': updated_table_rows,
    }
