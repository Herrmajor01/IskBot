"""
Основной модуль Telegram-бота для
автоматизации расчёта процентов по ст. 395 ГК РФ.
"""

import json
import logging
import os
import re
import shutil
import uuid
from datetime import datetime, timedelta
from decimal import Decimal, InvalidOperation, ROUND_DOWN, ROUND_HALF_UP
from typing import Any, Dict, List, Optional, Set, Tuple
from xml.etree import ElementTree as ET

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from dotenv import load_dotenv
from telegram import (InlineKeyboardButton, InlineKeyboardMarkup, InputFile,
                      Update)
from telegram.ext import (Application, CallbackQueryHandler, CommandHandler,
                          ContextTypes, ConversationHandler, MessageHandler,
                          filters)

from cal import calculate_duty
from calc_395 import (calc_395_on_periods, calculate_full_395,
                      get_key_rates_from_395gk, split_period_by_key_rate)
from llm_fallback import (
    apply_llm_fallback,
    extract_document_groups_llm,
    extract_payment_terms_llm,
    extract_transport_details_llm,
    extract_transport_details_vision,
    extract_text_from_image_llm,
    get_vision_config,
    match_cargo_to_application_llm,
    proofread_text_with_llm,
)
from pdf_extractor import (
    estimate_text_quality,
    extract_claim_document_with_vision,
)
from document_awareness import (
    analyze_documents_for_special_cases,
    adjust_claim_data,
    generate_awareness_text_block,
)
from sliding_window_parser import parse_documents_with_sliding_window
from external_claim_parser import (
    parse_external_claim,
    parse_document_packages,
    link_documents_full,
    ExternalClaimData,
)

load_dotenv()
TOKEN = os.getenv("TELEGRAM_BOT_TOKEN")

logging.basicConfig(
    level=logging.INFO,
    filename='bot.log',
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)
logger.info("Bot script started")

# Состояния диалога
(
    ASK_FLOW,
    ASK_DOCUMENT,
    ASK_JURISDICTION,
    ASK_CUSTOM_COURT,
    ASK_CLAIM_STATUS,
    ASK_TRACK,
    ASK_RECEIVE_DATE,
    ASK_SEND_DATE,
    ASK_PRETENSION_FIELD,
    ASK_EXTERNAL_CLAIM_DOCUMENT,  # Для внешних претензий
    ASK_EXTERNAL_CLAIM_FIELD,     # Для заполнения пропущенных полей
    ASK_BIRTH_DATE,
    ASK_BIRTH_PLACE,
) = range(13)

PRETENSION_FIELD_ORDER = [
    "plaintiff_name",
    "defendant_name",
    "debt",
    "payment_days",
    "shipping_method",
    "docs_track_number",
    "docs_received_date",
]

PRETENSION_FIELD_DEFS = {
    "plaintiff_name": {
        "prompt": "Укажите отправителя претензии (истца):",
        "required": True,
    },
    "defendant_name": {
        "prompt": "Укажите получателя претензии (должника):",
        "required": True,
    },
    "debt": {
        "prompt": "Укажите сумму задолженности (в рублях):",
        "required": True,
    },
    "payment_days": {
        "prompt": "Укажите срок оплаты в рабочих днях (только число):",
        "required": True,
    },
    "shipping_method": {
        "prompt": "Документы отправлялись почтой или СДЭК? Напишите: почта или сдэк.",
        "required": True,
    },
    "docs_received_date": {
        "prompt": (
            "Укажите даты получения оригиналов документов "
            "(ДД.ММ.ГГГГ, можно несколько через запятую):"
        ),
        "required": True,
    },
    "docs_track_number": {
        "prompt": (
            "Укажите трек-номера отправки оригиналов документов "
            "(можно несколько через запятую, или напишите «пропустить»):"
        ),
        "required": False,
    },
}


def resolve_court_from_dadata(
    court_name: str,
    court_address: str
) -> Tuple[str, str]:
    if not court_name:
        return court_name, court_address
    suggestion = fetch_dadata_court_suggest(
        court_name,
        court_type="AS"
    )
    if not suggestion:
        return court_name, court_address
    parsed = parse_dadata_court(suggestion)
    name = parsed.get("name") or court_name
    address = parsed.get("address") or court_address
    return name, address


def get_court_by_address(defendant_address: str) -> Tuple[str, str]:
    """
    Определяет суд по адресу ответчика.

    Args:
        defendant_address: Адрес ответчика

    Returns:
        Кортеж (название суда, адрес суда)
    """
    from courts_code import ARBITRATION_COURTS, CITY_TO_REGION

    defendant_address_lower = defendant_address.lower()

    # Сначала ищем по названию региона
    for region, court_info in ARBITRATION_COURTS.items():
        if region.lower() in defendant_address_lower:
            return resolve_court_from_dadata(
                court_info["name"],
                court_info["address"]
            )

    # Если регион не найден, ищем по городам
    for city, region in CITY_TO_REGION.items():
        if city in defendant_address_lower:
            if region in ARBITRATION_COURTS:
                court_info = ARBITRATION_COURTS[region]
                return resolve_court_from_dadata(
                    court_info["name"],
                    court_info["address"]
                )

    # Если ничего не найдено, возвращаем общий ответ
    return resolve_court_from_dadata(
        "Арбитражный суд по месту нахождения ответчика",
        "Адрес суда не определен"
    )


def insert_interest_table(doc, details, total_interest: Optional[float] = None):
    """
    Вставляет таблицу процентов в документ
    Word вместо маркера {interest_table}.
    """
    placeholders = ['{{interest_table}}', '{interest_table}']
    headers = [
        'Сумма', 'Дата начала', 'Дата окончания', 'Дни',
        'Ставка', 'Формула', 'Проценты'
    ]
    for i, paragraph in enumerate(doc.paragraphs):
        if any(ph in paragraph.text for ph in placeholders):
            table = doc.add_table(rows=1, cols=len(headers))
            for col, header in enumerate(headers):
                cell = table.cell(0, col)
                cell.text = header
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(
                            qn('w:eastAsia'), 'Times New Roman'
                        )
            for row in details:
                cells = table.add_row().cells
                row_sum = row.get('sum', 0.0)
                date_from = row.get('date_from', '')
                date_to = row.get('date_to', '')
                if isinstance(date_from, datetime):
                    date_from = date_from.strftime('%d.%m.%Y')
                if isinstance(date_to, datetime):
                    date_to = date_to.strftime('%d.%m.%Y')
                cells[0].text = f"{row_sum:,.2f}".replace(',', ' ')
                cells[1].text = f"{date_from} г." if date_from else ''
                cells[2].text = f"{date_to} г." if date_to else ''
                cells[3].text = str(row.get('days', ''))
                cells[4].text = str(row.get('rate', ''))
                cells[5].text = str(row.get('formula', ''))
                cells[6].text = f"{row.get('interest', 0.0):,.2f}".replace(',', ' ')
                for cell in cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(10)
                            run.font.name = 'Times New Roman'
                            run._element.rPr.rFonts.set(
                                qn('w:eastAsia'), 'Times New Roman'
                            )
            if total_interest is None:
                total_interest = sum(
                    float(row.get('interest', 0.0) or 0.0)
                    for row in details
                )
            if details or total_interest:
                total_row = table.add_row().cells
                label_cell = total_row[0].merge(total_row[5])
                label_cell.text = 'Итого процентов'
                total_row[6].text = f"{total_interest:,.2f}".replace(',', ' ')
                for cell in total_row:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(10)
                            run.font.name = 'Times New Roman'
                            run.bold = True
                            run._element.rPr.rFonts.set(
                                qn('w:eastAsia'), 'Times New Roman'
                            )
            p = paragraph._element
            p.addnext(table._element)
            for placeholder in placeholders:
                paragraph.text = paragraph.text.replace(placeholder, '')
            break


def insert_interest_table_from_rows(doc, table_rows: List[List[str]]) -> bool:
    """
    Вставляет таблицу процентов в документ, используя готовые строки.
    """
    if not table_rows:
        return False

    placeholders = ['{{interest_table}}', '{interest_table}']
    idx = None
    for i, paragraph in enumerate(doc.paragraphs):
        if any(ph in paragraph.text for ph in placeholders):
            idx = i
            break
    if idx is None:
        return False

    max_cols = max(len(row) for row in table_rows)
    table = doc.add_table(rows=0, cols=max_cols)

    def style_row(row, bold: bool = False) -> None:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'
                    run.bold = bold
                    run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman'
                    )

    last_row_index = len(table_rows) - 1
    for row_index, row_values in enumerate(table_rows):
        cells = table.add_row().cells
        for col_index in range(max_cols):
            value = row_values[col_index] if col_index < len(row_values) else ''
            cells[col_index].text = value
        if row_index == 0 and max_cols >= 11:
            if cells[1].text == cells[2].text:
                cells[2].text = ''
            if cells[1].text == cells[3].text:
                cells[3].text = ''
            if cells[4].text == cells[5].text:
                cells[5].text = ''
            if cells[6].text == cells[7].text:
                cells[7].text = ''
        if row_index == last_row_index and cells[0].text.strip().lower().startswith('итого'):
            if cells[0].text == cells[1].text:
                cells[1].text = ''
            if cells[0].text == cells[2].text:
                cells[2].text = ''
        is_header = row_index <= 2
        style_row(table.rows[row_index], bold=is_header)

    # Merge header groups if values are repeated
    if len(table.rows) > 0 and max_cols >= 11:
        header = table.rows[0].cells
        if header[1].text == header[2].text == header[3].text and header[1].text:
            header[1].merge(header[3])
        if header[4].text == header[5].text and header[4].text:
            header[4].merge(header[5])
        if header[6].text == header[7].text and header[6].text:
            header[6].merge(header[7])

    # Merge total row label if repeated
    for row in table.rows[::-1]:
        if row.cells and row.cells[0].text.strip().lower().startswith('итого'):
            if (
                row.cells[0].text == row.cells[1].text
                and row.cells[1].text == row.cells[2].text
            ):
                row.cells[0].merge(row.cells[2])
            style_row(row, bold=True)
            break

    p = doc.paragraphs[idx]._element
    parent = p.getparent()
    parent.remove(p)
    parent.insert(idx, table._element)
    return True


def insert_pretension_interest_table(
    doc,
    details,
    total_interest: Optional[float] = None,
    note: Optional[str] = None
):
    placeholders = ['{{interest_table}}', '{interest_table}']
    if not details and (total_interest is None or float(total_interest or 0) <= 0):
        for paragraph in doc.paragraphs:
            if any(ph in paragraph.text for ph in placeholders):
                paragraph.text = note or (
                    "Срок оплаты не истёк, проценты по ст. 395 ГК РФ "
                    "не начисляются."
                )
                return False
        return False
    for paragraph in doc.paragraphs:
        if not any(ph in paragraph.text for ph in placeholders):
            continue
        table = doc.add_table(rows=3, cols=9)

        header_row = table.rows[0].cells
        header_row[0].text = 'Задолженность,\nруб.'
        header_row[1].text = 'Период просрочки'
        header_row[1].merge(header_row[3])
        header_row[4].text = 'Увеличение долга'
        header_row[4].merge(header_row[5])
        header_row[6].text = 'Ставка'
        header_row[7].text = 'Дней\nв\nгоду'
        header_row[8].text = 'Проценты,\nруб.'

        sub_header = table.rows[1].cells
        sub_header[0].text = 'Задолженность,\nруб.'
        sub_header[1].text = 'c'
        sub_header[2].text = 'по'
        sub_header[3].text = 'дни'
        sub_header[4].text = 'сумма,\nруб.'
        sub_header[5].text = 'дата'
        sub_header[6].text = 'Ставка'
        sub_header[7].text = 'Дней\nв\nгоду'
        sub_header[8].text = 'Проценты,\nруб.'

        index_row = table.rows[2].cells
        for idx, cell in enumerate(index_row, start=1):
            cell.text = f"[{idx}]"

        def style_cell(cell):
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), 'Times New Roman'
                    )

        for row in table.rows:
            for cell in row.cells:
                style_cell(cell)

        for row in details:
            cells = table.add_row().cells
            date_from = row.get('date_from')
            date_to = row.get('date_to')
            if isinstance(date_from, datetime):
                date_from = date_from.strftime('%d.%m.%Y')
            if isinstance(date_to, datetime):
                date_to = date_to.strftime('%d.%m.%Y')
            increase_sum = row.get('increase_sum') or 0.0
            increase_date = row.get('increase_date')
            if isinstance(increase_date, datetime):
                increase_date = increase_date.strftime('%d.%m.%Y')
            rate = row.get('rate', 0.0)
            rate_text = f"{rate:.2f}".replace('.', ',') + "%"
            interest_text = format_money(row.get('interest', 0.0), 2).replace('.', ',')
            cells[0].text = format_money(row.get('sum', 0.0), 0)
            cells[1].text = f"{date_from} г." if date_from else ''
            cells[2].text = f"{date_to} г." if date_to else ''
            cells[3].text = str(row.get('days', ''))
            cells[4].text = format_money(increase_sum, 0) if increase_sum else '0'
            cells[5].text = increase_date or '-'
            cells[6].text = rate_text
            cells[7].text = str(row.get('year_days', ''))
            cells[8].text = interest_text
            for cell in cells:
                style_cell(cell)

        if total_interest is None:
            total_interest = sum(
                float(item.get('interest', 0.0) or 0.0)
                for item in details
            )
        if details or total_interest:
            total_row = table.add_row().cells
            label_cell = total_row[0].merge(total_row[7])
            label_cell.text = 'Итого процентов'
            total_row[8].text = format_money(total_interest or 0.0, 2).replace('.', ',')
            for cell in total_row:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Times New Roman'
                        run.bold = True
                        run._element.rPr.rFonts.set(
                            qn('w:eastAsia'), 'Times New Roman'
                        )

        p = paragraph._element
        p.addnext(table._element)
        for placeholder in placeholders:
            paragraph.text = paragraph.text.replace(placeholder, '')
        break


def generate_claim_paragraph(user_data: dict) -> str:
    claim_date = user_data.get('claim_date', '').strip()
    claim_number = user_data.get('claim_number', '').strip()
    receive_date = user_data.get('postal_receive_date', '').strip()
    claim_status = user_data.get('claim_status', '')

    if claim_status == 'claim_received':
        return (
            f"{claim_date} Истцом в адрес Ответчика была направлена "
            "досудебная претензия с требованием погасить "
            "образовавшуюся задолженность. "
            f"Претензия была отправлена почтовым отправлением "
            f"с трек-номером {claim_number}. "
            f"{receive_date} Ответчик получил данное отправление."
        )
    elif claim_status == 'claim_not_received':
        return (
            f"{claim_date}. Истцом в адрес Ответчика была направлена "
            "досудебная претензия с требованием погасить "
            "образовавшуюся задолженность. "
            f"Претензия была отправлена почтовым отправлением "
            f"с трек-номером {claim_number}. "
            "В соответствии с п. 1 ст. 165.1 ГК РФ Истец считает, "
            "что претензия была получена Ответчиком. "
            "У Ответчика имелось достаточно времени для получения "
            "почтового отправления. "
            "Таким образом, Истцом был соблюден обязательный "
            "претензионный порядок (досудебный порядок урегулирования "
            "споров) в строгом соответствии с действующим "
            "законодательством РФ."
        )
    else:
        return (
            "Не указано г. Истцом в адрес Ответчика была направлена "
            "досудебная претензия."
        )


class RussianPostTrackingError(Exception):
    pass


_DADATA_CACHE: Dict[str, Dict[str, Any]] = {}
_DADATA_COURT_CACHE: Dict[str, Dict[str, Any]] = {}


def get_russian_post_config() -> Dict[str, object]:
    login = os.getenv("RUSSIAN_POST_LOGIN", "").strip()
    password = os.getenv("RUSSIAN_POST_PASSWORD", "").strip()
    return {
        "enabled": bool(login and password),
        "login": login,
        "password": password,
        "endpoint": os.getenv(
            "RUSSIAN_POST_ENDPOINT",
            "https://tracking.russianpost.ru/rtm34"
        ).strip(),
        "language": os.getenv("RUSSIAN_POST_LANGUAGE", "RUS").strip() or "RUS",
        "message_type": os.getenv("RUSSIAN_POST_MESSAGE_TYPE", "0").strip() or "0",
        "timeout": int(os.getenv("RUSSIAN_POST_TIMEOUT", "30") or "30"),
    }


def normalize_tracking_number(value: str) -> str:
    return re.sub(r'[^0-9A-Za-z]', '', value or '').upper()


def is_valid_tracking_number(value: str) -> bool:
    if not value:
        return False
    if re.fullmatch(r'\d{10,20}', value):
        return True
    return bool(re.fullmatch(r'[A-Z]{2}\d{9}[A-Z]{2}', value))


def is_valid_inn(value: str) -> bool:
    return bool(re.fullmatch(r"\d{10}|\d{12}", value or ""))


def get_dadata_config() -> Dict[str, object]:
    token = os.getenv("DADATA_API_KEY", "").strip()
    secret = os.getenv("DADATA_SECRET", "").strip()
    return {
        "enabled": bool(token),
        "token": token,
        "secret": secret,
        "endpoint": os.getenv(
            "DADATA_PARTY_ENDPOINT",
            "https://suggestions.dadata.ru/suggestions/api/4_1/rs/findById/party"
        ).strip(),
        "timeout": int(os.getenv("DADATA_TIMEOUT", "15") or "15"),
    }


def get_dadata_court_config() -> Dict[str, object]:
    token = os.getenv("DADATA_API_KEY", "").strip()
    secret = os.getenv("DADATA_SECRET", "").strip()
    return {
        "enabled": bool(token),
        "token": token,
        "secret": secret,
        "suggest_endpoint": os.getenv(
            "DADATA_COURT_SUGGEST_ENDPOINT",
            "https://suggestions.dadata.ru/suggestions/api/4_1/rs/suggest/court"
        ).strip(),
        "find_endpoint": os.getenv(
            "DADATA_COURT_BY_ID_ENDPOINT",
            "https://suggestions.dadata.ru/suggestions/api/4_1/rs/findById/court"
        ).strip(),
        "timeout": int(os.getenv("DADATA_TIMEOUT", "15") or "15"),
    }


def normalize_shipping_method(value: Optional[str]) -> str:
    if not value:
        return ""
    cleaned = str(value).strip().lower()
    if cleaned in ("сдэк", "cdek", "sdek"):
        return "сдэк"
    if cleaned in ("почта", "post", "russian post", "russianpost"):
        return "почта"
    return cleaned


def normalize_shipping_source(value: Optional[str]) -> str:
    cleaned = normalize_shipping_method(value)
    if cleaned == "сдэк":
        return "cdek"
    if cleaned == "почта":
        return "post"
    return cleaned


def fetch_dadata_party_by_inn(
    inn: str,
    kpp: Optional[str] = None,
    branch_type: str = "MAIN"
) -> Optional[Dict[str, Any]]:
    inn_clean = re.sub(r"[^\d]", "", str(inn or ""))
    if not is_valid_inn(inn_clean):
        return None
    cache_key = f"{inn_clean}:{kpp or ''}:{branch_type}"
    if cache_key in _DADATA_CACHE:
        return _DADATA_CACHE[cache_key]

    config = get_dadata_config()
    if not config.get("enabled"):
        return None

    payload: Dict[str, Any] = {"query": inn_clean}
    if kpp:
        payload["kpp"] = re.sub(r"[^\d]", "", str(kpp))
    if branch_type:
        payload["branch_type"] = branch_type
    payload["type"] = "INDIVIDUAL" if len(inn_clean) == 12 else "LEGAL"

    headers = {
        "Content-Type": "application/json",
        "Accept": "application/json",
        "Authorization": f"Token {config['token']}",
    }
    if config.get("secret"):
        headers["X-Secret"] = str(config["secret"])

    try:
        resp = requests.post(
            str(config["endpoint"]),
            json=payload,
            headers=headers,
            timeout=int(config.get("timeout") or 15)
        )
        resp.raise_for_status()
        data = resp.json()
    except Exception as exc:
        logger.warning("Dadata request failed for INN %s: %s", inn_clean, exc)
        return None

    suggestions = None
    if isinstance(data, dict):
        suggestions = data.get("suggestions")
    if suggestions is None and isinstance(data, list):
        suggestions = data
    if not suggestions:
        return None

    suggestion = suggestions[0]
    if isinstance(suggestion, dict):
        _DADATA_CACHE[cache_key] = suggestion
        return suggestion
    return None


def parse_dadata_party(suggestion: Dict[str, Any]) -> Dict[str, str]:
    data = suggestion.get("data", {}) if isinstance(suggestion, dict) else {}
    name_block = data.get("name") or {}
    fio_block = data.get("fio") or {}
    org_type = data.get("type") or ""

    name_full = (
        name_block.get("short_with_opf")
        or name_block.get("full_with_opf")
        or suggestion.get("value")
        or ""
    )
    if org_type == "INDIVIDUAL":
        fio_parts = [
            fio_block.get("surname", ""),
            fio_block.get("name", ""),
            fio_block.get("patronymic", ""),
        ]
        fio = " ".join(part for part in fio_parts if part).strip()
        if fio:
            name_full = f"ИП {fio}"

    address_block = data.get("address") or {}
    address = (
        address_block.get("unrestricted_value")
        or address_block.get("value")
        or ""
    )

    return {
        "name": name_full,
        "inn": str(data.get("inn") or ""),
        "kpp": str(data.get("kpp") or ""),
        "ogrn": str(data.get("ogrn") or ""),
        "address": str(address or ""),
        "type": str(org_type or ""),
        "branch_type": str(data.get("branch_type") or ""),
    }


def fetch_dadata_court_suggest(
    query: str,
    court_type: Optional[str] = None
) -> Optional[Dict[str, Any]]:
    if not query:
        return None
    cache_key = f"suggest:{query}:{court_type or ''}"
    if cache_key in _DADATA_COURT_CACHE:
        return _DADATA_COURT_CACHE[cache_key]

    config = get_dadata_court_config()
    if not config.get("enabled"):
        return None

    payload: Dict[str, Any] = {"query": query}
    if court_type:
        payload["filters"] = [{"court_type": court_type}]

    headers = {
        "Content-Type": "application/json",
        "Accept": "application/json",
        "Authorization": f"Token {config['token']}",
    }
    if config.get("secret"):
        headers["X-Secret"] = str(config["secret"])

    try:
        resp = requests.post(
            str(config["suggest_endpoint"]),
            json=payload,
            headers=headers,
            timeout=int(config.get("timeout") or 15)
        )
        resp.raise_for_status()
        data = resp.json()
    except Exception as exc:
        logger.warning("Dadata court suggest failed for '%s': %s", query, exc)
        return None

    suggestions = None
    if isinstance(data, dict):
        suggestions = data.get("suggestions")
    if suggestions is None and isinstance(data, list):
        suggestions = data
    if not suggestions:
        return None

    suggestion = suggestions[0]
    if isinstance(suggestion, dict):
        _DADATA_COURT_CACHE[cache_key] = suggestion
        return suggestion
    return None


def fetch_dadata_court_by_code(code: str) -> Optional[Dict[str, Any]]:
    code_clean = (code or "").strip()
    if not code_clean:
        return None
    cache_key = f"code:{code_clean}"
    if cache_key in _DADATA_COURT_CACHE:
        return _DADATA_COURT_CACHE[cache_key]

    config = get_dadata_court_config()
    if not config.get("enabled"):
        return None

    payload: Dict[str, Any] = {"query": code_clean}
    headers = {
        "Content-Type": "application/json",
        "Accept": "application/json",
        "Authorization": f"Token {config['token']}",
    }
    if config.get("secret"):
        headers["X-Secret"] = str(config["secret"])

    try:
        resp = requests.post(
            str(config["find_endpoint"]),
            json=payload,
            headers=headers,
            timeout=int(config.get("timeout") or 15)
        )
        resp.raise_for_status()
        data = resp.json()
    except Exception as exc:
        logger.warning("Dadata court findById failed for '%s': %s", code_clean, exc)
        return None

    suggestions = None
    if isinstance(data, dict):
        suggestions = data.get("suggestions")
    if suggestions is None and isinstance(data, list):
        suggestions = data
    if not suggestions:
        return None

    suggestion = suggestions[0]
    if isinstance(suggestion, dict):
        _DADATA_COURT_CACHE[cache_key] = suggestion
        return suggestion
    return None


def parse_dadata_court(suggestion: Dict[str, Any]) -> Dict[str, str]:
    data = suggestion.get("data", {}) if isinstance(suggestion, dict) else {}
    name = data.get("name") or suggestion.get("value") or ""
    address = data.get("legal_address") or data.get("address") or ""
    return {
        "name": str(name or ""),
        "address": str(address or ""),
        "code": str(data.get("code") or ""),
        "court_type": str(data.get("court_type") or ""),
        "court_type_name": str(data.get("court_type_name") or ""),
        "website": str(data.get("website") or ""),
    }


def build_russian_post_request(barcode: str, config: Dict[str, object]) -> str:
    language = config.get("language", "RUS")
    message_type = config.get("message_type", "0")
    login = config.get("login", "")
    password = config.get("password", "")
    return (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<soap:Envelope xmlns:soap="http://www.w3.org/2003/05/soap-envelope" '
        'xmlns:oper="http://russianpost.org/operationhistory" '
        'xmlns:data="http://russianpost.org/operationhistory/data" '
        'xmlns:soapenv="http://schemas.xmlsoap.org/soap/envelope/">'
        '<soap:Header/>'
        '<soap:Body>'
        '<oper:getOperationHistory>'
        '<data:OperationHistoryRequest>'
        f'<data:Barcode>{barcode}</data:Barcode>'
        f'<data:MessageType>{message_type}</data:MessageType>'
        f'<data:Language>{language}</data:Language>'
        '</data:OperationHistoryRequest>'
        '<data:AuthorizationHeader soapenv:mustUnderstand="1">'
        f'<data:login>{login}</data:login>'
        f'<data:password>{password}</data:password>'
        '</data:AuthorizationHeader>'
        '</oper:getOperationHistory>'
        '</soap:Body>'
        '</soap:Envelope>'
    )


def parse_russian_post_date(value: str) -> Optional[datetime]:
    if not value:
        return None
    cleaned = value.strip()
    try:
        return datetime.fromisoformat(cleaned.replace('Z', '+00:00'))
    except ValueError:
        match = re.search(
            r'\d{4}-\d{2}-\d{2}T\d{2}:\d{2}:\d{2}',
            cleaned
        )
        if match:
            try:
                return datetime.strptime(
                    match.group(0),
                    "%Y-%m-%dT%H:%M:%S"
                )
            except ValueError:
                return None
    return None


def extract_fault_message(root: ET.Element) -> Optional[str]:
    for fault_name in ("AuthorizationFault", "OperationHistoryFault", "LanguageFault"):
        fault_elem = root.find(f".//{{*}}{fault_name}")
        if fault_elem is not None:
            for tag in ("Description", "Message", "FaultString"):
                text = fault_elem.findtext(f".//{{*}}{tag}")
                if text:
                    return text.strip()
            return fault_name
    fault_elem = root.find(".//{*}Fault")
    if fault_elem is not None:
        text = fault_elem.findtext(".//{*}Text") or fault_elem.findtext(".//{*}faultstring")
        if text:
            return text.strip()
        return "Ошибка сервиса отслеживания"
    return None


def fetch_russian_post_operations(barcode: str) -> List[Dict[str, object]]:
    config = get_russian_post_config()
    if not config.get("enabled"):
        raise RussianPostTrackingError("Не настроен доступ к API Почты России.")

    payload = build_russian_post_request(barcode, config)
    try:
        response = requests.post(
            config.get("endpoint"),
            data=payload.encode("utf-8"),
            headers={"Content-Type": "application/soap+xml; charset=utf-8"},
            timeout=config.get("timeout", 30),
        )
    except requests.RequestException as exc:
        raise RussianPostTrackingError(
            "Сервис отслеживания недоступен. Попробуйте позже."
        ) from exc

    if response.status_code != 200:
        raise RussianPostTrackingError(
            f"Сервис отслеживания вернул код {response.status_code}."
        )

    try:
        root = ET.fromstring(response.text)
    except ET.ParseError as exc:
        raise RussianPostTrackingError(
            "Не удалось разобрать ответ сервиса отслеживания."
        ) from exc

    fault_message = extract_fault_message(root)
    if fault_message:
        raise RussianPostTrackingError(fault_message)

    records = []
    for record in root.findall(".//{*}historyRecord"):
        oper_date_raw = record.findtext(".//{*}OperDate")
        oper_date = parse_russian_post_date(oper_date_raw)
        if not oper_date:
            continue
        oper_type = record.findtext(".//{*}OperType/{*}Name") or ""
        oper_attr = record.findtext(".//{*}OperAttr/{*}Name") or ""
        records.append({
            "date": oper_date,
            "oper_type": oper_type.strip(),
            "oper_attr": oper_attr.strip(),
        })

    return sorted(records, key=lambda item: item["date"])


def extract_tracking_dates(records: List[Dict[str, object]]) -> Tuple[Optional[str], Optional[str]]:
    if not records:
        return None, None

    def normalize(text: str) -> str:
        return (text or "").lower().replace("ё", "е")

    send_date = None
    for record in records:
        if "прием" in normalize(str(record.get("oper_type", ""))):
            send_date = record["date"]
            break

    if not send_date:
        send_date = records[0]["date"]

    receive_date = None
    for record in records:
        combined = normalize(
            f"{record.get('oper_type', '')} {record.get('oper_attr', '')}"
        )
        if "неудач" in combined or "отправител" in combined:
            continue
        if "вруч" in combined or "получен" in combined:
            receive_date = record["date"]

    send_str = send_date.strftime("%d.%m.%Y") if send_date else None
    receive_str = receive_date.strftime("%d.%m.%Y") if receive_date else None
    return send_str, receive_str


def format_header_paragraph(paragraph, label, value, postfix=None):
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if label:
        run_label = paragraph.add_run(label)
        run_label.bold = True
        run_label.font.name = 'Times New Roman'
        run_label.font.size = Pt(12)
    run_value = paragraph.add_run(value)
    run_value.bold = True
    run_value.font.name = 'Times New Roman'
    run_value.font.size = Pt(12)
    if postfix:
        run_post = paragraph.add_run(postfix)
        run_post.bold = True
        run_post.font.name = 'Times New Roman'
        run_post.font.size = Pt(12)


def format_header_address(paragraph, value):
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(value)
    run.bold = False
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)


def format_placeholder_paragraph(paragraph, placeholder, value, bold=False):
    text = ''.join(run.text for run in paragraph.runs)
    idx = text.find(placeholder)
    if idx == -1:
        return
    before = text[:idx]
    after = text[idx+len(placeholder):]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if before:
        run_before = paragraph.add_run(before)
        run_before.bold = bold
        run_before.font.name = 'Times New Roman'
        run_before.font.size = Pt(12)
    run_value = paragraph.add_run(value)
    run_value.bold = bold
    run_value.font.name = 'Times New Roman'
    run_value.font.size = Pt(12)
    if after:
        run_after = paragraph.add_run(after)
        run_after.bold = bold
        run_after.font.name = 'Times New Roman'
        run_after.font.size = Pt(12)


def format_placeholder_paragraph_plain(paragraph, value):
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(value)
    run.bold = False
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)


def format_document_list(document_string: str) -> str:
    """
    Форматирует список документов с переносами строк.

    Args:
        document_string: Строка с документами, разделенными точкой с запятой

    Returns:
        Отформатированная строка с переносами строк
    """
    if not document_string or document_string == 'Не указано':
        return document_string

    # Разделяем по точке с запятой и убираем пустые строки
    documents = [doc.strip()
                 for doc in document_string.split(';') if doc.strip()]

    if not documents:
        return document_string

    # Форматируем каждый документ без маркера
    formatted_docs = []
    for doc in documents:
        # Убираем лишние пробелы
        formatted_docs.append(doc.strip())

    # Соединяем документы с переносами строк
    return ';\n'.join(formatted_docs) + ';'


def normalize_str(value: Optional[str], default: str = 'Не указано') -> str:
    if value is None:
        return default
    text = str(value).strip()
    return text if text else default


def maybe_proofread_text(
    value: Optional[str],
    protected_values: Optional[List[str]] = None
) -> str:
    if value is None:
        return value
    text = str(value).strip()
    if not text or text == "Не указано":
        return text
    return proofread_text_with_llm(text, protected_values=protected_values) or text


def add_prefix_if_missing(value: str, prefix: str) -> str:
    if not value or value == 'Не указано':
        return value
    trimmed = value.strip()
    if trimmed.lower().startswith(prefix.strip().lower()):
        return trimmed
    return f"{prefix}{trimmed}"


def parse_amount(value: Optional[str], default: float = 0.0) -> float:
    if value is None:
        return default
    if isinstance(value, (int, float)):
        return float(value)
    cleaned = re.sub(r'\s+', '', str(value)).replace(',', '.')
    try:
        return float(cleaned)
    except ValueError:
        return default


def extract_last_amount_from_text(text: str) -> Optional[float]:
    if not text:
        return None
    candidates = re.findall(r'\d[\d\s]*[.,]\d{2}', text)
    if not candidates:
        return None
    for raw in reversed(candidates):
        value = parse_amount(raw)
        if value > 0:
            return value
    return None


def inflect_ruble_word(amount: float) -> str:
    value = int(abs(amount))
    last_two = value % 100
    last = value % 10
    if 11 <= last_two <= 14:
        return "рублей"
    if last == 1:
        return "рубль"
    if 2 <= last <= 4:
        return "рубля"
    return "рублей"


def replace_ruble_words(text: str) -> str:
    def repl(match: re.Match) -> str:
        number_text = match.group(1)
        amount = parse_amount(number_text)
        word = inflect_ruble_word(amount)
        return f"{number_text} {word}"

    pattern = r'(\d[\d\s]*(?:[.,]\d+)?)\s+рубл(?:ь|я|ей)\b'
    return re.sub(pattern, repl, text)


def parse_amount_decimal(
    value: Optional[str],
    default: Decimal = Decimal("0")
) -> Decimal:
    if value is None:
        return default
    if isinstance(value, Decimal):
        return value
    if isinstance(value, (int, float)):
        return Decimal(str(value))
    cleaned = re.sub(r'\s+', '', str(value)).replace(',', '.')
    try:
        return Decimal(cleaned)
    except InvalidOperation:
        return default


def normalize_payment_terms(text: str) -> str:
    if not text:
        return text
    normalized = re.sub(r'\s+', ' ', str(text)).strip()
    lower = normalized.lower()
    if 'условия оплаты' in lower or 'оплаты по договор' in lower:
        dash_match = re.search(r'\s[–\-]\s*', normalized)
        colon_index = normalized.find(':')
        if colon_index != -1 and (dash_match is None or colon_index < dash_match.start()):
            split_match = re.split(r':\s*', normalized, maxsplit=1)
        else:
            split_match = re.split(r'\s[–\-]\s*', normalized, maxsplit=1)
        if len(split_match) == 2:
            normalized = split_match[1].strip()
    if re.search(r'\bг\.$', normalized):
        return normalized
    normalized = re.sub(r'[.;:]+$', '', normalized).strip()
    if normalized.startswith(('«', '"')) and normalized.endswith(('»', '"')):
        normalized = normalized[1:-1].strip()
    if '«' in normalized or '»' in normalized or '"' in normalized:
        normalized = normalized.replace('«', '').replace('»', '').replace('"', '')
    return normalized


def parse_prepayment_terms_details(
    terms_text: str
) -> Tuple[float, int, Optional[str], Optional[int]]:
    if not terms_text:
        return 0.0, 0, None, None
    text = terms_text
    lower = text.lower()

    prepay_amount = 0.0
    prepay_days = 0
    prepay_base: Optional[str] = None
    remainder_days: Optional[int] = None

    prepay_load_pattern = re.compile(
        r'(\d[\d\s]*[.,]?\d*)\s*(?:руб|р\.)?'
        r'[^\n]{0,60}?в\s+течение\s+(\d+)\s+рабоч\w*\s+дн\w*'
        r'\s+с\s+даты\s+погрузк',
        re.IGNORECASE
    )
    prepay_label_pattern = re.compile(
        r'(?:сумма\s+предоплаты|предоплата|аванс)\s*[:\-]?\s*'
        r'(\d[\d\s]*[.,]?\d*)',
        re.IGNORECASE
    )
    load_days_pattern = re.compile(
        r'в\s+течение\s+(\d+)\s+рабоч\w*\s+дн\w*\s+с\s+даты\s+погрузк',
        re.IGNORECASE
    )
    remainder_pattern = re.compile(
        r'остаток[^\d]{0,60}?'
        r'(?:не\s+позднее|в\s+течение)?\s*(\d+)\s+рабоч\w*\s+дн\w*'
        r'\s+с\s+даты\s+получен\w*\s+документ',
        re.IGNORECASE
    )
    docs_days_pattern = re.compile(
        r'оплат[аы][^\d]{0,60}?'
        r'(?:не\s+позднее|в\s+течение)?\s*(\d+)\s+рабоч\w*\s+дн\w*'
        r'\s+с\s+даты\s+получен\w*\s+документ',
        re.IGNORECASE
    )

    match = prepay_load_pattern.search(text)
    if match:
        prepay_amount = parse_amount(match.group(1))
        try:
            prepay_days = int(match.group(2))
        except ValueError:
            prepay_days = 0
        prepay_base = "load"
    else:
        match = prepay_label_pattern.search(text)
        if match:
            prepay_amount = parse_amount(match.group(1))
            prepay_base = "load" if "погруз" in lower else None
            match_days = load_days_pattern.search(text)
            if match_days:
                try:
                    prepay_days = int(match_days.group(1))
                except ValueError:
                    prepay_days = 0

    match = remainder_pattern.search(text)
    if match:
        try:
            remainder_days = int(match.group(1))
        except ValueError:
            remainder_days = None
    else:
        match = docs_days_pattern.search(text)
        if match:
            try:
                remainder_days = int(match.group(1))
            except ValueError:
                remainder_days = None

    return prepay_amount, prepay_days, prepay_base, remainder_days


def build_prepayment_terms_text(
    prepay_amount: float,
    prepay_days: int,
    prepay_base: Optional[str],
    remainder_days: Optional[int]
) -> str:
    parts: List[str] = []
    if prepay_amount > 0:
        amount_text = f"{format_money_ru(prepay_amount, 2)} руб."
        if prepay_days > 0:
            days_word = "рабочего дня" if prepay_days == 1 else "рабочих дней"
            base_text = "с даты погрузки" if prepay_base == "load" else ""
            segment = f"{amount_text} в течение {prepay_days} {days_word}"
            if base_text:
                segment += f" {base_text}"
            parts.append(segment)
        else:
            parts.append(f"Сумма предоплаты: {amount_text}")
    if remainder_days:
        days_word = "рабочего дня" if remainder_days == 1 else "рабочих дней"
        remainder_text = (
            f"Остаток не позднее {remainder_days} {days_word} "
            "с даты получения документов, подтверждающих перевозку"
        )
        if parts:
            parts.append(remainder_text)
        else:
            parts.append(
                f"Оплата не позднее {remainder_days} {days_word} "
                "с даты получения документов, подтверждающих перевозку"
            )
    return ". ".join(parts).strip()


def normalize_company_name(value: Optional[str]) -> str:
    if not value:
        return ""
    text = re.sub(r'\s+', ' ', str(value)).strip()
    text = text.replace('«"', '«').replace('"»', '»')
    # Убираем лишние двойные кавычки внутри
    text = text.replace('"', '')
    text = re.sub(r'«{2,}', '«', text)
    text = re.sub(r'»{2,}', '»', text)
    replacements = [
        (r'общество с ограниченной ответственностью', 'ООО'),
        (r'публичное акционерное общество', 'ПАО'),
        (r'закрытое акционерное общество', 'ЗАО'),
        (r'открытое акционерное общество', 'ОАО'),
        (r'акционерное общество', 'АО'),
        (r'индивидуальный предприниматель', 'ИП'),
    ]
    for pattern, abbr in replacements:
        text = re.sub(pattern, abbr, text, flags=re.IGNORECASE)

    def titleize_token(token: str) -> str:
        if not token:
            return token
        if re.fullmatch(r'[A-ZА-ЯЁ]{1,3}', token):
            return token
        if re.fullmatch(r'[A-ZА-ЯЁ]{4,}', token):
            return token.title()
        if re.fullmatch(r'[a-zа-яё]{2,}', token):
            return token.capitalize()
        return token

    def titleize_name(name: str) -> str:
        parts: List[str] = []
        for chunk in re.split(r'(\s+)', name.strip()):
            if not chunk or chunk.isspace():
                parts.append(chunk)
                continue
            if '-' in chunk:
                subparts = [titleize_token(p) for p in chunk.split('-')]
                parts.append('-'.join(subparts))
            else:
                parts.append(titleize_token(chunk))
        return ''.join(parts).strip()

    opf_match = re.match(r'^(ООО|ЗАО|ПАО|ОАО|АО|ИП)\b\s*(.+)?$', text, re.IGNORECASE)
    if opf_match:
        opf = opf_match.group(1).upper()
        rest = (opf_match.group(2) or '').strip()
        if opf == 'ИП':
            name = rest.replace('«', '').replace('»', '').strip()
            name = titleize_name(name) if name else ''
            return f"{opf} {name}".strip()
        if rest:
            quoted = re.search(r'«(.+?)»', rest)
            name = quoted.group(1) if quoted else rest
            name = name.replace('«', '').replace('»', '').strip()
            name = titleize_name(name)
            return f"{opf} «{name}»"
        return opf

    quoted = re.search(r'«(.+?)»', text)
    if quoted:
        name = titleize_name(quoted.group(1))
        return f"«{name}»"

    return titleize_name(text).strip()


def format_company_name_full_upper(value: Optional[str]) -> str:
    if not value:
        return ""
    text = re.sub(r'[«»"]', '', str(value)).strip()
    text = re.sub(r'\s+', ' ', text)
    full_map = {
        "ООО": "ОБЩЕСТВО С ОГРАНИЧЕННОЙ ОТВЕТСТВЕННОСТЬЮ",
        "ЗАО": "ЗАКРЫТОЕ АКЦИОНЕРНОЕ ОБЩЕСТВО",
        "ПАО": "ПУБЛИЧНОЕ АКЦИОНЕРНОЕ ОБЩЕСТВО",
        "ОАО": "ОТКРЫТОЕ АКЦИОНЕРНОЕ ОБЩЕСТВО",
        "АО": "АКЦИОНЕРНОЕ ОБЩЕСТВО",
        "ИП": "ИНДИВИДУАЛЬНЫЙ ПРЕДПРИНИМАТЕЛЬ",
    }
    match = re.match(r'^(ООО|ЗАО|ПАО|ОАО|АО|ИП)\b\s*(.+)?$', text, re.IGNORECASE)
    if match:
        abbr = match.group(1).upper()
        rest = (match.group(2) or "").strip()
        full = full_map.get(abbr, abbr)
        if rest:
            return f"{full} {rest}".upper()
        return full
    for abbr, full in full_map.items():
        text = re.sub(
            r'\b' + re.escape(full) + r'\b',
            full,
            text,
            flags=re.IGNORECASE
        )
    return text.upper().strip()


def get_ogrn_label(name: str, inn_value: str) -> str:
    inn_clean = re.sub(r'[^\d]', '', inn_value or '')
    if (
        'ИП' in name
        or 'Индивидуальный предприниматель' in name
        or len(inn_clean) == 12
    ):
        return 'ОГРНИП'
    return 'ОГРН'


def get_first_list_value(values) -> str:
    if not values:
        return ''
    for val in values:
        if val and str(val).strip():
            return str(val).strip()
    return ''


def join_list_values(values) -> str:
    if not values:
        return ''
    if isinstance(values, str):
        return values.strip()
    cleaned = [str(val).strip() for val in values if str(val).strip()]
    return ', '.join(cleaned)


def normalize_document_item(value: str) -> str:
    normalized = re.sub(r'\s+', ' ', str(value)).strip()
    if normalized.endswith(';'):
        normalized = normalized[:-1].strip()
    return normalized


def format_document_item(value: str) -> str:
    normalized = normalize_document_item(value)
    if not normalized:
        return ''
    if re.match(r'^\s*(\d+[\.\)]|-)\s+', normalized):
        return normalized
    return f"- {normalized}"


def build_documents_list(claim_data: dict) -> str:
    items = []
    for key in [
        'contract_applications',
        'invoice_blocks',
        'upd_blocks',
        'cargo_docs',
        'contracts'
    ]:
        value = claim_data.get(key, '')
        if not value or value == 'Не указано':
            continue
        parts = [part.strip() for part in str(value).split(';') if part.strip()]
        items.extend(parts)
    if not items:
        attachments = claim_data.get('attachments', [])
        if isinstance(attachments, str):
            attachments = [attachments]
        for item in attachments:
            cleaned = normalize_document_item(item)
            if cleaned and cleaned != 'Не указано':
                items.append(cleaned)
    if not items:
        return 'Не указано'
    unique = []
    seen = set()
    for item in items:
        cleaned = normalize_document_item(item)
        if not cleaned or cleaned == 'Не указано':
            continue
        if cleaned not in seen:
            seen.add(cleaned)
            unique.append(cleaned)
    if not unique:
        return 'Не указано'
    if len(unique) == 1:
        return unique[0]
    formatted = [format_document_item(item) for item in unique]
    return '\n'.join([item for item in formatted if item])


def extract_documents_list_structure(text: str) -> Optional[List[Tuple[int, str]]]:
    lines = [line.strip() for line in str(text).splitlines()]

    start_idx = None
    for i, line in enumerate(lines):
        if 'подтверждается документами' in line.lower():
            start_idx = i + 1
            break

    if start_idx is not None:
        block = []
        for line in lines[start_idx:]:
            if not line:
                continue
            if re.match(r'^\d+(\.\d+)?\.', line):
                break
            if re.search(
                r'качество исполнения|отправка оригиналов|расчет процентов|приложен',
                line,
                re.IGNORECASE,
            ):
                break
            block.append(line.strip())
        if block:
            return [(0, item) for item in block]

    for i, line in enumerate(lines):
        if re.search(r'основани[яе].*задолж', line, re.IGNORECASE):
            start_idx = i + 1
            break

    if start_idx is None:
        return None

    def split_document_line(value: str) -> List[str]:
        parts: List[str] = []
        for chunk in re.split(r';', value):
            chunk = chunk.strip()
            if not chunk:
                continue
            subparts = re.split(r'\s*[•\u2022\u00B7]\s*', chunk)
            for subpart in subparts:
                cleaned = subpart.strip()
                if cleaned:
                    parts.append(cleaned)
        return parts or [value]

    block = []
    for line in lines[start_idx:]:
        if not line:
            continue
        if re.search(r'итого\s+задолж', line, re.IGNORECASE):
            break
        if re.match(r'^\d+\.\d+', line):
            break
        if re.search(
            r'качество исполнения|отправка оригиналов|расчет процентов|приложен',
            line,
            re.IGNORECASE,
        ):
            break
        for part in split_document_line(line):
            if part:
                block.append(part)

    if not block:
        return None

    def strip_list_prefix(value: str) -> str:
        stripped = re.sub(r'^\s*\d+[\.\)]\s*', '', value)
        stripped = re.sub(r'^\s*[-–—\u2022\u00B7]\s*', '', stripped)
        return stripped.strip()

    def is_document_line(value: str) -> bool:
        lower = value.lower()
        if '№' in value:
            return True
        if 'комплект сопроводительных документов' in lower:
            return True
        if re.search(r'\d{2}\.\d{2}\.\d{2,4}', value):
            if re.search(
                r'(заявк|сч[её]т|акт|упд|накладн|передаточн|'
                r'разноглас|возврат|транспортн|товарн)',
                lower,
            ):
                return True
        if re.match(
            r'^(заявк|сч[её]т|акт|упд|накладн|договор|транспортн|товарн)',
            lower,
        ):
            return True
        return False

    groups = []
    current = None
    for line in block:
        cleaned = strip_list_prefix(line).rstrip(';').strip()
        if not cleaned:
            continue
        if not is_document_line(cleaned):
            continue
        if re.match(r'^(заявк|договор-?заявк)', cleaned, re.IGNORECASE):
            if current:
                groups.append(current)
            current = {'header': cleaned, 'items': []}
            continue
        if current:
            current['items'].append(cleaned)
        else:
            groups.append({'header': cleaned, 'items': []})

    if current:
        groups.append(current)

    if not groups:
        return None

    structured = []
    for index, group in enumerate(groups, 1):
        header = group['header']
        if group['items']:
            structured.append((0, f"{index}. {header}"))
            for item in group['items']:
                structured.append((1, item))
        else:
            structured.append((0, f"{index}. {header}"))

    return structured


def expand_placeholder_map(replacements: dict) -> dict:
    expanded = {}
    for key, value in replacements.items():
        expanded[key] = value
        if key.startswith('{') and key.endswith('}'):
            name = key[1:-1]
            expanded[f"{{{{{name}}}}}"] = value
    return expanded


def generate_debt_text(claim_data: dict) -> str:
    """
    Форматирует сумму задолженности для вставки в шаблон.

    Args:
        claim_data: Данные о требованиях

    Returns:
        Сумма задолженности в формате строки
    """
    debt_amount = parse_amount_decimal(claim_data.get('debt', '0'))
    rubles_str, kopeks_str = split_rubles_kopeks(debt_amount)
    return f"{rubles_str},{kopeks_str}"


def generate_payment_terms(claim_data: dict) -> str:
    """
    Генерирует текст о порядке оплаты по приоритету:
    1. Если в payment_terms есть текст — возвращает payment_terms
    2. Если есть только дата — возвращает строку с датой
    3. Если есть только дни — возвращает строку с днями
    4. Если ничего нет — стандартный текст
    """
    payment_days = claim_data.get('payment_days')
    payment_due_date = claim_data.get('payment_due_date')
    payment_terms = claim_data.get('payment_terms', '')

    # Если в тексте требования явно есть условия — используем их как есть
    if payment_terms:
        return normalize_payment_terms(payment_terms)
    # Если есть только дата
    if payment_due_date and not payment_days:
        return normalize_payment_terms(
            f"Срок оплаты не позднее {payment_due_date} г."
        )
    # Если есть только дни
    if payment_days and not payment_due_date:
        return normalize_payment_terms(
            f"Оплата производится в течение {payment_days} банковских дней "
            "безналичным расчетом после получения оригиналов документов."
        )
    # Если ничего нет — стандарт
    return normalize_payment_terms(
        "Оплата производится безналичным расчетом после получения "
        "оригиналов документов."
    )


def parse_date_str(value: Optional[str]) -> Optional[datetime]:
    if not value:
        return None
    match = re.search(r'\d{2}\.\d{2}\.\d{4}', str(value))
    if not match:
        return None
    try:
        return datetime.strptime(match.group(0), "%d.%m.%Y")
    except ValueError:
        return None


def parse_ru_text_date(value: Optional[str]) -> Optional[datetime]:
    if not value:
        return None
    match = re.search(
        r'(\d{1,2})\s+([А-Яа-яЁё\.]+)\s+(\d{4})',
        str(value)
    )
    if not match:
        return None
    day_raw, month_raw, year_raw = match.groups()
    month = _parse_ru_month(month_raw)
    if not month:
        return None
    try:
        return datetime(int(year_raw), month, int(day_raw))
    except ValueError:
        return None


def format_money(amount: float, decimals: int = 2) -> str:
    if decimals <= 0:
        return f"{amount:,.0f}".replace(',', ' ')
    return f"{amount:,.{decimals}f}".replace(',', ' ')


def format_money_ru(amount: float, decimals: int = 2) -> str:
    return format_money(amount, decimals).replace('.', ',')


def split_rubles_kopeks(amount: Decimal) -> Tuple[str, str]:
    quantized = amount.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    rubles_value = int(quantized.to_integral_value(rounding=ROUND_DOWN))
    kopeks_value = int((quantized - Decimal(rubles_value)) * 100)
    rubles_str = f"{rubles_value:,}".replace(',', ' ')
    return rubles_str, f"{kopeks_value:02d}"


def format_russian_date(value: Optional[datetime] = None) -> str:
    target = value or datetime.today()
    months = [
        'января', 'февраля', 'марта', 'апреля', 'мая', 'июня',
        'июля', 'августа', 'сентября', 'октября', 'ноября', 'декабря'
    ]
    month_name = months[target.month - 1]
    return f"«{target:%d}» {month_name} {target.year} г."


WORK_CALENDAR_CACHE = os.path.join(
    os.path.dirname(__file__),
    "work_calendar_cache.json"
)


def _parse_ru_month(value: str) -> Optional[int]:
    normalized = str(value).strip().lower().rstrip(".")
    mapping = {
        "января": 1,
        "янв": 1,
        "февраля": 2,
        "фев": 2,
        "марта": 3,
        "мар": 3,
        "апреля": 4,
        "апр": 4,
        "мая": 5,
        "май": 5,
        "июня": 6,
        "июн": 6,
        "июля": 7,
        "июл": 7,
        "августа": 8,
        "авг": 8,
        "сентября": 9,
        "сен": 9,
        "сент": 9,
        "октября": 10,
        "окт": 10,
        "ноября": 11,
        "ноя": 11,
        "нояб": 11,
        "декабря": 12,
        "дек": 12,
    }
    return mapping.get(normalized)


def fetch_work_calendar(year: int) -> Dict[datetime.date, bool]:
    url = (
        f"https://calendar.yoip.ru/work/{year}-proizvodstvennyj-calendar.html"
    )
    response = requests.get(url, timeout=15)
    response.raise_for_status()
    html = response.text

    pattern = re.compile(
        r'title="[^"]*(\d{1,2})&nbsp;([А-Яа-яЁё]+)\s+'
        r'(\d{4})&nbsp;года\.\s*Это\s+'
        r'(выходной|рабочий)\s+день',
        re.IGNORECASE
    )
    calendar: Dict[datetime.date, bool] = {}
    for match in pattern.finditer(html):
        day_raw, month_raw, year_raw, kind = match.groups()
        month = _parse_ru_month(month_raw)
        if not month:
            continue
        try:
            date_obj = datetime(
                int(year_raw),
                month,
                int(day_raw)
            ).date()
        except ValueError:
            continue
        is_working = "рабоч" in kind.lower()
        calendar[date_obj] = is_working
    return calendar


def load_work_calendar(year: int) -> Dict[datetime.date, bool]:
    try:
        if os.path.exists(WORK_CALENDAR_CACHE):
            with open(WORK_CALENDAR_CACHE, "r", encoding="utf-8") as handle:
                payload = json.load(handle)
            if isinstance(payload, dict) and str(year) in payload:
                cached = payload[str(year)]
                if isinstance(cached, dict):
                    result = {}
                    for key, value in cached.items():
                        try:
                            date_obj = datetime.strptime(
                                key, "%Y-%m-%d"
                            ).date()
                        except ValueError:
                            continue
                        result[date_obj] = bool(value)
                    if result:
                        return result
    except Exception as exc:
        logging.warning("Ошибка чтения календаря рабочих дней: %s", exc)

    try:
        calendar = fetch_work_calendar(year)
    except Exception as exc:
        logging.warning(
            "Не удалось загрузить производственный календарь: %s",
            exc
        )
        return {}

    try:
        payload = {}
        if os.path.exists(WORK_CALENDAR_CACHE):
            with open(WORK_CALENDAR_CACHE, "r", encoding="utf-8") as handle:
                payload = json.load(handle)
        payload[str(year)] = {
            date_obj.strftime("%Y-%m-%d"): int(is_working)
            for date_obj, is_working in calendar.items()
        }
        with open(WORK_CALENDAR_CACHE, "w", encoding="utf-8") as handle:
            json.dump(payload, handle, ensure_ascii=False, indent=2)
    except Exception as exc:
        logging.warning(
            "Ошибка сохранения календаря рабочих дней: %s",
            exc
        )

    return calendar


def is_working_day(value: datetime, calendar: Dict[datetime.date, bool]) -> bool:
    if calendar:
        lookup = calendar.get(value.date())
        if lookup is not None:
            return lookup
    return value.weekday() < 5


def add_working_days(
    start_date: datetime,
    days: int,
    calendar: Dict[datetime.date, bool]
) -> datetime:
    if days <= 0:
        return start_date
    current = start_date
    added = 0
    while added < days:
        current += timedelta(days=1)
        if is_working_day(current, calendar):
            added += 1
    return current


def extract_pdf_pages(file_path: str) -> Tuple[List[str], List[int]]:
    """
    Извлекает текст из PDF с использованием гибридного подхода:
    1. pdfplumber (лучше для таблиц)
    2. PyMuPDF как fallback
    3. Vision LLM для страниц с низким качеством (если доступен)
    """
    # Пробуем использовать улучшенный экстрактор
    try:
        from pdf_extractor import extract_pdf_hybrid
        results = extract_pdf_hybrid(file_path, quality_threshold=0.6)
        if results:
            pages = [r.get("text", "") for r in results]
            low_text_pages = [
                r["page_num"] for r in results
                if r.get("text_quality", 1.0) < 0.6 or len(r.get("text", "")) < 40
            ]
            return pages, low_text_pages
    except ImportError:
        logger.debug("pdf_extractor не доступен, используем PyMuPDF")
    except Exception as e:
        logger.warning(f"pdf_extractor error: {e}, fallback to PyMuPDF")

    # Fallback на PyMuPDF
    try:
        import fitz  # PyMuPDF
    except Exception as exc:
        raise RuntimeError(
            f"Не удалось импортировать PyMuPDF для чтения PDF: {exc}"
        ) from exc
    doc = fitz.open(file_path)
    pages = []
    low_text_pages: List[int] = []
    min_chars = 40
    quality_threshold_raw = os.getenv("VISION_TEXT_QUALITY_THRESHOLD", "0.7")
    try:
        quality_threshold = float(quality_threshold_raw)
    except ValueError:
        quality_threshold = 0.7
    for idx in range(doc.page_count):
        page = doc.load_page(idx)
        page_text = page.get_text("text") or ""
        cleaned = page_text.strip()
        if len(cleaned) < min_chars:
            low_text_pages.append(idx + 1)
        else:
            quality = estimate_text_quality(cleaned)
            if quality < quality_threshold:
                low_text_pages.append(idx + 1)
        pages.append(cleaned)
    doc.close()
    return pages, low_text_pages


def extract_pdf_text(file_path: str) -> Tuple[str, List[int]]:
    try:
        pages, low_text_pages = extract_pdf_pages(file_path)
        texts = [
            f"[Страница {idx + 1}]\n{page}"
            for idx, page in enumerate(pages)
        ]
        return "\n\n".join(texts).strip(), low_text_pages
    except Exception as exc:
        raise RuntimeError(
            f"Не удалось прочитать PDF: {exc}"
        ) from exc


def render_pdf_pages(
    file_path: str,
    pages: List[int],
    max_pages: int = 3
) -> List[str]:
    try:
        import fitz  # PyMuPDF
    except Exception:
        return []

    os.makedirs("uploads", exist_ok=True)
    doc = fitz.open(file_path)
    image_paths = []
    for page_number in pages[:max_pages]:
        try:
            page = doc.load_page(page_number - 1)
        except Exception:
            continue
        pix = page.get_pixmap(dpi=150)
        output_name = f"scan_{uuid.uuid4().hex}_p{page_number}.png"
        output_path = os.path.join("uploads", output_name)
        pix.save(output_path)
        image_paths.append(output_path)
    doc.close()
    return image_paths


def apply_vision_ocr_to_pages(
    file_path: str,
    pages: List[str],
    low_text_pages: List[int]
) -> List[int]:
    config = get_vision_config()
    if not config.get("enabled") or not low_text_pages:
        return []

    max_pages = config.get("max_pages") or 0
    target_pages = (
        low_text_pages[:max_pages] if max_pages > 0 else low_text_pages
    )
    image_paths = render_pdf_pages(
        file_path,
        target_pages,
        max_pages=len(target_pages)
    )
    processed: List[int] = []
    for page_number, image_path in zip(target_pages, image_paths):
        vision_text = extract_text_from_image_llm(image_path)
        if vision_text:
            existing = pages[page_number - 1] if page_number - 1 < len(pages) else ""
            combined = vision_text.strip()
            if existing and existing.strip() not in combined:
                combined = f"{existing.strip()}\n{combined}"
            pages[page_number - 1] = combined
            processed.append(page_number)
        try:
            os.remove(image_path)
        except OSError:
            pass
    return processed


def find_postal_candidate_pages(
    pages: List[str],
    limit: int = 2
) -> List[int]:
    """
    Ищет страницы, похожие на почтовые квитанции/отчёты по ключевым словам.
    Возвращает номера страниц (1-based).
    """
    if not pages:
        return []
    hints = (
        "квитанц",
        "рпо",
        "почт",
        "отслежив",
        "идентификатор",
        "отчет",
        "отправлен",
        "письмо",
        "посылк",
    )
    scored: List[Tuple[int, int]] = []
    for idx, page in enumerate(pages, start=1):
        lower = page.lower()
        compact = re.sub(r'\s+', '', lower)
        score = sum(1 for hint in hints if hint in lower)
        if "квитанц" in compact:
            score += 2
        if score > 0:
            scored.append((score, idx))
    scored.sort(key=lambda item: (item[0], item[1]), reverse=True)
    return [idx for _score, idx in scored[: max(limit, 1)]]


def find_keyword_candidate_pages(
    pages: List[str],
    keywords: Tuple[str, ...],
    limit: int = 2
) -> List[int]:
    if not pages or not keywords:
        return []
    scored: List[Tuple[int, int]] = []
    for idx, page in enumerate(pages, start=1):
        lower = page.lower()
        score = sum(1 for key in keywords if key in lower)
        if score > 0:
            scored.append((score, idx))
    scored.sort(key=lambda item: (item[0], item[1]), reverse=True)
    return [idx for _score, idx in scored[: max(limit, 1)]]


def collect_targeted_ocr_pages(
    pages: List[str],
    low_text_pages: Optional[List[int]] = None
) -> List[int]:
    """
    Собирает небольшую выборку страниц для OCR по ключевым словам:
    заявки, накладные, счета, УПД и почтовые квитанции.
    """
    if not pages:
        return []
    low_text_pages = low_text_pages or []
    candidates: List[int] = []
    candidates.extend(find_postal_candidate_pages(pages, limit=2))
    candidates.extend(find_keyword_candidate_pages(
        pages,
        ("заявк", "реквизиты заявки", "экспедитор", "перевозчик"),
        limit=2
    ))
    candidates.extend(find_keyword_candidate_pages(
        pages,
        ("накладн", "транспортн", "товарн", "ттн", "реестр сопровод", "экспедиторск"),
        limit=2
    ))
    candidates.extend(find_keyword_candidate_pages(
        pages,
        ("счет", "счёт", "покупатель", "поставщик"),
        limit=1
    ))
    candidates.extend(find_keyword_candidate_pages(
        pages,
        ("упд", "передаточн"),
        limit=1
    ))
    candidates.extend(find_keyword_candidate_pages(
        pages,
        ("гарантийн", "гарант", "обязуемся оплатить", "гарантия оплаты"),
        limit=1
    ))
    unique = []
    seen = set(low_text_pages)
    for page in candidates:
        if page in seen:
            continue
        seen.add(page)
        unique.append(page)
    return unique


def collect_fallback_doc_pages(
    pages: List[str],
    limit: int = 1
) -> List[int]:
    if not pages:
        return []
    keywords = (
        "договор", "счет", "счёт", "упд", "накладн", "акт",
        "квитанц", "платеж", "платёж", "письмо", "гарант",
        "заявк", "перевоз", "товар", "услуг", "рпо",
    )
    candidates = find_keyword_candidate_pages(pages, keywords, limit=limit)
    if candidates:
        return candidates
    return [1]


def collect_vision_doc_pages(
    pages: List[str],
    processed_low_pages: Optional[List[int]] = None,
    limit: int = 2
) -> List[int]:
    if not pages:
        return []
    processed_low_pages = processed_low_pages or []
    selected: List[int] = []
    seen = set()
    for page_num in processed_low_pages:
        if page_num not in seen:
            seen.add(page_num)
            selected.append(page_num)
    targeted = collect_targeted_ocr_pages(pages, [])
    for page_num in targeted:
        if page_num not in seen:
            seen.add(page_num)
            selected.append(page_num)
    if not selected:
        selected.extend(collect_fallback_doc_pages(pages, limit=limit))
    return selected[: max(1, limit)]


def _format_date_value(value: Any) -> str:
    if isinstance(value, datetime):
        return value.strftime("%d.%m.%Y")
    if not value:
        return ""
    value_str = str(value).strip()
    parsed = parse_date_str(value_str)
    if not parsed:
        parsed = parse_ru_text_date(value_str)
    if parsed:
        return parsed.strftime("%d.%m.%Y")
    return value_str


def build_vision_hint_lines(data: Dict[str, Any]) -> List[str]:
    if not data:
        return []

    lines: List[str] = []

    doc_type = str(
        data.get("document_type")
        or data.get("document_type_raw")
        or ""
    ).strip()
    doc_type_lower = doc_type.lower()

    if doc_type:
        lines.append(f"Документ: {doc_type}")

    def pick_first(*values: Any) -> str:
        for val in values:
            if val:
                return str(val).strip()
        return ""

    document_number = pick_first(
        data.get("document_number"),
        data.get("invoice_number"),
        data.get("upd_number"),
        data.get("contract_number"),
        data.get("payment_order_number"),
    )
    document_date = pick_first(
        data.get("document_date"),
        data.get("invoice_date"),
        data.get("upd_date"),
        data.get("contract_date"),
        data.get("payment_order_date"),
    )
    document_date = _format_date_value(document_date)

    if any(token in doc_type_lower for token in ("заявк", "application")):
        app_number = pick_first(
            data.get("application_number"),
            data.get("document_number"),
        )
        app_date = _format_date_value(
            pick_first(data.get("application_date"), data.get("document_date"))
        )
        if app_number and app_date:
            lines.append(f"Заявка № {app_number} от {app_date}")

    if any(token in doc_type_lower for token in ("счет", "счёт", "invoice")):
        number = pick_first(data.get("invoice_number"), document_number)
        if number and document_date:
            lines.append(f"Счет № {number} от {document_date}")

    if "упд" in doc_type_lower or "передаточн" in doc_type_lower:
        number = pick_first(data.get("upd_number"), document_number)
        if number and document_date:
            lines.append(f"УПД № {number} от {document_date}")

    if any(token in doc_type_lower for token in ("накладн", "ттн", "транспорт", "transport")):
        if document_number and document_date:
            lines.append(f"Транспортная накладная № {document_number} от {document_date}")

    if any(token in doc_type_lower for token in ("акт", "выполненных работ", "оказанных услуг")):
        if document_number and document_date:
            lines.append(f"Акт № {document_number} от {document_date}")

    if "договор" in doc_type_lower:
        if document_number and document_date:
            lines.append(f"Договор № {document_number} от {document_date}")

    if any(token in doc_type_lower for token in ("платеж", "платёж", "payment_order")):
        if document_number and document_date:
            lines.append(f"Платежное поручение № {document_number} от {document_date}")

    if "гарант" in doc_type_lower or "письмо" in doc_type_lower:
        if document_date:
            lines.append(f"Гарантийное письмо от {document_date}")
        else:
            lines.append("Гарантийное письмо")

    amount = data.get("amount")
    if amount:
        lines.append(f"Сумма: {amount} руб.")

    payment_terms = data.get("payment_terms")
    if payment_terms:
        lines.append(f"Условия оплаты: {payment_terms}")

    track_number = (
        data.get("track_number")
        or data.get("tracking_number")
        or data.get("rpo")
    )
    if track_number:
        normalized = normalize_tracking_number(str(track_number))
        if normalized:
            lines.append(f"Идентификатор: {normalized}")
            lines.append(f"РПО: {normalized}")

    sender_name = data.get("sender_name")
    sender_inn = data.get("sender_inn")
    if sender_name:
        if sender_inn:
            lines.append(f"Грузоотправитель: {sender_name}, ИНН {sender_inn}")
        else:
            lines.append(f"Грузоотправитель: {sender_name}")

    receiver_name = data.get("receiver_name")
    receiver_inn = data.get("receiver_inn")
    if receiver_name:
        if receiver_inn:
            lines.append(f"Грузополучатель: {receiver_name}, ИНН {receiver_inn}")
        else:
            lines.append(f"Грузополучатель: {receiver_name}")

    carrier_name = data.get("carrier_name")
    carrier_inn = data.get("carrier_inn")
    if carrier_name:
        if carrier_inn:
            lines.append(f"Перевозчик: {carrier_name}, ИНН {carrier_inn}")
        else:
            lines.append(f"Перевозчик: {carrier_name}")

    return lines


def apply_vision_document_extraction(
    file_path: str,
    pages: List[str],
    page_numbers: List[int]
) -> List[int]:
    processed: List[int] = []
    for page_num in page_numbers:
        page_idx = page_num - 1
        if page_idx < 0 or page_idx >= len(pages):
            continue
        vision_data = extract_claim_document_with_vision(
            file_path,
            page_idx
        )
        if not vision_data:
            continue
        hint_lines = build_vision_hint_lines(vision_data)
        full_text = str(vision_data.get("full_text") or "").strip()
        combined_parts = []
        if hint_lines:
            combined_parts.append("\n".join(hint_lines))
        if full_text:
            combined_parts.append(full_text)
        combined_text = "\n".join(combined_parts).strip()
        if combined_text:
            existing = pages[page_idx].strip()
            if combined_text not in existing:
                pages[page_idx] = (existing + "\n" + combined_text).strip()
            processed.append(page_num)
    return processed


def extract_payment_terms_from_text(
    text: str,
    allow_llm: bool = True
) -> Tuple[Optional[str], Optional[int]]:
    if not text:
        return None, None
    if re.search(r'предоплат|аванс|остаток|погрузк', text, re.IGNORECASE):
        prepay_amount, prepay_days, prepay_base, remainder_days = parse_prepayment_terms_details(text)
        if prepay_amount > 0 or remainder_days:
            terms_text = build_prepayment_terms_text(
                prepay_amount,
                prepay_days,
                prepay_base,
                remainder_days
            )
            if terms_text:
                return normalize_payment_terms(terms_text), None
    match = re.search(
        r'(Оплата[\s\S]{0,200}?(\d[\d\s]{0,4})\s*'
        r'(?:\([^)]*\)\s*)?(?:рабоч|банков|календар)[^\n]{0,120})',
        text,
        re.IGNORECASE
    )
    if not match:
        if not re.search(r'оплат|срок|расчет|рабоч|банков|остаток|не позднее', text, re.IGNORECASE):
            return None, None
        lines = text.splitlines()
        extracted_lines = []
        for idx, line in enumerate(lines):
            if re.search(r'оплат|срок|расчет|рабоч|банков|остаток|не позднее', line, re.IGNORECASE):
                start = max(0, idx - 1)
                end = min(len(lines), idx + 3)
                extracted_lines.extend(lines[start:end])
        cleaned_lines = []
        seen = set()
        for line in extracted_lines:
            stripped = line.strip()
            if not stripped or stripped in seen:
                continue
            seen.add(stripped)
            cleaned_lines.append(stripped)
        if not cleaned_lines:
            return None, None
        context = "\n".join(cleaned_lines)
        if len(context) > 4000:
            context = context[:4000]
        if allow_llm:
            llm_terms, llm_days = extract_payment_terms_llm(context)
            if llm_terms:
                return normalize_payment_terms(llm_terms), llm_days
            if llm_days:
                return None, llm_days
        return None, None
    terms = normalize_payment_terms(match.group(1))
    try:
        days = int(re.sub(r"[^\d]", "", match.group(2)))
    except ValueError:
        days = None
    return terms, days


PLATE_PATTERN = re.compile(
    r'[A-ZА-Я]\s*\d{3}\s*[A-ZА-Я]{2}\s*\d{2,3}',
    re.IGNORECASE
)


def _build_label_pattern(label: str) -> str:
    parts = [re.escape(part) for part in label.split()]
    return r'\s+'.join(parts)


def clean_text_value(value: Optional[str]) -> str:
    if not value:
        return ""
    cleaned = re.sub(r'\s+', ' ', str(value))
    return cleaned.strip(" \t\r\n-–—:;")


def extract_line_value(text: str, labels: List[str]) -> Optional[str]:
    for label in labels:
        label_pattern = _build_label_pattern(label)
        boundary = (
            r'\b' if re.search(r'[A-Za-zА-Яа-я0-9]', label) else ''
        )
        pattern_start = rf'(?m)^\s*{boundary}{label_pattern}{boundary}\s*[:\-]?\s*([^\n]+)$'
        match = re.search(pattern_start, text, re.IGNORECASE)
        if match:
            value = clean_text_value(match.group(1))
            if value:
                return value
        pattern_same_line = rf'{boundary}{label_pattern}{boundary}\s*[:\-]?\s*([^\n]+)'
        match = re.search(pattern_same_line, text, re.IGNORECASE)
        if match:
            value = clean_text_value(match.group(1))
            if value:
                return value
        pattern_next_line = rf'{boundary}{label_pattern}{boundary}\s*[:\-]?\s*\n\s*([^\n]+)'
        match = re.search(pattern_next_line, text, re.IGNORECASE)
        if match:
            value = clean_text_value(match.group(1))
            if value:
                return value
    return None


def extract_date_near_labels(text: str, labels: List[str]) -> Optional[datetime]:
    for label in labels:
        label_pattern = _build_label_pattern(label)
        boundary = (
            r'\b' if re.search(r'[A-Za-zА-Яа-я0-9]', label) else ''
        )
        for line in text.splitlines():
            if re.search(rf'{boundary}{label_pattern}{boundary}', line, re.IGNORECASE):
                date_value = parse_date_str(line)
                if date_value:
                    return date_value
        match = re.search(
            rf'{boundary}{label_pattern}{boundary}(.{{0,80}})',
            text,
            re.IGNORECASE
        )
        if match:
            date_value = parse_date_str(match.group(1))
            if date_value:
                return date_value
    return None


def extract_plate_near_labels(text: str, labels: List[str]) -> Optional[str]:
    for label in labels:
        label_pattern = _build_label_pattern(label)
        match = re.search(rf'{label_pattern}[^\n]*', text, re.IGNORECASE)
        if not match:
            continue
        snippet = match.group(0)
        plate_match = PLATE_PATTERN.search(snippet)
        if not plate_match:
            tail = text[match.end():]
            next_line = tail.splitlines()[0] if tail else ""
            plate_match = PLATE_PATTERN.search(next_line)
        if plate_match:
            return clean_text_value(plate_match.group(0))
    return None


def extract_application_number_from_text(text: str) -> Optional[str]:
    matches = re.findall(
        r'Заявк[аеи]?\s*№\s*([A-Za-zА-Яа-я0-9/\\-]+)',
        text,
        re.IGNORECASE
    )
    if len(matches) == 1:
        return matches[0].strip()
    return None


def normalize_application_number(value: Optional[str]) -> str:
    if not value:
        return ""
    return re.sub(r'[^A-Za-zА-Яа-я0-9]', '', str(value)).upper()


def extract_transport_details(text: str, allow_llm: bool = True) -> Dict[str, Any]:
    details: Dict[str, Any] = {}

    application_number = extract_application_number_from_text(text)
    if application_number:
        details["application_number"] = application_number

    # Расширенные паттерны для дат (включая табличные форматы)
    load_date = extract_date_near_labels(
        text,
        [
            "дата погрузки",
            "дата загрузки",
            "погрузка",
            "загрузка",
            "время начала",  # Из таблицы ТН
            "прибытия под погрузку",  # Из ТН
            "фактические дата и время прибытия",  # Из ТН
        ],
    )
    unload_date = extract_date_near_labels(
        text,
        [
            "дата разгрузки",
            "дата выгрузки",
            "разгрузка",
            "выгрузка",
            "время окончания",  # Из таблицы ТН
            "доставки груза",  # Из ТН
            "сдачи груза",  # Из ТН
        ],
    )

    # Парсинг табличного формата с переносами строк
    # "Погрузка\n15.11.2025 01:00" или "Погрузка | 15.11.2025"
    if not load_date or not unload_date:
        table_patterns = [
            # "Погрузка\n15.11.2025 01:00" - с переносом
            re.compile(
                r'(погрузка|разгрузка)\s*\n\s*(\d{2}[./]\d{2}[./]\d{4})',
                re.IGNORECASE
            ),
            # "Погрузка | 15.11.2025" - через разделитель
            re.compile(
                r'(погрузка|разгрузка)\s*[\|\t]\s*(\d{2}[./]\d{2}[./]\d{4})',
                re.IGNORECASE
            ),
            # "Тип: Погрузка ... 15.11.2025" - дата рядом
            re.compile(
                r'(погрузка|разгрузка)[^\d]{0,50}(\d{2}[./]\d{2}[./]\d{4})',
                re.IGNORECASE
            ),
        ]
        for pattern in table_patterns:
            for match in pattern.finditer(text):
                op_type = match.group(1).lower()
                date_str = match.group(2).replace('/', '.')
                parsed = parse_date_str(date_str)
                if parsed:
                    if "погрузк" in op_type and not load_date:
                        load_date = parsed
                    elif "разгрузк" in op_type and not unload_date:
                        unload_date = parsed
    if load_date:
        details["load_date"] = load_date
    if unload_date:
        details["unload_date"] = unload_date

    driver = extract_line_value(
        text,
        [
            "фио водителя",
            "ф.и.о. водителя",
            "водитель",
            "водитель экспедитор",
        ],
    )
    if driver:
        driver = re.split(r'[;,]|тел|\bинн\b', driver, 1, re.IGNORECASE)[0]
        driver = driver.strip()
        # Фильтруем невалидные значения (обрезанные слова, скобки и т.п.)
        if driver and len(driver) > 5 and not driver.startswith('('):
            details["driver_name"] = driver

    # Табличный формат: "Ф.И.О. водителя, ИНН\nЕгиазарян Вазген..."
    # Или: "Ф.И.О. водителя, ИНН Егиазарян Вазген..." (pdfplumber на одной строке)
    if not details.get("driver_name"):
        driver_patterns_table = [
            # С переносом строки
            re.compile(
                r'Ф\.?И\.?О\.?\s*водител[яь][^\n]*\n+([А-ЯЁ][а-яё]+\s+'
                r'[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+)',
                re.IGNORECASE
            ),
            # На одной строке (pdfplumber): "Ф.И.О. водителя, ИНН Егиазарян Вазген Шагенович"
            re.compile(
                r'Ф\.?И\.?О\.?\s*водител[яь][,\s]+(?:ИНН\s+)?'
                r'([А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+)',
                re.IGNORECASE
            ),
        ]
        for dp in driver_patterns_table:
            dm = dp.search(text)
            if dm:
                details["driver_name"] = dm.group(1).strip()
                break

    # Формат ТН: ФИО водителя указано ПЕРЕД строкой
    # "(реквизиты, позволяющие идентифицировать водителя(-ей)"
    # Пример: "БЕЗРОДНЫЙ АЛЕКСЕЙ НИКОЛАЕВИЧ, ИНН 110104211012\n(реквизиты...водителя"
    # Или: "...110104211012 (реквизиты...водителя" (pdfplumber без переноса)
    if not details.get("driver_name"):
        driver_patterns = [
            # ФИО ПЕРЕД "(реквизиты, позволяющие идентифицировать водителя" - с переносом
            re.compile(
                r'([А-ЯЁ][А-ЯЁа-яё]+\s+[А-ЯЁ][А-ЯЁа-яё]+\s+[А-ЯЁ][А-ЯЁа-яё]+)'
                r'[,\s]+(?:ИНН\s*\d+)?[^\n]*\n'
                r'\(реквизиты[^\)]*идентифицировать\s+(?:Перевозчика|водител)',
                re.IGNORECASE | re.MULTILINE
            ),
            # ФИО ПЕРЕД "(реквизиты...водителя" - БЕЗ переноса (pdfplumber)
            re.compile(
                r'([А-ЯЁ][А-ЯЁа-яё]+\s+[А-ЯЁ][А-ЯЁа-яё]+\s+[А-ЯЁ][А-ЯЁа-яё]+)'
                r'[,\s]+(?:ИНН\s*\d+)?[^\(]*'
                r'\(реквизиты[^\)]*идентифицировать\s+водител',
                re.IGNORECASE
            ),
            # ФИО после Перевозчика, на строке перед водителем
            re.compile(
                r'Перевозчик[а-я]*\)\s*[\n\s]+'
                r'([А-ЯЁ][А-ЯЁа-яё]+\s+[А-ЯЁ][А-ЯЁа-яё]+\s+[А-ЯЁ][А-ЯЁа-яё]+)',
                re.IGNORECASE | re.MULTILINE
            ),
        ]
        for dp in driver_patterns:
            dm = dp.search(text)
            if dm:
                driver_candidate = dm.group(1).strip()
                # Убираем ИНН и прочее после ФИО
                driver_candidate = re.split(
                    r'[;,]|тел|\bинн\b|\bогрн', driver_candidate, 1, re.IGNORECASE
                )[0].strip()
                if driver_candidate and len(driver_candidate) > 5:
                    details["driver_name"] = driver_candidate
                    break

    vehicle = extract_plate_near_labels(
        text,
        [
            "гос. номер",
            "гос номер",
            "г/н",
            "транспортное средство",
            "автомобиль",
            "тягач",
            "тс",
            "марка и номер тс",  # Из таблицы заявки
        ],
    )
    if vehicle:
        details["vehicle_plate"] = vehicle

    # Формат ТН: "(регистрационный номер транспортного средства)"
    # После этой строки идет номер: "Е 015 ВК 11"
    if not details.get("vehicle_plate"):
        plate_patterns = [
            # После "регистрационный номер транспортного средства"
            re.compile(
                r'регистрационный\s+номер\s+транспортного\s+средства\s*\)?\s*\n?'
                r'([А-ЯЁA-Z]\s*\d{3}\s*[А-ЯЁA-Z]{2}\s*\d{2,3})',
                re.IGNORECASE
            ),
            # Госномер в формате "Е015ВК11" или "Е 015 ВК 11"
            re.compile(
                r'(?:7\.\s*Транспортное\s+средство|номер\s+тс)[^\n]*\n'
                r'[^\n]*\n\s*([А-ЯЁA-Z]\s*\d{3}\s*[А-ЯЁA-Z]{2}\s*\d{2,3})',
                re.IGNORECASE
            ),
        ]
        for pp in plate_patterns:
            pm = pp.search(text)
            if pm:
                plate_candidate = pm.group(1).strip()
                # Нормализуем номер (убираем лишние пробелы)
                plate_candidate = re.sub(r'\s+', ' ', plate_candidate)
                if plate_candidate:
                    details["vehicle_plate"] = plate_candidate
                    break

    trailer = extract_plate_near_labels(
        text,
        [
            "прицеп",
            "полуприцеп",
            "марка и номер прицепа",  # Из таблицы заявки
        ],
    )
    if trailer:
        details["trailer_plate"] = trailer

    # Табличный формат: "Марка и номер прицепа\nКОГОЕ ЕТ 9226 77"
    if not details.get("trailer_plate"):
        trailer_match = re.search(
            r'(?:марка\s+и\s+номер\s+прицеп|полуприцеп)[а-я]*\s*\n+'
            r'([А-ЯЁA-Z]{2,}\s*[А-ЯЁA-Z]{0,2}\s*\d{4}\s*\d{2,3})',
            text,
            re.IGNORECASE
        )
        if trailer_match:
            details["trailer_plate"] = trailer_match.group(1).strip()

    sender = extract_line_value(text, ["грузоотправитель", "отправитель"])
    if sender:
        details["sender_name"] = sender

    receiver = extract_line_value(text, ["грузополучатель", "получатель"])
    if receiver:
        details["receiver_name"] = receiver

    load_address = extract_line_value(
        text,
        [
            "место погрузки",
            "пункт погрузки",
            "адрес погрузки",
            "место загрузки",
            "пункт загрузки",
            "адрес загрузки",
        ],
    )
    if load_address:
        details["load_address"] = load_address

    unload_address = extract_line_value(
        text,
        [
            "место разгрузки",
            "пункт разгрузки",
            "адрес разгрузки",
            "место выгрузки",
            "пункт выгрузки",
            "адрес выгрузки",
        ],
    )
    if unload_address:
        details["unload_address"] = unload_address

    route = extract_line_value(text, ["маршрут", "направление"])
    if route and ("-" in route or "—" in route):
        parts = re.split(r'\s*[-—]\s*', route)
        if len(parts) >= 2:
            if "load_address" not in details:
                details["load_address"] = parts[0].strip()
            if "unload_address" not in details:
                details["unload_address"] = parts[-1].strip()

    # LLM-fallback для незаполненных полей
    transport_fields = [
        "driver_name", "vehicle_plate", "trailer_plate",
        "load_date", "unload_date",
        "load_address", "unload_address",
        "sender_name", "receiver_name"
    ]
    missing_count = sum(1 for f in transport_fields if not details.get(f))
    if allow_llm and missing_count >= 3:  # Если много полей не найдено - пробуем LLM
        try:
            details = extract_transport_details_llm(text, details)
        except Exception as e:
            logger.warning(f"LLM transport fallback error: {e}")

    for date_field in ("load_date", "unload_date"):
        value = details.get(date_field)
        if isinstance(value, str):
            parsed = parse_date_str(value)
            if parsed:
                details[date_field] = parsed

    return details


def _number_like_in_text(number: str, text: str) -> bool:
    if not number or not text:
        return False
    if number in text:
        return True
    compact_number = re.sub(r"\s+", "", str(number))
    if not compact_number:
        return False
    compact_text = re.sub(r"\s+", "", text)
    if compact_number in compact_text:
        return True
    pattern = r"\s*".join(re.escape(ch) for ch in compact_number)
    return bool(re.search(pattern, text))


def extract_application_amount(text: str) -> Optional[float]:
    if not text:
        return None

    def find_amount(line: str) -> Optional[float]:
        candidates = re.findall(r"\d[\d\s]*[.,]?\d{0,2}", line)
        for candidate in candidates:
            value = parse_amount(candidate)
            if value >= 1000:
                return value
        return None

    lines = [line.strip() for line in text.splitlines() if line.strip()]
    for idx, line in enumerate(lines):
        lower = line.lower()
        if "стоим" in lower and any(
            token in lower for token in ("услуг", "перевоз", "фрахт", "тариф")
        ):
            value = find_amount(line)
            if not value and idx + 1 < len(lines):
                value = find_amount(lines[idx + 1])
            if value:
                return value
        if "цена" in lower and any(
            token in lower for token in ("услуг", "перевоз", "фрахт")
        ):
            value = find_amount(line)
            if not value and idx + 1 < len(lines):
                value = find_amount(lines[idx + 1])
            if value:
                return value

    pattern = re.compile(
        r'(?:стоимость|цена)[^\n]{0,80}\n?\s*([\d\s.,]{3,})',
        re.IGNORECASE
    )
    for match in pattern.finditer(text):
        value = parse_amount(match.group(1))
        if value >= 1000:
            return value

    return None


def extract_application_payment_terms(
    pages: List[str],
    applications: List[Dict[str, Any]],
    allow_llm: bool = True
) -> Dict[str, Dict[str, Any]]:
    terms_map: Dict[str, Dict[str, Any]] = {}
    if not pages or not applications:
        return terms_map
    for app in applications:
        number = app.get("number")
        label = app.get("label")
        if not number or not label:
            continue
        for page in pages:
            if not _number_like_in_text(number, page):
                continue
            terms, days = extract_payment_terms_from_text(
                page,
                allow_llm=allow_llm
            )
            if terms or days:
                terms_map[label] = {
                    "terms": terms,
                    "days": days,
                }
                break
    return terms_map


def extract_applications_from_pages(
    pages: List[str],
    allow_llm: bool = True
) -> List[Dict[str, Any]]:
    """
    Извлекает заявки из страниц PDF.

    ВАЖНО: Заявка может занимать несколько страниц (обычно 2).
    Заголовок на первой странице, данные о транспорте - на второй.
    Функция группирует страницы по номеру заявки и объединяет данные.
    """
    applications = []
    seen = set()

    # Паттерны для поиска заголовка заявки
    pattern = re.compile(
        r'Реквизиты\s+заявки\s*([A-Za-zА-Яа-я0-9/\\-]+)\s*от\s*'
        r'(\d{2}\.\d{2}\.\d{4})',
        re.IGNORECASE
    )
    fallback_pattern = re.compile(
        r'Заявка\s*№\s*([A-Za-zА-Яа-я0-9/\\-]+)\s*от\s*'
        r'(\d{2}\.\d{2}\.\d{4})',
        re.IGNORECASE
    )

    # Шаг 1: Найти все заявки и их страницы
    app_pages: Dict[Tuple[str, str], List[str]] = {}

    detail_markers = (
        "заявк",
        "водител",
        "гос",
        "прицеп",
        "полуприцеп",
        "маршрут",
        "погруз",
        "разгруз",
        "грузоотправ",
        "грузополуч",
        "транспортное средство",
        "тягач",
        "тс",
    )
    exclude_markers = (
        "счет",
        "упд",
        "универсальн",
        "акт",
        "накладн",
        "отчет об отслеживании",
        "платеж",
        "договор оказания юридических услуг",
    )

    for i, page in enumerate(pages):
        # Ищем номер заявки на странице
        match = pattern.search(page) or fallback_pattern.search(page)
        if match:
            number, date_str = match.groups()
            key = (number.strip(), date_str)
            if key not in app_pages:
                app_pages[key] = []
            app_pages[key].append(page)
        else:
            # Страница без заголовка - возможно продолжение
            # Присоединяем к последней найденной заявке
            if app_pages:
                last_key = list(app_pages.keys())[-1]
                lowered = page.lower()
                has_detail = any(marker in lowered for marker in detail_markers)
                has_excluded = any(marker in lowered for marker in exclude_markers)
                # Проверяем что это действительно продолжение:
                # либо есть номер заявки, либо страница похожа на приложение
                if last_key[0] in page or (has_detail and not has_excluded):
                    app_pages[last_key].append(page)

    # Шаг 2: Извлечь данные из объединённых страниц
    for (number, date_str), page_list in app_pages.items():
        if (number, date_str) in seen:
            continue
        seen.add((number, date_str))

        # Объединяем все страницы заявки
        combined_text = "\n".join(page_list)
        page_details = extract_transport_details(
            combined_text,
            allow_llm=allow_llm
        )

        app = {
            "number": number,
            "date": parse_date_str(date_str),
            "label": f"Заявка № {number} от {date_str}",
        }
        amount_value = extract_application_amount(combined_text)
        if amount_value:
            app["amount"] = amount_value
        if page_details:
            app.update(page_details)
        applications.append(app)

    return applications


def extract_invoices_from_pages(pages: List[str]) -> List[Dict[str, Any]]:
    invoices = []
    seen = set()
    pattern = re.compile(
        r'Сч[её]т(?:\s+на\s+оплату)?\s*(?:№|No|Nо|N)\s*([A-Za-zА-Яа-я0-9/\\-]+)'
        r'\s*от\s*([^\n]+)',
        re.IGNORECASE
    )
    amount_patterns = [
        re.compile(
            r'Всего\s+к\s+оплате:\s*([\s\S]{0,60})',
            re.IGNORECASE
        ),
        re.compile(
            r'Итого:\s*([\s\S]{0,60})',
            re.IGNORECASE
        ),
        re.compile(
            r'на\s+сумм[ау]\s*([\s\S]{0,60})',
            re.IGNORECASE
        ),
        re.compile(
            r'сумм[ау]\s*([\s\S]{0,60})\s*руб',
            re.IGNORECASE
        ),
    ]
    for page in pages:
        match = pattern.search(page)
        if not match:
            continue
        number, date_raw = match.groups()
        date_candidate = parse_date_str(date_raw)
        if not date_candidate:
            date_candidate = parse_ru_text_date(date_raw)
        date_str = date_candidate.strftime("%d.%m.%Y") if date_candidate else ""
        key = (number, date_str)
        if key in seen:
            continue
        seen.add(key)
        amount_value = None
        for amount_pattern in amount_patterns:
            amount_matches = amount_pattern.findall(page)
            if not amount_matches:
                continue
            for match in reversed(amount_matches):
                amount_value = extract_last_amount_from_text(match)
                if amount_value:
                    break
            if amount_value:
                break
        invoices.append({
            "number": number,
            "date": parse_date_str(date_str),
            "label": f"Счет № {number} от {date_str}",
            "amount": amount_value,
        })
    return invoices


def extract_upd_from_pages(pages: List[str]) -> List[Dict[str, Any]]:
    upd_docs = []
    seen = set()
    pattern = re.compile(
        r'(?:УПД|универсальн[^\n]*передаточн[^\n]*документ)'
        r'[^\n]*?№\s*([A-Za-zА-Яа-я0-9/\\-]+)'
        r'[^\n]*?от\s*([^\n]+)',
        re.IGNORECASE
    )
    amount_patterns = [
        re.compile(
            r'(?:Итого\s+к\s+оплате|Всего\s+к\s+оплате)\s*'
            r'([\s\S]{0,60})',
            re.IGNORECASE
        ),
        re.compile(
            r'на\s+сумм[ау]\s*([\s\S]{0,60})',
            re.IGNORECASE
        ),
        re.compile(
            r'сумм[ау]\s*([\s\S]{0,60})\s*руб',
            re.IGNORECASE
        ),
    ]
    for page in pages:
        if "упд" not in page.lower() and "передаточн" not in page.lower():
            continue
        for match in pattern.finditer(page):
            number, date_raw = match.groups()
            date_candidate = parse_date_str(date_raw)
            if not date_candidate:
                date_candidate = parse_ru_text_date(date_raw)
            date_str = date_candidate.strftime("%d.%m.%Y") if date_candidate else ""
            key = (number, date_str)
            if key in seen:
                continue
            seen.add(key)
            amount_value = None
            for amount_pattern in amount_patterns:
                amount_matches = amount_pattern.findall(page)
                if not amount_matches:
                    continue
                for match in reversed(amount_matches):
                    amount_value = extract_last_amount_from_text(match)
                    if amount_value:
                        break
                if amount_value:
                    break
            upd_docs.append({
                "number": number.strip(),
                "date": date_candidate,
                "label": f"УПД № {number.strip()} от {date_str}",
                "amount": amount_value,
            })
    return upd_docs


def extract_legal_docs_from_pages(pages: List[str]) -> Dict[str, Any]:
    result: Dict[str, Any] = {}
    for page in pages:
        lowered = page.lower()
        lines = [line.strip() for line in page.splitlines() if line.strip()]
        has_legal_context = any(
            token in lowered
            for token in ("юридичес", "представител", "юрист", "договор оказания юрид")
        )

        if "договор" in lowered and "юрид" in lowered:
            for idx, line in enumerate(lines):
                lower_line = line.lower()
                if "договор" not in lower_line or "юрид" not in lower_line:
                    continue
                num_match = re.search(
                    r'№\s*([A-Za-zА-Яа-я0-9/\\-]+)',
                    line
                )
                if num_match and not result.get("legal_contract_number"):
                    result["legal_contract_number"] = num_match.group(1).strip()
                if not result.get("legal_contract_date"):
                    snippet = line
                    if idx + 1 < len(lines):
                        snippet = snippet + " " + lines[idx + 1]
                    if idx + 2 < len(lines):
                        snippet = snippet + " " + lines[idx + 2]
                    date_match = re.search(
                        r'от\s+(\d{1,2}[./]\d{1,2}[./]\d{4}|\d{1,2}\s+[А-Яа-яЁё\.]+\s+\d{4})',
                        snippet,
                        re.IGNORECASE
                    )
                    if date_match:
                        date_str = date_match.group(1)
                        date_val = parse_date_str(date_str.replace('/', '.')) or parse_ru_text_date(date_str)
                        if date_val:
                            result["legal_contract_date"] = date_val.strftime("%d.%m.%Y")
                break

        if re.search(r'сч[её]т\s+на\s+оплат', lowered) and "юрид" in lowered:
            for idx, line in enumerate(lines):
                lower_line = line.lower()
                if "счет" not in lower_line or "юрид" not in lower_line:
                    continue
                inv_match = re.search(
                    r'счет[^\n]{0,40}?(?:№|N[оo0])\s*([A-Za-zА-Яа-я0-9/\\-]+)',
                    line,
                    re.IGNORECASE
                )
                inv_date_match = re.search(
                    r'от\s+(\d{1,2}[./]\d{1,2}[./]\d{4}|\d{1,2}\s+[А-Яа-яЁё\.]+\s+\d{4})',
                    line,
                    re.IGNORECASE
                )
                if inv_match and not result.get("legal_invoice_number"):
                    result["legal_invoice_number"] = inv_match.group(1).strip()
                if inv_date_match and not result.get("legal_invoice_date"):
                    date_str = inv_date_match.group(1)
                    date_val = parse_date_str(date_str.replace('/', '.')) or parse_ru_text_date(date_str)
                    if date_val:
                        result["legal_invoice_date"] = date_val.strftime("%d.%m.%Y")
                break

        if any(token in lowered for token in ("платежн", "платёжн", "п/п", "поручен")):
            if has_legal_context:
                for line in lines:
                    lower_line = line.lower()
                    if "платеж" not in lower_line and "платёж" not in lower_line and "п/п" not in lower_line:
                        continue
                    pay_match = re.search(
                        r'(?:плат[её]жн\w*\s+поручен\w*|п/п)\s*№?\s*([A-Za-zА-Яа-я0-9/\\-]+)',
                        line,
                        re.IGNORECASE
                    )
                    if pay_match and not result.get("legal_payment_number"):
                        result["legal_payment_number"] = pay_match.group(1).strip()
                    if not result.get("legal_payment_date"):
                        date_match = re.search(
                            r'(\d{1,2}[./]\d{1,2}[./]\d{4})',
                            line
                        )
                        if date_match:
                            date_val = parse_date_str(date_match.group(1).replace('/', '.'))
                            if date_val:
                                result["legal_payment_date"] = date_val.strftime("%d.%m.%Y")
                    break
                if not result.get("legal_payment_date"):
                    first_date = extract_first_date(page)
                    if first_date:
                        result["legal_payment_date"] = first_date.strftime("%d.%m.%Y")
                amount_match = re.search(
                    r'(?:сумм[ау]|на\s+сумму)\s*([\d\s]+[.,]?\d*)\s*(?:руб|р\.)',
                    page,
                    re.IGNORECASE
                )
                if amount_match and not result.get("legal_fees"):
                    amount_val = parse_amount(amount_match.group(1))
                    if amount_val > 0:
                        result["legal_fees"] = amount_val

        if (
            has_legal_context
            and not result.get("legal_fees")
            and "сумм" in lowered
        ):
            amount_match = re.search(
                r'(?:сумм[ау]|на\s+сумму)\s*([\d\s]+[.,]?\d*)\s*(?:руб|р\.)',
                page,
                re.IGNORECASE
            )
            if amount_match:
                amount_val = parse_amount(amount_match.group(1))
                if amount_val > 0:
                    result["legal_fees"] = amount_val

        if has_legal_context and not result.get("legal_fees"):
            for line in lines:
                lower_line = line.lower()
                if "стоимость" in lower_line and "услуг" in lower_line:
                    amount_matches = re.findall(
                        r'(\d[\d\s]{2,}[.,]?\d*)',
                        line
                    )
                    if amount_matches:
                        amount_val = parse_amount(amount_matches[-1])
                        if amount_val > 0:
                            result["legal_fees"] = amount_val
                            break

    return result


def extract_first_date(text: str) -> Optional[datetime]:
    match = re.search(r'\d{2}[./]\d{2}[./]\d{4}', text)
    if match:
        return parse_date_str(match.group(0).replace('/', '.'))
    match = re.search(
        r'\d{1,2}\s+[А-Яа-яЁё\.]+\s+\d{4}',
        text
    )
    if match:
        return parse_ru_text_date(match.group(0))
    return None


def extract_cargo_docs_from_pages(
    pages: List[str],
    allow_llm: bool = True
) -> List[Dict[str, Any]]:
    cargo_docs = []
    seen = set()
    support_doc_types = {
        "Реестр сопроводительных документов",
        "Инструкция для водителя",
        "Маршрутный лист",
        "Акт проведения дезинфекции автотранспорта",
        "Чек-лист проверки температуры и санитарного состояния ТС",
        "Акт осмотра продукции перед отгрузкой",
        "Перечень материальных ценностей",
        "Доверенность",
        "Акт контроля погрузки/разгрузки продукции",
    }
    doc_markers = (
        "накладн",
        "экспедиторск",
        "реестр сопровод",
        "реестр",
        "инструкция",
        "инструкция для водителя",
        "маршрутный лист",
        "дезинфекц",
        "чек-лист",
        "чеклист",
        "торг-12",
        "торг-13",
        "м-15",
        "перечень материальных ценностей",
        "доверенност",
        "акт осмотра",
        "акт контроля",
    )
    last_context: Dict[str, Any] = {}
    patterns = [
        (
            "Транспортная накладная",
            re.compile(
                r'транспортн[^\n]{0,40}?накладн[^\n]{0,40}?(?:№|N[оo0])\s*'
                r'[:№]?\s*([A-Za-zА-Яа-я0-9/\\-]+)',
                re.IGNORECASE
            ),
        ),
        (
            "Товарно-транспортная накладная",
            re.compile(
                r'товарно[-\s]*транспортн[^\n]{0,40}?накладн[^\n]{0,40}?'
                r'(?:№|N[оo0])\s*[:№]?\s*([A-Za-zА-Яа-я0-9/\\-]+)',
                re.IGNORECASE
            ),
        ),
        # Формат с номером на отдельной строке:
        # "Транспортная накладная\n...\n№\nМ 96028/8117/0022"
        (
            "Транспортная накладная",
            re.compile(
                r'транспортная\s+накладная\s*\n'
                r'(?:[^\n]*\n){0,5}?'  # До 5 строк между
                r'№\s*\n\s*([A-Za-zА-Яа-я0-9/\s-]+)',
                re.IGNORECASE | re.MULTILINE
            ),
        ),
        # Формат pdfplumber: "Транспортная накладная Заказ (заявка)\nДата ... № М 96028/8117/0022"
        (
            "Транспортная накладная",
            re.compile(
                r'транспортная\s+накладная\s+(?:заказ|форма)[^\n]*\n'
                r'[^\n]*№\s*([A-Za-zА-Яа-я0-9/\s-]+?)(?:\s+Дата|\s*\n)',
                re.IGNORECASE | re.MULTILINE
            ),
        ),
    ]

    for page_index, page in enumerate(pages, 1):
        lowered = page.lower()
        if "акт сверки" in lowered:
            continue
        if not any(marker in lowered for marker in doc_markers):
            continue

        allow_llm_page = allow_llm and (
            "транспортная накладная" in lowered
            or "товарно-транспортная накладная" in lowered
            or "ттн" in lowered
        )
        page_details = extract_transport_details(
            page,
            allow_llm=allow_llm_page
        )
        page_details = {key: value for key, value in page_details.items() if value}
        if page_details:
            last_context = page_details.copy()

        def add_cargo_doc(
            doc_type: str,
            number: str,
            date_value: Optional[datetime],
            label: str
        ) -> None:
            date_str = (
                date_value.strftime("%d.%m.%Y") if date_value else ""
            )
            if number:
                key = (doc_type, number, date_str)
            else:
                key = (doc_type, label, date_str, page_index)
            if key in seen:
                return
            seen.add(key)
            entry = {
                "doc_type": doc_type,
                "number": number,
                "date": date_value,
                "label": label,
                "source_page": page_index,
            }
            details = page_details.copy()
            if doc_type in support_doc_types and last_context:
                filled = len([v for v in details.values() if v])
                if filled < 2:
                    merged = last_context.copy()
                    merged.update(details)
                    details = merged
            if details:
                entry.update(details)
            cargo_docs.append(entry)

        # Ищем дату накладной из специфичных паттернов
        # "Дата 15/11/2025" или "Дата | 15.11.2025" в начале документа
        doc_date_str = None
        doc_date_patterns = [
            # "Дата 15/11/2025" или "Дата 15.11.2025"
            re.compile(
                r'(?:^|\n)\s*Дата\s+(\d{1,2}[./]\d{1,2}[./]\d{4})',
                re.IGNORECASE | re.MULTILINE
            ),
            # "Дата | 15.11.2025"
            re.compile(
                r'Дата\s*[\|\t:]\s*(\d{1,2}[./]\d{1,2}[./]\d{4})',
                re.IGNORECASE
            ),
            # "от 15.11.2025" в контексте накладной
            re.compile(
                r'накладн[а-я]*\s+(?:№\s*[^\n]+?\s+)?от\s+(\d{2}[./]\d{2}[./]\d{4})',
                re.IGNORECASE
            ),
        ]
        for dp in doc_date_patterns:
            dm = dp.search(page)
            if dm:
                date_candidate = dm.group(1).replace('/', '.')
                # Проверяем что это не дата постановления (2021 год и ранее)
                if date_candidate and not date_candidate.endswith(('2021', '2020', '2019', '2018')):
                    doc_date_str = date_candidate
                    break

        for doc_type, pattern in patterns:
            for match in pattern.finditer(page):
                number = match.group(1).strip()
                # Очищаем номер от лишних символов и переносов
                number = re.sub(r'\s*\n.*', '', number)  # Убираем всё после \n
                number = number.strip()
                # Пропускаем ложные номера
                if not number or number.lower() in ('экземпляр', 'форма', 'дата', '1', '2', '3', '4'):
                    continue
                # Пропускаем слишком короткие номера (менее 3 символов без пробелов)
                if len(number.replace(' ', '')) < 3:
                    continue
                snippet = page[max(0, match.start() - 80):match.end() + 160]
                date_match = re.search(r'\d{2}\.\d{2}\.\d{4}', snippet)
                # Используем дату из сниппета только если она не из постановления
                date_str = ""
                if date_match:
                    candidate = date_match.group(0)
                    if not candidate.endswith(('2021', '2020', '2019', '2018')):
                        date_str = candidate
                # Если не нашли в сниппете, используем дату документа
                if not date_str and doc_date_str:
                    date_str = doc_date_str
                date_value = parse_date_str(date_str) if date_str else None
                label = (
                    f"{doc_type} № {number} от {date_str}"
                    if date_str else f"{doc_type} № {number}"
                )
                add_cargo_doc(doc_type, number, date_value, label)

        if "транспортная накладная" in lowered:
            for line in page.splitlines():
                if "накладн" not in line.lower() or "№" not in line:
                    continue
                numbers = re.findall(r'\d{7,}', line)
                if not numbers:
                    continue
                # Ищем дату в строке, избегая дат постановлений
                date_match = re.search(r'\d{2}\.\d{2}\.\d{4}', line)
                date_str = ""
                if date_match:
                    candidate = date_match.group(0)
                    if not candidate.endswith(('2021', '2020', '2019', '2018')):
                        date_str = candidate
                # Используем дату документа если не нашли
                if not date_str and doc_date_str:
                    date_str = doc_date_str
                date_value = parse_date_str(date_str) if date_str else None
                for number in numbers:
                    label = (
                        f"Транспортная накладная № {number} от {date_str}"
                        if date_str else f"Транспортная накладная № {number}"
                    )
                    add_cargo_doc("Транспортная накладная", number, date_value, label)

        if "акт контроля" in lowered:
            if re.search(
                r'акт\s+контроля\s+погрузки\s*/\s*разгрузки\s+продукции',
                page,
                re.IGNORECASE
            ):
                add_cargo_doc(
                    "Акт контроля погрузки/разгрузки продукции",
                    "",
                    extract_first_date(page),
                    "Акт контроля погрузки/разгрузки продукции"
                )

        if "экспедиторская расписка" in lowered:
            receipt_patterns = [
                re.compile(
                    r'экспедиторск\w+\s+расписк\w+[^\n]*?№\s*([A-Za-zА-Яа-я0-9/\\-]+)'
                    r'(?:[^\n]*?от\s*([0-9]{1,2}[./][0-9]{1,2}[./][0-9]{4}'
                    r'|[0-9]{1,2}\s+[А-Яа-яЁё\\.]+\s+[0-9]{4}))?',
                    re.IGNORECASE
                ),
                re.compile(
                    r'экспедиторск\w+\s+расписк\w+\s*'
                    r'(\d{2}[./]\d{2}[./]\d{4})\s*([A-Za-zА-Яа-я0-9/\\-]+)',
                    re.IGNORECASE
                ),
            ]
            date_candidate = None
            number = ""
            match = None
            for pattern in receipt_patterns:
                match = pattern.search(page)
                if match:
                    break
            if match:
                if len(match.groups()) >= 2 and match.group(2):
                    if re.match(r'\d{1,2}[./]\d{1,2}[./]\d{4}', match.group(2)):
                        date_candidate = parse_date_str(match.group(2).replace('/', '.'))
                        number = match.group(1).strip()
                    elif re.search(r'[А-Яа-яЁё]', match.group(2)):
                        date_candidate = parse_ru_text_date(match.group(2))
                        number = match.group(1).strip()
                    else:
                        number = match.group(1).strip()
                else:
                    date_candidate = parse_date_str(match.group(1).replace('/', '.'))
                    if len(match.groups()) > 1:
                        number = (match.group(2) or "").strip()
            if not date_candidate:
                date_candidate = extract_first_date(page)
            date_str = date_candidate.strftime("%d.%m.%Y") if date_candidate else ""
            label = (
                f"Экспедиторская расписка № {number} от {date_str}"
                if number and date_str else (
                    f"Экспедиторская расписка № {number}"
                    if number else "Экспедиторская расписка"
                )
            )
            add_cargo_doc(
                "Экспедиторская расписка",
                number,
                date_candidate,
                label
            )

        if "реестр" in lowered and "сопровод" in lowered:
            date_candidate = extract_first_date(page)
            date_str = date_candidate.strftime("%d.%m.%Y") if date_candidate else ""
            label = (
                f"Реестр сопроводительных документов от {date_str}"
                if date_str else "Реестр сопроводительных документов"
            )
            add_cargo_doc(
                "Реестр сопроводительных документов",
                "",
                date_candidate,
                label
            )

        if "инструкция для водителя" in lowered:
            add_cargo_doc(
                "Инструкция для водителя",
                "",
                extract_first_date(page),
                "Инструкция для водителя"
            )

        if "маршрутный лист" in lowered:
            match = re.search(
                r'маршрутный\s+лист\s*№\s*([A-Za-zА-Яа-я0-9/\\-]+)\s*'
                r'от\s*([^\n]+)',
                page,
                re.IGNORECASE
            )
            number = ""
            date_candidate = None
            if match:
                number = match.group(1).strip()
                date_raw = match.group(2).strip()
                date_candidate = (
                    parse_date_str(date_raw) or parse_ru_text_date(date_raw)
                )
            if not date_candidate:
                date_candidate = extract_first_date(page)
            date_str = date_candidate.strftime("%d.%m.%Y") if date_candidate else ""
            label = (
                f"Маршрутный лист № {number} от {date_str}"
                if number and date_str else (
                    f"Маршрутный лист № {number}"
                    if number else "Маршрутный лист"
                )
            )
            add_cargo_doc("Маршрутный лист", number, date_candidate, label)

        if "дезинфекц" in lowered and "акт" in lowered:
            date_candidate = extract_first_date(page)
            date_str = date_candidate.strftime("%d.%m.%Y") if date_candidate else ""
            label = (
                f"Акт проведения дезинфекции автотранспорта от {date_str}"
                if date_str else "Акт проведения дезинфекции автотранспорта"
            )
            add_cargo_doc(
                "Акт проведения дезинфекции автотранспорта",
                "",
                date_candidate,
                label
            )

        if ("чек-лист" in lowered or "чеклист" in lowered) and "температур" in lowered:
            match = re.search(
                r'ттн\s*№\s*([A-Za-zА-Яа-я0-9/\\-]+)',
                page,
                re.IGNORECASE
            )
            number = match.group(1).strip() if match else ""
            date_candidate = extract_first_date(page)
            label = (
                f"Чек-лист проверки температуры и санитарного состояния ТС (ТТН № {number})"
                if number else "Чек-лист проверки температуры и санитарного состояния ТС"
            )
            add_cargo_doc(
                "Чек-лист проверки температуры и санитарного состояния ТС",
                number,
                date_candidate,
                label
            )

        if "акт осмотра продукции" in lowered:
            date_candidate = extract_first_date(page)
            date_str = date_candidate.strftime("%d.%m.%Y") if date_candidate else ""
            label = (
                f"Акт осмотра продукции перед отгрузкой от {date_str}"
                if date_str else "Акт осмотра продукции перед отгрузкой"
            )
            add_cargo_doc(
                "Акт осмотра продукции перед отгрузкой",
                "",
                date_candidate,
                label
            )

        if "перечень материальных ценностей" in lowered:
            match = re.search(
                r'к\s+транспортной\s+накладной\s*№\s*([A-Za-zА-Яа-я0-9/\\-]+)',
                page,
                re.IGNORECASE
            )
            number = match.group(1).strip() if match else ""
            date_candidate = extract_first_date(page)
            date_str = date_candidate.strftime("%d.%m.%Y") if date_candidate else ""
            label = (
                f"Перечень материальных ценностей к ТН № {number} от {date_str}"
                if number and date_str else (
                    f"Перечень материальных ценностей к ТН № {number}"
                    if number else "Перечень материальных ценностей"
                )
            )
            add_cargo_doc(
                "Перечень материальных ценностей",
                number,
                date_candidate,
                label
            )

        if "м-15" in lowered and "накладная" in lowered:
            number_match = re.search(
                r'накладная\s*N[оo0]?\s*([A-Za-zА-Яа-я0-9/\\-]+)',
                page,
                re.IGNORECASE
            )
            number = number_match.group(1).strip() if number_match else ""
            date_match = re.search(
                r'дата\s*составления\s*(\d{2}[./]\d{2}[./]\d{4})',
                page,
                re.IGNORECASE
            )
            date_candidate = (
                parse_date_str(date_match.group(1).replace('/', '.'))
                if date_match else None
            )
            if not date_candidate:
                date_candidate = extract_first_date(page)
            date_str = date_candidate.strftime("%d.%m.%Y") if date_candidate else ""
            label = (
                f"Накладная М-15 № {number} от {date_str}"
                if number and date_str else (
                    f"Накладная М-15 № {number}"
                    if number else "Накладная М-15"
                )
            )
            add_cargo_doc("Накладная М-15", number, date_candidate, label)

        if "торг-13" in lowered:
            match = re.search(
                r'накладная\s+([A-Za-zА-Яа-я0-9/\\-]+)\s+'
                r'(\d{2}[./]\d{2}[./]\d{4})',
                page,
                re.IGNORECASE
            )
            number = ""
            date_candidate = None
            if match:
                number = match.group(1).strip()
                date_candidate = parse_date_str(match.group(2).replace('/', '.'))
            if not date_candidate:
                date_candidate = extract_first_date(page)
            date_str = date_candidate.strftime("%d.%m.%Y") if date_candidate else ""
            label = (
                f"Накладная ТОРГ-13 № {number} от {date_str}"
                if number and date_str else (
                    f"Накладная ТОРГ-13 № {number}"
                    if number else "Накладная ТОРГ-13"
                )
            )
            add_cargo_doc("Накладная ТОРГ-13", number, date_candidate, label)

        if ("товарная накладная" in lowered or "торг-12" in lowered) and "м-15" not in lowered:
            lines = [line.strip() for line in page.splitlines() if line.strip()]
            for idx, line in enumerate(lines):
                line_lower = line.lower()
                if "товарная накладная" not in line_lower and "торг-12" not in line_lower:
                    continue
                # Если в строке встречается "заявка", не считаем это накладной
                if "заявк" in line_lower:
                    continue
                match = re.search(
                    r'товарная\s+накладная[^\n]*№\s*([A-Za-zА-Яа-я0-9/\\-]+)'
                    r'[^\n]*?от\s*(\d{2}[./]\d{2}[./]\d{4})',
                    line,
                    re.IGNORECASE
                )
                if not match and "торг-12" in line_lower:
                    match = re.search(
                        r'№\s*([A-Za-zА-Яа-я0-9/\\-]+)'
                        r'[^\n]*?от\s*(\d{2}[./]\d{2}[./]\d{4})',
                        line,
                        re.IGNORECASE
                    )
                if not match and idx + 1 < len(lines):
                    next_line = lines[idx + 1]
                    next_lower = next_line.lower()
                    if "заявк" in next_lower:
                        continue
                    if "торг-12" in line_lower or "торг-12" in next_lower:
                        match = re.search(
                            r'№\s*([A-Za-zА-Яа-я0-9/\\-]+)\s*'
                            r'от\s*(\d{2}[./]\d{2}[./]\d{4})',
                            next_line,
                            re.IGNORECASE
                        )
                if not match:
                    continue
                number = match.group(1).strip()
                if len(number.replace(' ', '')) < 3:
                    continue
                date_candidate = parse_date_str(match.group(2).replace('/', '.'))
                date_str = date_candidate.strftime("%d.%m.%Y") if date_candidate else ""
                label = f"Товарная накладная № {number} от {date_str}"
                add_cargo_doc("Товарная накладная", number, date_candidate, label)

        if "доверенность" in lowered and "срок действия" in lowered:
            date_candidate = extract_first_date(page)
            date_str = date_candidate.strftime("%d.%m.%Y") if date_candidate else ""
            label = (
                f"Доверенность от {date_str}"
                if date_str else "Доверенность"
            )
            add_cargo_doc("Доверенность", "", date_candidate, label)

    # Постобработка: объединяем данные для накладных с одинаковыми номерами
    # (одна накладная может занимать несколько страниц PDF)
    merged_docs: Dict[object, Dict[str, Any]] = {}

    def format_date_key(value: Optional[datetime]) -> str:
        if isinstance(value, datetime):
            return value.strftime("%d.%m.%Y")
        return str(value) if value else ""

    for doc in cargo_docs:
        number = doc.get("number", "")
        doc_type = doc.get("doc_type") or doc.get("label") or ""
        if not number:
            date_key = format_date_key(doc.get("date"))
            meta_key = (
                doc_type,
                date_key,
                doc.get("vehicle_plate") or "",
                doc.get("trailer_plate") or "",
                doc.get("driver_name") or "",
                doc.get("sender_name") or "",
                doc.get("receiver_name") or "",
                format_date_key(doc.get("load_date")),
                format_date_key(doc.get("unload_date")),
            )
            if not any(meta_key[1:]):
                key = ("label", doc.get("label", ""), doc.get("source_page"))
            else:
                key = ("label",) + meta_key
            if key not in merged_docs:
                merged_docs[key] = doc
            continue

        key = (doc_type, number)
        if key not in merged_docs:
            merged_docs[key] = doc.copy()
        else:
            # Объединяем: берём непустые значения
            existing = merged_docs[key]
            for field, value in doc.items():
                if value and not existing.get(field):
                    existing[field] = value

    return list(merged_docs.values())


def enrich_cargo_docs_with_vision(
    cargo_docs: List[Dict[str, Any]],
    files: List[Dict[str, Any]],
    low_pages_info: List[Tuple[str, str, List[int]]]
) -> List[Dict[str, Any]]:
    """
    Обогащает данные грузовых документов с помощью Vision LLM
    для страниц с низким качеством текста.

    Args:
        cargo_docs: Список извлечённых документов
        files: Список файлов с путями
        low_pages_info: Информация о страницах с низким качеством
                        [(path, name, [page_nums]), ...]

    Returns:
        Обогащённый список документов
    """
    # Проверяем конфигурацию Vision
    config = get_vision_config()
    if not config.get("enabled") or not config.get("model"):
        logger.debug("Vision LLM not configured, skip enrichment")
        return cargo_docs

    # Создаём маппинг: путь → страницы с плохим качеством
    low_quality_map: Dict[str, List[int]] = {}
    for path, _name, pages in low_pages_info:
        low_quality_map[path] = pages

    # Поля, которые обычно отсутствуют
    key_fields = ["driver_name", "vehicle_plate", "trailer_plate"]

    enriched = []
    for doc in cargo_docs:
        # Проверяем, есть ли отсутствующие ключевые поля
        missing_fields = [f for f in key_fields if not doc.get(f)]
        if not missing_fields:
            enriched.append(doc)
            continue

        # Ищем файл с этим документом
        # (упрощённо - берём первый файл с низким качеством)
        for path, pages in low_quality_map.items():
            if not pages:
                continue

            logger.info(
                f"Vision enrichment for {doc.get('label')}, "
                f"missing: {missing_fields}"
            )

            # Пробуем Vision LLM для первой страницы с низким качеством
            try:
                vision_data = extract_transport_details_vision(
                    pdf_path=path,
                    page_num=pages[0] - 1,  # 0-indexed
                    existing_details=doc,
                    doc_type="transport"
                )

                if vision_data:
                    # Обновляем документ данными из Vision
                    for key, value in vision_data.items():
                        if value and not doc.get(key):
                            doc[key] = value
                            logger.info(f"Vision filled: {key}={value}")
                    break
            except Exception as e:
                logger.warning(f"Vision enrichment error: {e}")

        enriched.append(doc)

    return enriched


def extract_cdek_shipments_from_pages(pages: List[str]) -> List[Dict[str, Any]]:
    shipments = []
    seen = set()
    for page in pages:
        if "сдэк" not in page.lower() and "cdek" not in page.lower():
            continue
        track_match = re.search(r'Накладная\s*(\d{8,})', page)
        if not track_match:
            continue
        track_number = normalize_tracking_number(track_match.group(1))
        date_match = re.search(
            r'Подтверждено\s+по\s+CDEK\s+ID\s*(\d{2}\.\d{2}\.\d{4})',
            page,
            re.IGNORECASE
        )
        if date_match:
            date_str = date_match.group(1)
        else:
            date_match = re.search(r'\d{2}\.\d{2}\.\d{4}', page)
            date_str = date_match.group(0) if date_match else ""
        key = (track_number, date_str)
        if key in seen:
            continue
        seen.add(key)
        shipments.append({
            "track_number": track_number,
            "received_date": parse_date_str(date_str),
            "received_date_str": date_str,
            "source": "cdek",
        })
    return shipments


def extract_postal_shipments_from_pages(pages: List[str]) -> List[Dict[str, Any]]:
    shipments = []
    seen = set()
    delivery_keywords = (
        "вручение адресату",
        "вручено адресату",
        "получено адресатом",
        "вручение адресату почтальоном",
        "вручено адресату почтальоном",
        "доставлено адресату",
        "вручение получателю",
    )
    page_keywords = (
        "отчет об отслеживании",
        "почта россии",
        "идентификатор",
        "трек",
        "tracking",
        "рпо",
        "квитанц",
    )
    line_keywords = (
        "идентификатор",
        "трек",
        "отчет об отслеживании",
        "отслеживании",
        "рпо",
        "квитанц",
        "tracking",
    )

    def normalize_track(value: str) -> str:
        return normalize_tracking_number(value)

    def extract_numbers_from_line(line: str) -> List[str]:
        numbers = re.findall(r"\b\d{8,20}\b", line)
        normalized = []
        for num in numbers:
            candidate = normalize_track(num)
            if (
                candidate
                and is_valid_tracking_number(candidate)
                and (
                    not candidate.isdigit()
                    or len(candidate) in (13, 14)
                )
            ):
                normalized.append(candidate)
        return normalized

    for page in pages:
        lowered = page.lower()
        if not any(keyword in lowered for keyword in page_keywords):
            continue

        lines = [line.strip() for line in page.splitlines() if line.strip()]
        track_numbers: List[str] = []
        for idx, line in enumerate(lines):
            lower_line = line.lower()
            if any(keyword in lower_line for keyword in line_keywords):
                track_numbers.extend(extract_numbers_from_line(line))
                if not track_numbers and idx + 1 < len(lines):
                    track_numbers.extend(extract_numbers_from_line(lines[idx + 1]))

        if not track_numbers:
            for match in re.finditer(
                r'идентификатор[^\d]{0,20}(\d{8,20})',
                page,
                re.IGNORECASE
            ):
                candidate = normalize_track(match.group(1))
                if candidate and is_valid_tracking_number(candidate):
                    track_numbers.append(candidate)

        if not track_numbers and any(
            token in lowered
            for token in ("почта россии", "russian post", "отчет об отслеживании")
        ):
            track_numbers.extend(
                normalize_track(num)
                for num in re.findall(r"\b\d{13,14}\b", page)
            )

        track_numbers = [num for num in track_numbers if num]
        if not track_numbers:
            continue

        # Уникализируем, сохраняя порядок
        unique_numbers: List[str] = []
        seen_numbers = set()
        for num in track_numbers:
            if num in seen_numbers:
                continue
            seen_numbers.add(num)
            unique_numbers.append(num)
        track_numbers = unique_numbers

        received_date = None
        for idx, line in enumerate(lines):
            lower_line = line.lower()
            if any(key in lower_line for key in delivery_keywords):
                candidate = parse_ru_text_date(line)
                if not candidate:
                    for back in range(1, 3):
                        if idx - back < 0:
                            continue
                        candidate = parse_ru_text_date(lines[idx - back])
                        if candidate:
                            break
                received_date = candidate
                if received_date:
                    break
        if not received_date:
            match = re.search(
                r'(\d{1,2}\s+[А-Яа-яЁё\.]+\s+\d{4}).{0,40}'
                r'(Вручение адресату|Вручено адресату|Получено адресатом)',
                page,
                re.IGNORECASE | re.DOTALL
            )
            if match:
                received_date = parse_ru_text_date(match.group(1))
        received_date_str = (
            received_date.strftime("%d.%m.%Y")
            if received_date else ""
        )
        for track_number in track_numbers:
            key = (track_number, received_date_str)
            if key in seen:
                continue
            seen.add(key)
            shipments.append({
                "track_number": track_number,
                "received_date": received_date,
                "received_date_str": received_date_str,
                "source": "post",
            })
    return shipments


def extract_party_from_page(page: str, role: str) -> Optional[Dict[str, str]]:
    role_label = "Экспедитор" if role == "defendant" else "Перевозчик"
    if role_label.lower() not in page.lower():
        return None
    name_match = re.search(
        rf'{role_label}:\s*([^\n]+)', page, re.IGNORECASE
    )
    inn_match = re.search(
        rf'{role_label}:[^\n]*ИНН\s*(\d{{10,12}})',
        page,
        re.IGNORECASE
    )
    address_match = re.search(
        rf'{role_label}:[^\n]*Адрес:\s*([^\n]+)',
        page,
        re.IGNORECASE
    )
    name = name_match.group(1).strip() if name_match else ""
    inn = inn_match.group(1).strip() if inn_match else ""
    address = address_match.group(1).strip() if address_match else ""
    if not any([name, inn, address]):
        return None
    return {"name": name, "inn": inn, "address": address}


def extract_party_from_labels(
    pages: List[str],
    labels: List[str]
) -> Optional[Dict[str, str]]:
    label_set = {label.lower() for label in labels}
    label_norms = {re.sub(r'[^a-zа-яё]', '', label) for label in label_set}
    skip_tokens = {
        "инн",
        "кпп",
        "огрн",
        "огрнип",
        "адрес",
        "грузоотправитель",
        "грузополучатель",
        "датаивремя",
        "статус",
        "место",
        "вес",
        "видоп",
    } | label_norms
    def clean_party_name(value: str) -> str:
        if not value:
            return ""
        cleaned = re.split(
            r'\s+ИНН\b|\s+ИНН/КПП\b',
            value,
            maxsplit=1,
            flags=re.IGNORECASE
        )[0]
        cleaned = cleaned.split(",", 1)[0].strip()
        cleaned = re.sub(r'[\"«»]', '', cleaned).strip()
        return cleaned

    for page in pages:
        lines = [line.strip() for line in page.splitlines() if line.strip()]
        for idx, line in enumerate(lines):
            lower_line = line.lower()
            normalized_line = re.sub(r'[^a-zа-яё]', '', lower_line)
            if not any(normalized_line.startswith(label) for label in label_norms):
                continue
            name = line.split(":", 1)[1].strip() if ":" in line else ""
            if not name or not re.search(r"[A-Za-zА-Яа-яЁё]", name):
                for offset in range(1, 4):
                    if idx + offset >= len(lines):
                        break
                    candidate = lines[idx + offset].strip()
                    candidate_norm = re.sub(r'[^a-zа-яё]', '', candidate.lower())
                    if any(candidate_norm.startswith(token) for token in skip_tokens):
                        continue
                    if not re.search(r"[A-Za-zА-Яа-яЁё]", candidate):
                        continue
                    name = candidate
                    break

            name = clean_party_name(name)
            raw_value = line.split(":", 1)[1].strip() if ":" in line else ""
            raw_inn_match = re.search(
                r'\bИНН\s*[:\s]*(\d{10,12})',
                raw_value,
                re.IGNORECASE
            )
            raw_kpp_match = re.search(
                r'\bКПП\s*[:\s]*(\d{9})',
                raw_value,
                re.IGNORECASE
            )
            raw_ogrn_match = re.search(
                r'\bОГРН(?:ИП)?\s*[:\s]*(\d{13,15})',
                raw_value,
                re.IGNORECASE
            )
            inn = raw_inn_match.group(1) if raw_inn_match else ""
            kpp = raw_kpp_match.group(1) if raw_kpp_match else ""
            ogrn = raw_ogrn_match.group(1) if raw_ogrn_match else ""
            address = ""
            if raw_kpp_match:
                address = raw_value[raw_kpp_match.end():].strip(" ,")
            elif raw_inn_match:
                address = raw_value[raw_inn_match.end():].strip(" ,")
            if address:
                address = re.split(
                    r'\bтел\b|тел\.|телефон',
                    address,
                    1,
                    flags=re.IGNORECASE
                )[0].strip(" ,")
                if idx + 1 < len(lines):
                    next_line = lines[idx + 1].strip()
                    next_norm = re.sub(r'[^a-zа-яё]', '', next_line.lower())
                    if (
                        next_line
                        and not any(next_norm.startswith(token) for token in skip_tokens)
                        and not any(next_norm.startswith(label) for label in label_norms)
                        and re.search(r'\bул\b|\bдом\b|\d{5}|г\.?', next_line.lower())
                    ):
                        separator = ", " if address and not address.endswith(",") else " "
                        address = f"{address}{separator}{next_line}".strip()
                        address = re.split(
                            r'\bтел\b|тел\.|телефон',
                            address,
                            1,
                            flags=re.IGNORECASE
                        )[0].strip(" ,")
            for j in range(idx, min(idx + 12, len(lines))):
                current = lines[j]
                lower_current = current.lower()
                if "инн/кпп" in lower_current:
                    numbers = re.findall(r"\d{9,12}", current)
                    if not numbers and j + 1 < len(lines):
                        numbers = re.findall(r"\d{9,12}", lines[j + 1])
                    if numbers:
                        if not inn:
                            inn = numbers[0]
                        if len(numbers) > 1 and not kpp:
                            kpp = numbers[1]
                if lower_current.startswith("инн"):
                    numbers = re.findall(r"\d{10,12}", current)
                    if not numbers and j + 1 < len(lines):
                        numbers = re.findall(r"\d{10,12}", lines[j + 1])
                    if numbers and not inn:
                        inn = numbers[0]
                if lower_current.startswith("кпп"):
                    numbers = re.findall(r"\d{9}", current)
                    if not numbers and j + 1 < len(lines):
                        numbers = re.findall(r"\d{9}", lines[j + 1])
                    if numbers and not kpp:
                        kpp = numbers[0]
                if lower_current.startswith("огрн"):
                    numbers = re.findall(r"\d{13,15}", current)
                    if not numbers and j + 1 < len(lines):
                        numbers = re.findall(r"\d{13,15}", lines[j + 1])
                    if numbers and not ogrn:
                        ogrn = numbers[0]
                if lower_current.startswith("адрес"):
                    addr = current.split(":", 1)[1].strip() if ":" in current else ""
                    if not addr and j + 1 < len(lines):
                        addr = lines[j + 1].strip()
                    if addr and not address:
                        address = addr

            if address:
                address = re.sub(r'\s+,', ',', address)
                address = re.sub(r'г\.\s*,', 'г. ', address)
                address = re.sub(r'\s{2,}', ' ', address).strip()
            if any([name, inn, kpp, ogrn, address]):
                return {
                    "name": name,
                    "inn": inn,
                    "kpp": kpp,
                    "ogrn": ogrn,
                    "address": address,
                }
    return None


def extract_parties_from_pages(pages: List[str]) -> Dict[str, Dict[str, str]]:
    result: Dict[str, Dict[str, str]] = {}

    def has_org_token(name: str) -> bool:
        if not name:
            return False
        return bool(
            re.search(r'\b(ООО|ИП|ПАО|ЗАО|ОАО|АО)\b', name)
            or "Индивидуальный предприниматель" in name
        )

    def merge_party(
        primary: Optional[Dict[str, str]],
        fallback: Optional[Dict[str, str]]
    ) -> Optional[Dict[str, str]]:
        if not primary:
            return fallback
        if not fallback:
            return primary
        primary_name = primary.get("name", "")
        fallback_name = fallback.get("name", "")
        if not primary_name and fallback_name:
            primary["name"] = fallback_name
        elif fallback_name and has_org_token(fallback_name) and not has_org_token(primary_name):
            primary["name"] = fallback_name
        for field in ("inn", "kpp", "ogrn", "address"):
            if not primary.get(field) and fallback.get(field):
                primary[field] = fallback[field]
        return primary

    def build_inn_name_map(pages: List[str]) -> Dict[str, str]:
        inn_map: Dict[str, str] = {}
        org_pattern = re.compile(
            r'\b(ООО|ИП|АО|ПАО|ЗАО|ОАО)\s*[«\"]?([^\n\"»]+)[»\"]?'
            r'[^\n]{0,80}?\bИНН\s*(\d{10,12})',
            re.IGNORECASE
        )
        for page in pages:
            for match in org_pattern.finditer(page):
                org, name, inn = match.groups()
                cleaned = normalize_company_name(f"{org} «{name.strip()}»")
                if inn and cleaned:
                    inn_map[inn] = cleaned
        return inn_map

    def parse_party_from_text(value: str, inn_map: Dict[str, str]) -> Dict[str, str]:
        payload: Dict[str, str] = {"name": "", "inn": "", "kpp": "", "ogrn": "", "address": ""}
        if not value:
            return payload
        inn_match = re.search(r'\bИНН\s*[:\s]*(\d{10,12})', value, re.IGNORECASE)
        kpp_match = re.search(r'\bКПП\s*[:\s]*(\d{9})', value, re.IGNORECASE)
        ogrn_match = re.search(r'\bОГРН(?:ИП)?\s*[:\s]*(\d{13,15})', value, re.IGNORECASE)
        if inn_match:
            payload["inn"] = inn_match.group(1)
        else:
            digits = re.findall(r'\b\d{10,12}\b', value)
            if digits:
                payload["inn"] = digits[0]
        if kpp_match:
            payload["kpp"] = kpp_match.group(1)
        if ogrn_match:
            payload["ogrn"] = ogrn_match.group(1)
        name_part = re.split(r'\bИНН\b|\bКПП\b|\bОГРН\b', value, 1, flags=re.IGNORECASE)[0]
        name_part = re.sub(r'^(Заказчик|Исполнитель|Перевозчик|Экспедитор)\s*[:\-]?\s*', '', name_part, flags=re.IGNORECASE)
        name_part = normalize_company_name(name_part)
        if payload["inn"] and payload["inn"] in inn_map:
            payload["name"] = inn_map[payload["inn"]]
        else:
            payload["name"] = name_part
        if payload["name"] and not has_org_token(payload["name"]):
            payload["name"] = ""
        return payload

    def extract_party_from_role(
        pages: List[str],
        labels: Tuple[str, ...],
        inn_map: Dict[str, str]
    ) -> Optional[Dict[str, str]]:
        for page in pages:
            for label in labels:
                match = re.search(
                    rf'{label}\s*[:\-]\s*([^\n]+)',
                    page,
                    re.IGNORECASE
                )
                if match:
                    value = match.group(1).strip()
                    payload = parse_party_from_text(value, inn_map)
                    if any(payload.get(k) for k in ("name", "inn", "kpp", "ogrn")):
                        return payload
        return None

    invoice_pages = [
        page for page in pages
        if "счет на оплату" in page.lower() or "счёт на оплату" in page.lower()
    ]
    inn_name_map = build_inn_name_map(pages)
    seller = (
        extract_party_from_labels(invoice_pages, ["поставщик"])
        if invoice_pages else None
    )
    buyer = (
        extract_party_from_labels(invoice_pages, ["покупатель"])
        if invoice_pages else None
    )
    if not seller:
        seller = extract_party_from_labels(pages, ["продавец", "поставщик"])
    if not buyer:
        buyer = extract_party_from_labels(pages, ["покупатель"])
    postal_pages = [
        page for page in pages
        if "отчет об отслеживании" in page.lower()
        or "почты россии" in page.lower()
    ]
    sender = extract_party_from_labels(postal_pages, ["отправитель"])
    recipient = extract_party_from_labels(postal_pages, ["получатель"])

    result["plaintiff"] = merge_party(seller, sender) if seller or sender else None
    result["defendant"] = merge_party(buyer, recipient) if buyer or recipient else None
    result = {key: value for key, value in result.items() if value}

    # Ролевое извлечение: заявки и договоры
    role_plaintiff = extract_party_from_role(
        pages,
        ("исполнитель", "перевозчик"),
        inn_name_map
    )
    role_defendant = extract_party_from_role(
        pages,
        ("заказчик", "экспедитор"),
        inn_name_map
    )

    if "plaintiff" not in result or not result.get("plaintiff"):
        result["plaintiff"] = role_plaintiff
    if "defendant" not in result or not result.get("defendant"):
        result["defendant"] = role_defendant

    # Если нашли роли с ИНН, но у seller/buyer пустые ИНН — усилим
    if result.get("plaintiff") and role_plaintiff:
        if not result["plaintiff"].get("inn") and role_plaintiff.get("inn"):
            result["plaintiff"] = role_plaintiff
    if result.get("defendant") and role_defendant:
        if not result["defendant"].get("inn") and role_defendant.get("inn"):
            result["defendant"] = role_defendant

    if "plaintiff" not in result or "defendant" not in result:
        for page in pages:
            if "defendant" not in result:
                defendant = extract_party_from_page(page, "defendant")
                if defendant:
                    result["defendant"] = defendant
            if "plaintiff" not in result:
                plaintiff = extract_party_from_page(page, "plaintiff")
                if plaintiff:
                    result["plaintiff"] = plaintiff
    return result


def normalize_party_name(value: Optional[str]) -> str:
    if not value:
        return ""
    return re.sub(r'[^a-zа-я0-9]', '', str(value).lower())


def apply_extracted_parties(
    claim_data: Dict[str, Any],
    parties: Dict[str, Dict[str, str]]
) -> None:
    extracted_plaintiff = parties.get("plaintiff") or {}
    extracted_defendant = parties.get("defendant") or {}
    override_plaintiff = False
    override_defendant = False

    def apply_fields(prefix: str, payload: Dict[str, str], force: bool) -> None:
        mapping = {
            "name": f"{prefix}_name",
            "inn": f"{prefix}_inn",
            "kpp": f"{prefix}_kpp",
            "ogrn": f"{prefix}_ogrn",
            "address": f"{prefix}_address",
        }
        for key, target in mapping.items():
            if payload.get(key) and (force or is_missing_value(claim_data.get(target))):
                if key == "name":
                    claim_data[target] = normalize_company_name(payload.get(key))
                else:
                    claim_data[target] = payload.get(key)

    def enrich_with_dadata(prefix: str) -> None:
        inn_value = re.sub(r"[^\d]", "", str(claim_data.get(f"{prefix}_inn") or ""))
        if not is_valid_inn(inn_value):
            return
        kpp_value = re.sub(r"[^\d]", "", str(claim_data.get(f"{prefix}_kpp") or "")) or None
        suggestion = fetch_dadata_party_by_inn(inn_value, kpp=kpp_value, branch_type="MAIN")
        if not suggestion and kpp_value:
            suggestion = fetch_dadata_party_by_inn(inn_value, kpp=None, branch_type="MAIN")
        if not suggestion:
            return
        parsed = parse_dadata_party(suggestion)
        if parsed.get("inn") and parsed.get("inn") != inn_value:
            return
        if parsed.get("name"):
            claim_data[f"{prefix}_name"] = normalize_company_name(parsed.get("name", ""))
        if parsed.get("kpp"):
            claim_data[f"{prefix}_kpp"] = parsed["kpp"]
        if parsed.get("ogrn"):
            claim_data[f"{prefix}_ogrn"] = parsed["ogrn"]
        # Для ИП адрес берём из документа, не из DaData
        if parsed.get("type") != "INDIVIDUAL" and parsed.get("address"):
            claim_data[f"{prefix}_address"] = parsed["address"]

    current_plaintiff = normalize_party_name(
        claim_data.get("plaintiff_name")
    )
    current_defendant = normalize_party_name(
        claim_data.get("defendant_name")
    )
    current_plaintiff_inn = re.sub(r"[^\d]", "", str(claim_data.get("plaintiff_inn") or ""))
    current_defendant_inn = re.sub(r"[^\d]", "", str(claim_data.get("defendant_inn") or ""))
    extracted_plaintiff_name = normalize_party_name(
        extracted_plaintiff.get("name")
    )
    extracted_defendant_name = normalize_party_name(
        extracted_defendant.get("name")
    )
    extracted_plaintiff_inn = re.sub(r"[^\d]", "", str(extracted_plaintiff.get("inn") or ""))
    extracted_defendant_inn = re.sub(r"[^\d]", "", str(extracted_defendant.get("inn") or ""))

    if extracted_plaintiff_name and extracted_defendant_name:
        if not current_plaintiff or not current_defendant:
            claim_data["plaintiff_name"] = extracted_plaintiff.get("name")
            claim_data["defendant_name"] = extracted_defendant.get("name")
            override_plaintiff = True
            override_defendant = True
        elif current_plaintiff and current_defendant:
            if (
                extracted_plaintiff_inn
                and current_plaintiff_inn
                and extracted_plaintiff_inn != current_plaintiff_inn
            ):
                claim_data["plaintiff_name"] = extracted_plaintiff.get("name")
                override_plaintiff = True
            if (
                extracted_defendant_inn
                and current_defendant_inn
                and extracted_defendant_inn != current_defendant_inn
            ):
                claim_data["defendant_name"] = extracted_defendant.get("name")
                override_defendant = True
            if (
                current_plaintiff == current_defendant
                or (
                    current_plaintiff == extracted_defendant_name
                    and current_defendant == extracted_plaintiff_name
                )
            ):
                claim_data["plaintiff_name"] = extracted_plaintiff.get("name")
                claim_data["defendant_name"] = extracted_defendant.get("name")
                override_plaintiff = True
                override_defendant = True
            else:
                if current_plaintiff == extracted_defendant_name:
                    claim_data["plaintiff_name"] = extracted_plaintiff.get("name")
                    override_plaintiff = True
                if current_defendant == extracted_plaintiff_name:
                    claim_data["defendant_name"] = extracted_defendant.get("name")
                    override_defendant = True
    else:
        if not current_plaintiff and extracted_plaintiff.get("name"):
            claim_data["plaintiff_name"] = extracted_plaintiff.get("name")
        if not current_defendant and extracted_defendant.get("name"):
            claim_data["defendant_name"] = extracted_defendant.get("name")

    apply_fields("plaintiff", extracted_plaintiff, override_plaintiff)
    apply_fields("defendant", extracted_defendant, override_defendant)

    enrich_with_dadata("plaintiff")
    enrich_with_dadata("defendant")


def assign_invoices_to_applications(
    applications: List[Dict[str, Any]],
    invoices: List[Dict[str, Any]]
) -> Dict[str, Dict[str, Any]]:
    assignment = {}
    available = invoices.copy()
    for app in sorted(applications, key=lambda item: item.get("date") or datetime.min):
        if not available:
            break
        app_date = app.get("date")
        candidates = [
            inv for inv in available
            if inv.get("date") and app_date and inv["date"] >= app_date
        ]
        if candidates:
            chosen = min(candidates, key=lambda inv: inv["date"])
        else:
            chosen = min(
                available,
                key=lambda inv: inv.get("date") or datetime.max
            )
        assignment[app["label"]] = chosen
        available.remove(chosen)
    return assignment


def assign_upd_to_applications(
    applications: List[Dict[str, Any]],
    upd_docs: List[Dict[str, Any]]
) -> Dict[str, Dict[str, Any]]:
    assignment = {}
    available = upd_docs.copy()
    for app in sorted(applications, key=lambda item: item.get("date") or datetime.min):
        if not available:
            break
        app_date = app.get("date")
        candidates = [
            upd for upd in available
            if upd.get("date") and app_date and upd["date"] >= app_date
        ]
        if candidates:
            chosen = min(candidates, key=lambda upd: upd["date"])
        else:
            chosen = min(
                available,
                key=lambda upd: upd.get("date") or datetime.max
            )
        assignment[app["label"]] = chosen
        available.remove(chosen)
    return assignment


PLATE_TRANSLIT = str.maketrans({
    "А": "A",
    "В": "B",
    "Е": "E",
    "К": "K",
    "М": "M",
    "Н": "H",
    "О": "O",
    "Р": "P",
    "С": "C",
    "Т": "T",
    "У": "Y",
    "Х": "X",
})

ADDRESS_STOPWORDS = {
    "г", "город", "ул", "улица", "д", "дом", "кв", "квартира",
    "оф", "офис", "склад", "лит", "литер", "тер", "территория",
    "респ", "республика", "обл", "область", "район", "пр", "проспект",
    "пр-т", "ш", "шоссе", "пер", "переулок", "стр", "строение",
    "корп", "корпус", "вл", "владение",
}


def normalize_vehicle_plate(value: Optional[str]) -> str:
    if not value:
        return ""
    text = str(value).upper().translate(PLATE_TRANSLIT)
    return re.sub(r'[^A-Z0-9]', '', text)


def normalize_person_key(value: Optional[str]) -> str:
    if not value:
        return ""
    return re.sub(r'[^a-zа-я]', '', str(value).lower())


def normalize_address_tokens(value: Optional[str]) -> Set[str]:
    if not value:
        return set()
    text = str(value).lower().replace("ё", "е")
    text = re.sub(r'[^a-zа-я0-9]', ' ', text)
    tokens = [
        token for token in text.split()
        if token and token not in ADDRESS_STOPWORDS and len(token) > 1
    ]
    return set(tokens)


def score_date_match(
    first: Optional[datetime],
    second: Optional[datetime],
    full_points: int,
    near_points: int
) -> int:
    first = _coerce_date(first)
    second = _coerce_date(second)
    if not first or not second:
        return 0
    delta = abs((first - second).days)
    if delta == 0:
        return full_points
    if delta == 1:
        return near_points
    return 0


def score_token_overlap(tokens_a: Set[str], tokens_b: Set[str]) -> int:
    if not tokens_a or not tokens_b:
        return 0
    overlap = len(tokens_a & tokens_b)
    if overlap < 2:
        return 0
    return 6 + min(overlap - 2, 2) * 2


def _coerce_date(value: Any) -> Optional[datetime]:
    if not value:
        return None
    if isinstance(value, datetime):
        return value
    if hasattr(value, "year") and hasattr(value, "month") and hasattr(value, "day"):
        try:
            return datetime(value.year, value.month, value.day)
        except Exception:
            return None
    if isinstance(value, str):
        parsed = parse_date_str(value)
        if not parsed:
            parsed = parse_ru_text_date(value)
        return parsed
    return None


def parse_short_date(value: Optional[str]) -> Optional[datetime]:
    if not value:
        return None
    match = re.search(r'(\d{2})\.(\d{2})\.(\d{2})(\d{2})?', str(value))
    if not match:
        return None
    day, month, year1, year2 = match.groups()
    year_str = year1 + (year2 or "")
    try:
        year = int(year_str)
        if len(year_str) == 2:
            year = 2000 + year
        return datetime(int(year), int(month), int(day))
    except Exception:
        return None


def extract_reference_doc_numbers(text: str) -> List[str]:
    if not text:
        return []
    patterns = [
        r'(?:счет|счёт)\s*(?:на\s*оплату\s*)?(?:№|N[оo0]?)\s*([A-Za-zА-Яа-я0-9/\\-]+)',
        r'(?:упд|универсальн[^\n]{0,20}документ)\s*(?:№|N[оo0]?)\s*([A-Za-zА-Яа-я0-9/\\-]+)',
        r'заявк[ае]?\s*№\s*([A-Za-zА-Яа-я0-9/\\-]+)',
        r'накладн[а-я]*\s*№\s*([A-Za-zА-Яа-я0-9/\\-]+)',
    ]
    numbers: List[str] = []
    for pattern in patterns:
        for match in re.finditer(pattern, text, re.IGNORECASE):
            numbers.append(match.group(1))
    return numbers


def extract_reconciliation_payments(
    pages: List[str]
) -> List[Dict[str, Any]]:
    entries, _sales = extract_reconciliation_entries(pages)
    return [entry for entry in entries if entry.get("entry_type") == "payment"]


def extract_reconciliation_entries(
    pages: List[str]
) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    if not pages:
        return [], []
    target_pages = [
        page for page in pages
        if "акт сверки" in page.lower() or "сверк" in page.lower()
    ]
    if not target_pages:
        return [], []
    text = " ".join(target_pages)
    date_pattern = re.compile(r'\d{2}\.\d{2}\.\d{2,4}')
    matches = list(date_pattern.finditer(text))
    if not matches:
        return [], []

    entries: List[Dict[str, Any]] = []
    sales: List[Dict[str, Any]] = []
    seen_payments = set()
    seen_sales = set()

    for idx, match in enumerate(matches):
        start = match.start()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        segment = text[start:end]
        segment_clean = re.sub(r'\s+', ' ', segment).strip()
        if len(segment_clean) < 8:
            continue
        date_str = match.group(0)
        date_value = parse_date_str(date_str) or parse_short_date(date_str)
        if not date_value:
            continue

        amount_matches = re.findall(r'\d[\d\s]*[.,]\d{2}', segment_clean)
        if not amount_matches:
            continue
        amount = parse_amount(amount_matches[-1])
        if amount <= 0:
            continue

        lower = segment_clean.lower()
        entry_type = None
        if any(token in lower for token in ("оплата", "платеж", "платёж", "поступлен", "поступил", "перечисл", "п/п", "взаимозачет", "взаимозачёт")):
            entry_type = "payment"
        elif any(token in lower for token in ("продаж", "приход")):
            entry_type = "sale"

        if not entry_type:
            continue

        payment_number = None
        doc_number = None
        doc_date = None
        paren = re.search(r'\(([^)]+)\)', segment_clean)
        if paren:
            inner = paren.group(1)
            num_match = re.search(r'\b([A-Za-zА-Яа-я0-9/\\-]+)\b', inner)
            if num_match:
                doc_number = num_match.group(1)
            date_match = re.search(r'(\d{2}[./]\d{2}[./]\d{4})', inner)
            if date_match:
                doc_date = parse_date_str(date_match.group(1).replace('/', '.'))
            if entry_type == "payment":
                num_match = re.search(r'\b(\d{3,})\b', inner)
                if num_match:
                    payment_number = num_match.group(1)

        if not doc_number:
            num_match = re.search(r'№\s*([A-Za-zА-Яа-я0-9/\\-]+)', segment_clean)
            if num_match:
                doc_number = num_match.group(1)

        entry = {
            "entry_type": entry_type,
            "amount": amount,
            "date": date_value,
            "raw": segment_clean,
            "doc_number": doc_number,
            "doc_date": doc_date,
        }

        if entry_type == "payment":
            key = (
                date_value.strftime("%d.%m.%Y"),
                round(amount, 2),
                str(payment_number or ""),
            )
            if key in seen_payments:
                continue
            seen_payments.add(key)
            entry["payment_number"] = payment_number
            entry["source"] = "reconciliation"
            entry["reference_numbers"] = extract_reference_doc_numbers(segment_clean)
            entries.append(entry)
        else:
            key = (
                date_value.strftime("%d.%m.%Y"),
                round(amount, 2),
                str(doc_number or ""),
            )
            if key in seen_sales:
                continue
            seen_sales.add(key)
            sales.append(entry)

    return entries, sales


def match_reconciliation_payments_to_groups(
    groups: List[Dict[str, Any]],
    payments: List[Dict[str, Any]],
    sales: Optional[List[Dict[str, Any]]] = None
) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    if not groups or not payments:
        return [], payments or []

    def normalize_number(value: Optional[str]) -> str:
        if not value:
            return ""
        return re.sub(r'[^A-Za-zА-Яа-я0-9]', '', str(value)).upper()

    def extract_number_from_label(label: Optional[str]) -> str:
        if not label:
            return ""
        match = re.search(r'№\s*([A-Za-zА-Яа-я0-9/\\-]+)', label)
        if match:
            return normalize_number(match.group(1))
        return ""

    number_to_group: Dict[str, Optional[int]] = {}
    prepay_amount_map: Dict[float, List[int]] = {}
    invoice_amount_map: Dict[float, List[int]] = {}

    for idx, group in enumerate(groups):
        invoice_number = extract_number_from_label(group.get("invoice"))
        upd_number = extract_number_from_label(group.get("upd"))
        app_number = extract_number_from_label(group.get("application"))
        for number in (invoice_number, upd_number, app_number):
            if not number:
                continue
            existing = number_to_group.get(number)
            if existing is None and number in number_to_group:
                continue
            if existing is not None and existing != idx:
                number_to_group[number] = None
            else:
                number_to_group[number] = idx

        terms_text = normalize_payment_terms(group.get("payment_terms") or "")
        prepay_amount, _, _, _ = parse_prepayment_terms_details(terms_text)
        if prepay_amount > 0:
            key = round(prepay_amount, 2)
            prepay_amount_map.setdefault(key, []).append(idx)

        amount = group.get("amount") or 0.0
        if amount:
            key = round(float(amount), 2)
            invoice_amount_map.setdefault(key, []).append(idx)

    allocated: List[Dict[str, Any]] = []
    unassigned: List[Dict[str, Any]] = []
    sales = sales or []

    for payment in payments:
        assigned_idx: Optional[int] = None
        refs = payment.get("reference_numbers") or []
        for ref in refs:
            norm = normalize_number(ref)
            if not norm:
                continue
            target = number_to_group.get(norm)
            if target is None:
                continue
            assigned_idx = target
            break

        if assigned_idx is None:
            amount_key = round(float(payment.get("amount") or 0.0), 2)
            candidates = prepay_amount_map.get(amount_key) or []
            if len(candidates) == 1:
                assigned_idx = candidates[0]

        if assigned_idx is None:
            amount_key = round(float(payment.get("amount") or 0.0), 2)
            candidates = invoice_amount_map.get(amount_key) or []
            if len(candidates) == 1:
                assigned_idx = candidates[0]

        if assigned_idx is None:
            unassigned.append(payment)
            continue

        group = groups[assigned_idx]
        group.setdefault("payments", []).append(payment)
        payment["group_label"] = group.get("application") or group.get("invoice") or group.get("upd")
        allocated.append(payment)

    if unassigned and sales:
        sale_items: List[Dict[str, Any]] = []
        seen_sales = set()
        for entry in sales:
            doc_number = normalize_number(entry.get("doc_number"))
            if not doc_number:
                continue
            group_idx = number_to_group.get(doc_number)
            if group_idx is None:
                continue
            key = (
                doc_number,
                round(float(entry.get("amount") or 0.0), 2),
                entry.get("date").strftime("%d.%m.%Y") if hasattr(entry.get("date"), "strftime") else ""
            )
            if key in seen_sales:
                continue
            seen_sales.add(key)
            sale_items.append({
                "group_idx": group_idx,
                "amount": float(entry.get("amount") or 0.0),
                "date": entry.get("date"),
            })
        sale_items.sort(key=lambda item: item.get("date") or datetime.min)

        remaining_sales = sale_items[:]
        still_unassigned: List[Dict[str, Any]] = []
        for payment in unassigned:
            remaining = float(payment.get("amount") or 0.0)
            if remaining <= 0:
                continue
            while remaining > 0 and remaining_sales:
                sale = remaining_sales[0]
                if sale["amount"] <= 0:
                    remaining_sales.pop(0)
                    continue
                applied = min(remaining, sale["amount"])
                sale["amount"] -= applied
                remaining -= applied
                group = groups[sale["group_idx"]]
                payment_part = payment.copy()
                payment_part["amount"] = applied
                payment_part["allocation_method"] = "fifo"
                payment_part["group_label"] = (
                    group.get("application") or group.get("invoice") or group.get("upd")
                )
                group.setdefault("payments", []).append(payment_part)
                allocated.append(payment_part)
                if sale["amount"] <= 0:
                    remaining_sales.pop(0)
            if remaining > 0:
                payment["amount"] = remaining
                still_unassigned.append(payment)
        unassigned = still_unassigned

    return allocated, unassigned


def score_cargo_to_application(
    cargo: Dict[str, Any],
    app: Dict[str, Any]
) -> Tuple[int, List[str]]:
    """
    Скоринг сопоставления cargo-документа с заявкой.

    ВАЖНО: номер заявки из ТН используем как слабый сигнал,
    так как он часто относится к внутреннему заказу склада/покупателя.

    Сопоставление идёт ТОЛЬКО по:
    - Водитель (15 баллов)
    - Госномер ТС (20 баллов)
    - Госномер прицепа (10 баллов)
    - Даты погрузки/разгрузки (до 24 баллов)
    - Адреса (до 12 баллов)
    - Грузоотправитель/получатель (до 12 баллов)
    - Номер заявки (3 балла, слабый сигнал)
    """
    score = 0
    reasons: List[str] = []

    cargo_app_num = normalize_application_number(cargo.get("application_number"))
    app_num = normalize_application_number(app.get("number"))
    if cargo_app_num and app_num and cargo_app_num == app_num:
        score += 3
        reasons.append("номер заявки")

    cargo_driver = normalize_person_key(cargo.get("driver_name"))
    app_driver = normalize_person_key(app.get("driver_name"))
    if cargo_driver and app_driver and cargo_driver == app_driver:
        score += 15
        reasons.append("водитель")

    cargo_vehicle = normalize_vehicle_plate(cargo.get("vehicle_plate"))
    app_vehicle = normalize_vehicle_plate(app.get("vehicle_plate"))
    if cargo_vehicle and app_vehicle and cargo_vehicle == app_vehicle:
        score += 20
        reasons.append("транспорт")

    cargo_trailer = normalize_vehicle_plate(cargo.get("trailer_plate"))
    app_trailer = normalize_vehicle_plate(app.get("trailer_plate"))
    if cargo_trailer and app_trailer and cargo_trailer == app_trailer:
        score += 10
        reasons.append("прицеп")

    load_score = score_date_match(
        cargo.get("load_date"),
        app.get("load_date"),
        12,
        6,
    )
    if load_score:
        score += load_score
        reasons.append("дата погрузки")

    unload_score = score_date_match(
        cargo.get("unload_date"),
        app.get("unload_date"),
        12,
        6,
    )
    if unload_score:
        score += unload_score
        reasons.append("дата разгрузки")

    doc_score = score_date_match(
        cargo.get("date"),
        app.get("date"),
        6,
        3,
    )
    if doc_score:
        score += doc_score
        reasons.append("дата документа")

    load_tokens = normalize_address_tokens(cargo.get("load_address"))
    app_load_tokens = normalize_address_tokens(app.get("load_address"))
    load_addr_score = score_token_overlap(load_tokens, app_load_tokens)
    if load_addr_score:
        score += load_addr_score
        reasons.append("адрес погрузки")

    unload_tokens = normalize_address_tokens(cargo.get("unload_address"))
    app_unload_tokens = normalize_address_tokens(app.get("unload_address"))
    unload_addr_score = score_token_overlap(unload_tokens, app_unload_tokens)
    if unload_addr_score:
        score += unload_addr_score
        reasons.append("адрес разгрузки")

    cargo_sender = normalize_person_key(cargo.get("sender_name"))
    app_sender = normalize_person_key(app.get("sender_name"))
    if cargo_sender and app_sender and cargo_sender == app_sender:
        score += 6
        reasons.append("грузоотправитель")

    cargo_receiver = normalize_person_key(cargo.get("receiver_name"))
    app_receiver = normalize_person_key(app.get("receiver_name"))
    if cargo_receiver and app_receiver and cargo_receiver == app_receiver:
        score += 6
        reasons.append("грузополучатель")

    return score, reasons


def assign_cargo_to_applications(
    applications: List[Dict[str, Any]],
    cargo_docs: List[Dict[str, Any]]
) -> Dict[str, List[Dict[str, Any]]]:
    assignment: Dict[str, List[Dict[str, Any]]] = {
        app["label"]: [] for app in applications
    }
    apps_sorted = sorted(
        applications,
        key=lambda item: item.get("date") or datetime.min
    )
    for cargo in sorted(cargo_docs, key=lambda item: item.get("date") or datetime.min):
        cargo_date = cargo.get("date")
        chosen = None
        scored = []
        for app in apps_sorted:
            score, reasons = score_cargo_to_application(cargo, app)
            if score > 0:
                scored.append((score, app, reasons))
        if scored:
            scored.sort(
                key=lambda item: (
                    item[0],
                    item[1].get("date") or datetime.min,
                    item[1].get("label") or "",
                ),
                reverse=True,
            )
            best_score, best_app, reasons = scored[0]
            cargo["match_score"] = best_score
            cargo["match_reasons"] = reasons

            # Предупреждения о качестве сопоставления
            if len(scored) > 1:
                second_score = scored[1][0]
                if second_score == best_score or best_score - second_score < 5:
                    cargo["match_warning"] = "ambiguous"

            # Низкий скор — ненадёжное сопоставление
            if best_score < 15:
                cargo["match_warning"] = "low_confidence"
                logger.warning(
                    f"Низкий скор ({best_score}) для {cargo.get('label')}: "
                    f"привязан к {best_app.get('label')}, критерии: {reasons}"
                )

            chosen = best_app
            # LLM fallback для сложных случаев (низкий/неоднозначный скор)
            if cargo.get("match_warning") in ("low_confidence", "ambiguous"):
                llm_label, llm_conf = match_cargo_to_application_llm(
                    cargo,
                    apps_sorted
                )
                if llm_label and llm_conf >= 0.7:
                    llm_app = next(
                        (app for app in apps_sorted if app.get("label") == llm_label),
                        None
                    )
                    if llm_app:
                        chosen = llm_app
                        cargo["match_llm_confidence"] = llm_conf
                        if "llm" not in reasons:
                            reasons.append("llm")
                        if llm_conf >= 0.8:
                            cargo["match_score"] = max(cargo["match_score"], 15)
                            cargo.pop("match_warning", None)
        else:
            llm_label, llm_conf = match_cargo_to_application_llm(
                cargo,
                apps_sorted
            )
            if llm_label and llm_conf >= 0.7:
                llm_app = next(
                    (app for app in apps_sorted if app.get("label") == llm_label),
                    None
                )
                if llm_app:
                    cargo["match_score"] = max(cargo.get("match_score", 0), 15)
                    cargo["match_reasons"] = ["llm"]
                    cargo["match_llm_confidence"] = llm_conf
                    chosen = llm_app
            if cargo_date is None:
                cargo_label = (cargo.get("label") or "").lower()
                if "акт контроля" in cargo_label and apps_sorted:
                    chosen = apps_sorted[-1]
                elif apps_sorted:
                    chosen = apps_sorted[0]
            else:
                for app in reversed(apps_sorted):
                    app_date = app.get("date")
                    if app_date and cargo_date >= app_date:
                        chosen = app
                        break
            if chosen is None and apps_sorted:
                chosen = apps_sorted[0]
        if chosen:
            assignment[chosen["label"]].append(cargo)
    return assignment


_CARGO_SUPPORT_DOC_TYPES = {
    "Реестр сопроводительных документов",
    "Инструкция для водителя",
    "Маршрутный лист",
    "Акт проведения дезинфекции автотранспорта",
    "Чек-лист проверки температуры и санитарного состояния ТС",
    "Акт осмотра продукции перед отгрузкой",
    "Перечень материальных ценностей",
    "Доверенность",
    "Акт контроля погрузки/разгрузки продукции",
}

_CARGO_CORE_MARKERS = (
    "накладн",
    "экспедиторск",
    "торг-12",
    "торг 12",
    "торг-13",
    "м-15",
    "реестр сопровод",
    "cmr",
    "коносамент",
    "авианаклад",
)


def _is_placeholder_doc_number(value: Optional[str]) -> bool:
    if not value:
        return False
    digits = re.sub(r"\D", "", str(value))
    if not digits:
        return False
    if set(digits) == {"0"}:
        return True
    if len(digits) >= 4 and len(set(digits)) == 1:
        return True
    if digits in {"1234567890", "9876543210", "0123456789", "123456789012"}:
        return True
    if len(digits) >= 6:
        inc = all(
            (int(digits[i]) + 1) % 10 == int(digits[i + 1])
            for i in range(len(digits) - 1)
        )
        dec = all(
            (int(digits[i]) - 1) % 10 == int(digits[i + 1])
            for i in range(len(digits) - 1)
        )
        if inc or dec:
            return True
    return False


def _is_latin_heavy_label(value: str) -> bool:
    if not value:
        return False
    letters = [ch for ch in value if ch.isalpha()]
    if not letters:
        return False
    latin = sum(ch.isascii() and ch.isalpha() for ch in letters)
    cyrillic = sum(bool(re.match(r"[А-Яа-яЁё]", ch)) for ch in letters)
    if cyrillic:
        return False
    ratio = latin / len(letters) if letters else 0
    return ratio > 0.6


def filter_cargo_docs_for_output(
    cargo_docs: List[Dict[str, Any]],
    min_app_date: Optional[datetime] = None,
    max_app_date: Optional[datetime] = None
) -> List[Dict[str, Any]]:
    if not cargo_docs:
        return []
    core_docs: List[Dict[str, Any]] = []
    fallback_docs: List[Dict[str, Any]] = []
    for doc in cargo_docs:
        label = normalize_document_item(doc.get("label") or "")
        if not label:
            continue
        number = doc.get("number") or ""
        if _is_placeholder_doc_number(number):
            continue
        if _is_latin_heavy_label(label):
            lower_label = label.lower()
            if "cmr" not in lower_label and "tir" not in lower_label:
                continue
        doc_date = doc.get("date")
        if isinstance(doc_date, datetime) and min_app_date:
            if doc_date < min_app_date - timedelta(days=120):
                continue
        if isinstance(doc_date, datetime) and max_app_date:
            if doc_date > max_app_date + timedelta(days=120):
                continue
        doc_type = doc.get("doc_type") or ""
        lower_label = label.lower()
        is_core = (
            doc_type in _CARGO_SUPPORT_DOC_TYPES
        ) is False and any(marker in lower_label for marker in _CARGO_CORE_MARKERS)
        if is_core:
            core_docs.append(doc)
        else:
            fallback_docs.append(doc)
    return core_docs or fallback_docs


def get_matching_warnings(
    cargo_assignment: Dict[str, List[Dict[str, Any]]]
) -> List[str]:
    """
    Генерирует список предупреждений о качестве сопоставления документов.
    """
    warnings = []
    for app_label, cargo_list in cargo_assignment.items():
        for cargo in cargo_list:
            warning_type = cargo.get("match_warning")
            score = cargo.get("match_score", 0)
            reasons = cargo.get("match_reasons", [])

            if warning_type == "low_confidence":
                warnings.append(
                    f"⚠️ Низкая уверенность ({score} баллов):\n"
                    f"   {cargo.get('label', 'Документ')}\n"
                    f"   → {app_label}\n"
                    f"   Критерии: {', '.join(reasons) if reasons else 'нет'}"
                )
            elif warning_type == "ambiguous":
                warnings.append(
                    f"⚠️ Неоднозначное сопоставление ({score} баллов):\n"
                    f"   {cargo.get('label', 'Документ')}\n"
                    f"   → {app_label}\n"
                    f"   Есть другие подходящие заявки"
                )
    return warnings


def get_missing_document_warnings(
    groups: List[Dict[str, Any]]
) -> List[str]:
    warnings: List[str] = []
    if not groups:
        return warnings
    for group in groups:
        missing = []
        if not group.get("cargo_docs"):
            missing.append("накладные/сопроводительные документы")
        if (
            not group.get("invoice")
            and not group.get("upd")
            and (group.get("amount") or 0) <= 0
        ):
            missing.append("счет или УПД (стоимость перевозки)")
        if missing:
            label = (
                group.get("application")
                or group.get("invoice")
                or group.get("upd")
                or "перевозке"
            )
            warnings.append(
                f"⚠️ По {label} не найдены: {', '.join(missing)}."
            )
    return warnings


def build_pretension_groups(
    applications: List[Dict[str, Any]],
    invoices: List[Dict[str, Any]],
    cargo_docs: List[Dict[str, Any]],
    upd_docs: Optional[List[Dict[str, Any]]] = None,
    payment_terms_by_application: Optional[Dict[str, Dict[str, Any]]] = None
) -> List[Dict[str, Any]]:
    groups = []
    upd_docs = upd_docs or []
    payment_terms_by_application = payment_terms_by_application or {}
    app_dates = [app.get("date") for app in applications if app.get("date")]
    min_app_date = min(app_dates) if app_dates else None
    max_app_date = max(app_dates) if app_dates else None

    if not applications and (invoices or upd_docs):
        if invoices:
            for inv in invoices:
                amount = inv.get("amount") or 0.0
                groups.append({
                    "application": None,
                    "application_date": None,
                    "invoice": inv.get("label"),
                    "invoice_date": inv.get("date"),
                    "invoice_amount": amount,
                    "upd": None,
                    "upd_date": None,
                    "upd_amount": 0.0,
                    "cargo_docs": [],
                    "cargo_dates": [],
                    "amount": amount,
                    "payment_terms": None,
                    "payment_days": None,
                    "load_date": None,
                    "docs_track_number": None,
                    "docs_received_date": None,
                    "shipping_source": None,
                })
        else:
            for upd in upd_docs:
                amount = upd.get("amount") or 0.0
                groups.append({
                    "application": None,
                    "application_date": None,
                    "invoice": None,
                    "invoice_date": None,
                    "invoice_amount": 0.0,
                    "upd": upd.get("label"),
                    "upd_date": upd.get("date"),
                    "upd_amount": amount,
                    "cargo_docs": [],
                    "cargo_dates": [],
                    "amount": amount,
                    "payment_terms": None,
                    "payment_days": None,
                    "load_date": None,
                    "docs_track_number": None,
                    "docs_received_date": None,
                    "shipping_source": None,
                })
        return groups

    if not applications:
        return groups

    invoice_assignment = assign_invoices_to_applications(applications, invoices)
    upd_assignment = assign_upd_to_applications(applications, upd_docs)
    cargo_assignment = assign_cargo_to_applications(applications, cargo_docs)

    for app in applications:
        invoice = invoice_assignment.get(app["label"])
        upd = upd_assignment.get(app["label"])
        cargo_list = cargo_assignment.get(app["label"], [])
        filtered_cargo = filter_cargo_docs_for_output(
            cargo_list,
            min_app_date=min_app_date,
            max_app_date=max_app_date
        )
        cargo_dates = [item.get("date") for item in cargo_list if item.get("date")]
        load_candidates: List[datetime] = []
        app_load = app.get("load_date")
        if app_load:
            app_load_dt = _coerce_date(app_load)
            if app_load_dt:
                load_candidates.append(app_load_dt)
        for cargo in cargo_list:
            cargo_load = _coerce_date(cargo.get("load_date"))
            if cargo_load:
                load_candidates.append(cargo_load)
        if not load_candidates:
            for cargo in cargo_list:
                cargo_date = _coerce_date(cargo.get("date"))
                if cargo_date:
                    load_candidates.append(cargo_date)
        load_date = min(load_candidates) if load_candidates else None
        invoice_amount = invoice.get("amount") if invoice else 0.0
        upd_amount = upd.get("amount") if upd else 0.0
        app_amount = app.get("amount") or 0.0
        amount = invoice_amount or upd_amount or app_amount or 0.0
        terms_payload = payment_terms_by_application.get(app["label"], {}) or {}
        terms_text = normalize_payment_terms(terms_payload.get("terms") or "")
        terms_days = terms_payload.get("days")
        groups.append({
            "application": app["label"],
            "application_date": app.get("date"),
            "invoice": invoice.get("label") if invoice else None,
            "invoice_date": invoice.get("date") if invoice else None,
            "invoice_amount": invoice_amount or 0.0,
            "upd": upd.get("label") if upd else None,
            "upd_date": upd.get("date") if upd else None,
            "upd_amount": upd_amount or 0.0,
            "cargo_docs": [item["label"] for item in filtered_cargo],
            "cargo_docs_all": [item["label"] for item in cargo_list],
            "cargo_docs_details": cargo_list,
            "cargo_dates": cargo_dates,
            "amount": amount,
            "payment_terms": terms_text or None,
            "payment_days": terms_days,
            "load_date": load_date,
            "docs_track_number": None,
            "docs_received_date": None,
            "shipping_source": None,
        })
    return groups


def assign_shipments_to_groups(
    groups: List[Dict[str, Any]],
    shipments: List[Dict[str, Any]]
) -> None:
    if not groups or not shipments:
        return
    shipments_to_use = shipments
    if any(item.get("api_records", 0) > 0 for item in shipments):
        shipments_to_use = [
            item for item in shipments if item.get("api_records", 0) > 0
        ]
    sorted_groups = sorted(
        groups,
        key=lambda item: (
            max(item.get("cargo_dates") or []) if item.get("cargo_dates") else (
                item.get("invoice_date")
                or item.get("upd_date")
                or item.get("application_date")
                or datetime.min
            )
        )
    )
    def _shipment_received_date(item: Dict[str, Any]) -> Optional[datetime]:
        return _coerce_date(
            item.get("received_date")
            or item.get("received_date_str")
            or item.get("receive_date")
        )

    sorted_shipments = sorted(
        shipments_to_use,
        key=lambda item: _shipment_received_date(item) or datetime.max
    )
    unassigned = [
        group for group in sorted_groups
        if not group.get("docs_received_date") or not group.get("docs_track_number")
    ]
    prev_date = None
    for idx, shipment in enumerate(sorted_shipments):
        ship_date = _shipment_received_date(shipment)
        if not ship_date:
            continue
        bucket = []
        for group in list(unassigned):
            group_date = (
                max(group.get("cargo_dates") or []) if group.get("cargo_dates") else (
                    group.get("invoice_date")
                    or group.get("upd_date")
                    or group.get("application_date")
                )
            )
            if not group_date:
                continue
            if group_date <= ship_date and (prev_date is None or group_date > prev_date):
                bucket.append(group)
        if not bucket and idx == len(sorted_shipments) - 1:
            bucket = unassigned[:]
        for group in bucket:
            track_number = shipment.get("track_number")
            if track_number and not group.get("docs_track_number"):
                group["docs_track_number"] = track_number
            if ship_date and not group.get("docs_received_date"):
                group["docs_received_date"] = ship_date.strftime("%d.%m.%Y")
            if not group.get("shipping_source"):
                group["shipping_source"] = normalize_shipping_source(
                    shipment.get("source")
                )
            if group in unassigned:
                unassigned.remove(group)
        if bucket:
            prev_date = ship_date


def build_documents_list_structured_for_groups(
    groups: List[Dict[str, Any]]
) -> Optional[List[Tuple[int, str]]]:
    if not groups:
        return None
    structured = []
    for idx, group in enumerate(groups, 1):
        header = (
            group.get("application")
            or group.get("invoice")
            or group.get("upd")
            or "Перевозка"
        )
        structured.append((0, f"{idx}. {header}"))
        if group.get("invoice"):
            structured.append((1, f"{group['invoice']};"))
        if group.get("upd"):
            structured.append((1, f"{group['upd']};"))
        for doc in group.get("cargo_docs", []):
            structured.append((1, f"{doc};"))
        amount = group.get("amount") or 0.0
        if amount > 0:
            structured.append(
                (1, f"Цена перевозки {format_money(amount, 0)} руб.;")
            )
        track = group.get("docs_track_number")
        received = group.get("docs_received_date")
        source = normalize_shipping_source(group.get("shipping_source"))
        if track and received:
            if source == "cdek":
                structured.append(
                    (1, f"Отправление документов организацией СДЭК с номером {track} и получением {received};")
                )
            elif source == "post":
                structured.append(
                    (1, f"Отправление документов почтовым отправлением № {track} и получением {received};")
                )
            else:
                structured.append(
                    (1, f"Отправление документов с номером {track} и получением {received};")
                )
        elif track:
            if source == "cdek":
                structured.append(
                    (1, f"Отправление документов организацией СДЭК с номером {track};")
                )
            elif source == "post":
                structured.append(
                    (1, f"Отправление документов почтовым отправлением № {track};")
                )
            else:
                structured.append((1, f"Отправление документов с номером {track};"))
        elif received:
            if source == "cdek":
                structured.append(
                    (1, f"Получение документов по отправлению СДЭК {received};")
                )
            elif source == "post":
                structured.append(
                    (1, f"Получение документов по почтовому отправлению {received};")
                )
            else:
                structured.append((1, f"Получение документов {received};"))
    return structured


def build_shipping_summary(
    shipments: List[Dict[str, Any]],
    documents_count: Optional[int] = None
) -> str:
    if not shipments:
        return ""
    filtered_shipments = shipments
    if any(item.get("api_records", 0) > 0 for item in shipments):
        filtered_shipments = [
            item for item in shipments if item.get("api_records", 0) > 0
        ]
    numbers = []
    dates = []
    source = normalize_shipping_source(filtered_shipments[0].get("source"))
    for item in filtered_shipments:
        if item.get("track_number"):
            numbers.append(str(item["track_number"]))
        received_date = item.get("received_date")
        if isinstance(received_date, datetime):
            dates.append(received_date.strftime("%d.%m.%Y"))
        elif isinstance(received_date, str):
            parsed = parse_date_str(received_date)
            if parsed:
                dates.append(parsed.strftime("%d.%m.%Y"))
    numbers = list(dict.fromkeys(numbers))
    dates = list(dict.fromkeys(dates))
    if not numbers and not dates:
        return ""
    numbers_text = ", ".join(numbers)
    dates_text = ", ".join(dates)
    plural = documents_count is None or documents_count != 1
    doc_phrase = "перевозкам" if plural else "перевозке"
    if source == "cdek":
        prefix = (
            f"Документы по {doc_phrase} были отправлены организацией СДЭК."
        )
    elif source == "post":
        prefix = (
            f"Документы по {doc_phrase} были отправлены почтовыми отправлениями."
        )
    else:
        prefix = f"Документы по {doc_phrase} были отправлены."
    parts = [prefix]
    if numbers_text:
        parts.append(f"Номера отправлений: {numbers_text}.")
    if dates_text:
        parts.append(f"Отправления получены: {dates_text}.")
    return " ".join(parts)


def build_payment_terms_summary(
    groups: List[Dict[str, Any]],
    default_terms: str,
    default_days: int
) -> str:
    def finalize_line(text: str) -> str:
        cleaned = text.strip()
        if not cleaned:
            return cleaned
        if cleaned.endswith((".", "!", "?")):
            return cleaned
        return cleaned + "."

    lines = []
    grouped_terms: Dict[str, List[str]] = {}
    order: List[str] = []

    for group in groups:
        app_label = group.get("application")
        terms = normalize_payment_terms(group.get("payment_terms") or "")
        days = group.get("payment_days")
        if not terms:
            try:
                days_value = int(days)
            except (TypeError, ValueError):
                days_value = 0
            if days_value <= 0:
                days_value = default_days
            if days_value > 0:
                terms = (
                    "Оплата не позднее "
                    f"{days_value} рабочих дней с даты получения документов, "
                    "подтверждающих перевозку"
                )
        if not terms or terms == "Не указано":
            continue
        if terms not in grouped_terms:
            grouped_terms[terms] = []
            order.append(terms)
        if app_label:
            grouped_terms[terms].append(app_label)

    for terms in order:
        labels = grouped_terms.get(terms, [])
        terms_text = f"«{terms}»"
        if labels:
            unique_labels = []
            seen = set()
            for label in labels:
                if label and label not in seen:
                    seen.add(label)
                    unique_labels.append(label)
            if len(unique_labels) == 1:
                label_text = re.sub(
                    r'^Заявка\b',
                    'заявке',
                    unique_labels[0],
                    flags=re.IGNORECASE
                )
                lines.append(
                    finalize_line(
                        f"Условия оплаты по {label_text} –{terms_text}"
                    )
                )
            else:
                joined = ", ".join(unique_labels)
                lines.append(
                    finalize_line(
                        f"Условия оплаты по заявкам: {joined} –{terms_text}"
                    )
                )
        else:
            lines.append(finalize_line(f"Условия оплаты –{terms_text}"))

    if lines:
        return "\n\n".join(lines)
    fallback = normalize_payment_terms(default_terms or "")
    if not fallback or fallback == "Не указано":
        return default_terms
    if "условия оплаты" in fallback.lower():
        return fallback
    return f"Условия оплаты по договору-заявке –«{fallback}»"


def split_document_items(value: Any) -> List[str]:
    if not value or value == "Не указано":
        return []
    if isinstance(value, (list, tuple)):
        items = value
    else:
        items = re.split(r"[;\n]+", str(value))
    cleaned = []
    for item in items:
        text = normalize_document_item(item)
        if text:
            cleaned.append(fix_number_spacing(text))
    return cleaned


def build_document_groups_from_data(claim_data: Dict[str, Any]) -> List[Dict[str, Any]]:
    applications = split_document_items(claim_data.get("contract_applications"))
    invoices = split_document_items(claim_data.get("invoice_blocks"))
    acts = split_document_items(claim_data.get("upd_blocks"))
    cargo = split_document_items(claim_data.get("cargo_docs"))

    groups: List[Dict[str, Any]] = []
    if applications:
        for index, app in enumerate(applications):
            items = []
            if len(invoices) == len(applications):
                items.append(invoices[index])
            if len(acts) == len(applications):
                items.append(acts[index])
            if len(cargo) == len(applications):
                items.append(cargo[index])
            groups.append({"application": app, "documents": items})

        leftovers = []
        for docs in (invoices, acts, cargo):
            if docs and len(docs) != len(applications):
                leftovers.extend(docs)
        if leftovers:
            groups.append({"application": None, "documents": leftovers})
        return groups

    documents = invoices + acts + cargo
    if documents:
        groups.append({"application": None, "documents": documents})
    return groups


def build_document_groups(text: str, claim_data: Dict[str, Any]) -> List[Dict[str, Any]]:
    llm_payload = extract_document_groups_llm(text)
    groups: List[Dict[str, Any]] = []

    if llm_payload:
        for group in llm_payload.get("document_groups", []):
            application = group.get("application")
            documents = group.get("documents", []) or []
            if application or documents:
                groups.append({
                    "application": normalize_document_item(application) if application else None,
                    "documents": [normalize_document_item(doc) for doc in documents if doc],
                })
        ungrouped = llm_payload.get("ungrouped_documents", []) or []
        if ungrouped:
            groups.append({
                "application": None,
                "documents": [normalize_document_item(doc) for doc in ungrouped if doc],
            })

    if not groups:
        groups = build_document_groups_from_data(claim_data)

    # Remove duplicates while preserving order
    cleaned_groups = []
    seen = set()
    for group in groups:
        app = group.get("application")
        docs = group.get("documents", []) or []
        key = (app or "", tuple(docs))
        if key in seen:
            continue
        seen.add(key)
        cleaned_groups.append({"application": app, "documents": docs})

    return cleaned_groups


def build_documents_list_structured(
    groups: List[Dict[str, Any]]
) -> Optional[List[Tuple[int, str]]]:
    if not groups:
        return None
    structured: List[Tuple[int, str]] = []
    index = 1
    for group in groups:
        application = group.get("application")
        documents = group.get("documents", []) or []
        if application:
            structured.append((0, f"{index}. {application}"))
            for doc in documents:
                structured.append((1, doc))
            index += 1
        else:
            for doc in documents:
                structured.append((0, f"{index}. {doc}"))
                index += 1
    return structured or None


def build_party_block(
    label: str,
    name: str,
    inn: str,
    kpp: str,
    ogrn: str,
    ogrn_label: str,
    address: str,
    postal_address: str,
    is_ip: bool
) -> str:
    lines = [f"{label}: {name}"]
    if inn and inn != "Не указано":
        lines.append(f"ИНН {inn}")
    if not is_ip and kpp and kpp != "Не указано":
        lines.append(f"КПП {kpp}")
    if ogrn and ogrn != "Не указано":
        lines.append(f"{ogrn_label} {ogrn}")
    if address and address != "Не указано":
        lines.append(address)
    if (
        postal_address
        and postal_address != "Не указано"
        and postal_address != address
    ):
        lines.append(f"Почтовый адрес: {postal_address}")
    return "\n".join(lines)


def build_intro_paragraph(
    plaintiff_name_short: str,
    applications: List[str],
    cargo_docs: List[str],
    include_docs: bool = True
) -> str:
    parts = []
    if applications:
        parts.append(", ".join(applications))
    if cargo_docs:
        parts.append(", ".join(cargo_docs))
    documents_text = ", ".join(parts)
    if documents_text and include_docs:
        documents_text = f", что подтверждается {documents_text}."
    else:
        documents_text = "."
    app_phrase = "по указанным заявкам"
    if applications and len(applications) == 1:
        app_phrase = "по указанной заявке"
    return (
        f"{plaintiff_name_short} надлежащим образом выполнил(а) перевозку "
        f"{app_phrase}. Услуги оказаны своевременно и в полном объёме; "
        "приняты Заказчиком без замечаний и претензий"
        f"{documents_text}"
    )


def build_requirements_summary(
    debt_amount: float,
    total_interest: float,
    legal_fees: float
) -> str:
    parts = []
    if debt_amount > 0:
        parts.append(
            f"{format_money_ru(debt_amount, 2)} руб. — задолженность по оплате"
        )
    if total_interest > 0:
        parts.append(
            f"{format_money_ru(total_interest, 2)} руб. — проценты "
            "за пользование чужими денежными средствами"
        )
    if legal_fees > 0:
        parts.append(
            f"{format_money_ru(legal_fees, 2)} руб. — юридические услуги"
        )
    if not parts:
        return "Таким образом, размер требований составляет: Не указано."
    return "Таким образом, размер требований составляет: " + "; ".join(parts) + "."


def build_legal_fees_block(claim_data: Dict[str, Any]) -> str:
    legal_fees = parse_amount(claim_data.get("legal_fees", 0))
    contract_number = normalize_str(
        claim_data.get("legal_contract_number"),
        default=""
    )
    contract_date = normalize_str(
        claim_data.get("legal_contract_date"),
        default=""
    )
    payment_number = normalize_str(
        claim_data.get("legal_payment_number"),
        default=""
    )
    payment_date = normalize_str(
        claim_data.get("legal_payment_date"),
        default=""
    )
    if legal_fees <= 0 and not any([contract_number, contract_date, payment_number, payment_date]):
        return ""
    parts = [
        "Между Исполнителем и представителем заключён Договор оказания "
        "юридических услуг"
    ]
    if contract_number or contract_date:
        details = []
        if contract_number:
            details.append(f"№ {contract_number}")
        if contract_date:
            details.append(f"от {contract_date}")
        parts.append(" " + " ".join(details) + ".")
    else:
        parts.append(".")

    if payment_number or payment_date:
        payment_parts = []
        if payment_number:
            payment_parts.append(f"№ {payment_number}")
        if payment_date:
            payment_parts.append(f"от {payment_date}")
        parts.append(
            "Оплата по договору произведена Платёжным поручением "
            + " ".join(payment_parts)
            + (
                f" на сумму {format_money(legal_fees, 2)} руб."
                if legal_fees > 0 else ""
            )
        )
    else:
        if legal_fees > 0:
            parts.append(
                f"Оплата по договору произведена на сумму "
                f"{format_money(legal_fees, 2)} руб."
            )
    parts.append(
        "Расходы на представителя подлежат возмещению в порядке ст. 106, 110 "
        "АПК РФ, а также ст. 15, 393 ГК РФ и будут заявлены при обращении в суд."
    )
    return " ".join(parts)


def build_pretension_attachments(
    document_groups: List[Dict[str, Any]],
    claim_data: Dict[str, Any]
) -> List[str]:
    items: List[str] = []
    seen = set()

    def add_item(text: str) -> None:
        cleaned = normalize_document_item(text)
        key = normalize_attachment_text(cleaned)
        if not key or key in seen:
            return
        if len(cleaned) > 200:
            return
        seen.add(key)
        items.append(cleaned)

    def count_items(values: List[Optional[str]]) -> int:
        return len([val for val in values if val])

    app_count = count_items([g.get("application") for g in document_groups])
    invoice_count = count_items([g.get("invoice") for g in document_groups])
    upd_count = count_items([g.get("upd") for g in document_groups])
    cargo_count = 0
    for group in document_groups:
        if group.get("documents"):
            cargo_count += len(group.get("documents", []) or [])
            continue
        if group.get("cargo_docs_all"):
            cargo_count += len(group.get("cargo_docs_all", []) or [])
            continue
        cargo_count += len(group.get("cargo_docs", []) or [])

    if app_count:
        if app_count > 1:
            add_item("Комплект договор-заявок – копия")
        else:
            add_item("Копия договора-заявки")
    if invoice_count:
        if invoice_count > 1:
            add_item("Комплект счетов – копия")
        else:
            add_item("Копия счета")
    if upd_count:
        if upd_count > 1:
            add_item("Комплект УПД – копия")
        else:
            add_item("УПД – копия")
    if cargo_count:
        if cargo_count > 1:
            add_item("Комплект товаросопроводительных документов – копия")
        else:
            add_item("Копия товаросопроводительного документа")

    postal_number = get_first_list_value(claim_data.get("postal_numbers", []))
    postal_date = get_first_list_value(claim_data.get("postal_dates", []))
    if postal_number or postal_date or claim_data.get("shipments"):
        add_item("Чек отправления оригиналов документов по перевозкам – копия")

    contract_number = normalize_str(
        claim_data.get("legal_contract_number"),
        default=""
    )
    contract_date = normalize_str(
        claim_data.get("legal_contract_date"),
        default=""
    )
    legal_invoice_number = normalize_str(
        claim_data.get("legal_invoice_number"),
        default=""
    )
    legal_invoice_date = normalize_str(
        claim_data.get("legal_invoice_date"),
        default=""
    )
    payment_number = normalize_str(
        claim_data.get("legal_payment_number"),
        default=""
    )
    payment_date = normalize_str(
        claim_data.get("legal_payment_date"),
        default=""
    )

    if contract_number or contract_date:
        add_item("Договор на юр. услуги – копия")
    if legal_invoice_number or legal_invoice_date:
        add_item("Счет на оплату юр. услуг – копия")
    if payment_number or payment_date:
        add_item("Документы об оплате юр. услуг – копия")

    return items


def calculate_pretension_interest(
    debt_amount: float,
    start_date: datetime,
    end_date: Optional[datetime] = None,
    payments: Optional[List[Dict[str, Any]]] = None
) -> Dict[str, Any]:
    if debt_amount <= 0:
        return {"total_interest": 0.0, "detailed_calc": []}
    if end_date is None:
        end_date = datetime.today()
    if start_date > end_date:
        return {"total_interest": 0.0, "detailed_calc": []}

    if payments:
        events: Dict[datetime, float] = {}
        events[start_date] = events.get(start_date, 0.0) + debt_amount
        for payment in payments:
            amount = parse_amount(payment.get("amount"))
            date_value = _coerce_date(
                payment.get("date") or payment.get("payment_date")
            )
            if amount <= 0 or not date_value:
                continue
            event_date = date_value + timedelta(days=1)
            events[event_date] = events.get(event_date, 0.0) - amount
        if not events:
            return {"total_interest": 0.0, "detailed_calc": []}
        key_rates = get_key_rates_from_395gk()
        event_dates = sorted(events.keys())
        total_interest = 0.0
        detailed_calc: List[Dict[str, Any]] = []
        current_sum = 0.0

        for idx, event_date in enumerate(event_dates):
            delta = events[event_date]
            current_sum += delta
            if current_sum < 0:
                current_sum = 0.0
            period_start = event_date
            period_end = (
                event_dates[idx + 1] - timedelta(days=1)
                if idx + 1 < len(event_dates) else end_date
            )
            if period_end < period_start or current_sum <= 0:
                continue
            rate_periods = split_period_by_key_rate(period_start, period_end, key_rates)
            for rp_index, (start, end, rate) in enumerate(rate_periods):
                days = (end - start).days + 1
                year_days = 366 if (
                    start.year % 4 == 0 and
                    start.year % 100 != 0 or
                    start.year % 400 == 0
                ) else 365
                interest = current_sum * days * rate / 100 / year_days
                total_interest += interest
                detailed_calc.append({
                    "sum": current_sum,
                    "date_from": start,
                    "date_to": end,
                    "days": days,
                    "rate": rate,
                    "year_days": year_days,
                    "interest": interest,
                    "increase_sum": delta if rp_index == 0 else 0.0,
                    "increase_date": event_date if rp_index == 0 else None,
                })

        return {
            "total_interest": total_interest,
            "detailed_calc": detailed_calc,
        }

    key_rates = get_key_rates_from_395gk()
    periods = split_period_by_key_rate(start_date, end_date, key_rates)
    total_interest, detailed_calc = calc_395_on_periods(debt_amount, periods)
    return {
        "total_interest": total_interest,
        "detailed_calc": detailed_calc,
    }


def calculate_pretension_interest_schedule(
    groups: List[Dict[str, Any]],
    default_payment_days: int,
    payments: Optional[List[Dict[str, Any]]] = None
) -> Dict[str, Any]:
    def calculate_due_date(base_date: Optional[datetime], days: int) -> Optional[datetime]:
        if not base_date:
            return None
        if days < 0:
            days = 0
        calendar = load_work_calendar(base_date.year)
        return add_working_days(base_date, days, calendar) if days > 0 else base_date

    obligations: List[Dict[str, Any]] = []
    for group in groups:
        amount = float(group.get("amount") or 0)
        if amount <= 0:
            continue

        received_str = group.get("docs_received_date")
        received_date = _coerce_date(received_str)
        load_date = _coerce_date(group.get("load_date"))

        group_payment_days = group.get("payment_days")
        if group_payment_days is None:
            group_payment_days = default_payment_days
        try:
            group_payment_days = int(group_payment_days)
        except (TypeError, ValueError):
            group_payment_days = 0

        terms_text = normalize_payment_terms(group.get("payment_terms") or "")
        prepay_amount, prepay_days, prepay_base, remainder_days = parse_prepayment_terms_details(
            terms_text
        )

        remainder_amount = amount
        if prepay_amount > 0 and prepay_amount < amount:
            remainder_amount = amount - prepay_amount
        elif prepay_amount >= amount:
            prepay_amount = amount
            remainder_amount = 0.0

        group_label = group.get("application") or group.get("invoice") or group.get("upd")

        if prepay_amount > 0:
            base_date = load_date if prepay_base == "load" else load_date
            due_date = calculate_due_date(base_date, prepay_days)
            if due_date:
                obligations.append({
                    "group_label": group_label,
                    "due_date": due_date,
                    "amount": prepay_amount,
                })

        if remainder_amount > 0:
            if remainder_days is None:
                remainder_days = group_payment_days
            try:
                remainder_days = int(remainder_days)
            except (TypeError, ValueError):
                remainder_days = 0
            if remainder_days > 0 and received_date:
                due_date = calculate_due_date(received_date, remainder_days)
                if due_date:
                    obligations.append({
                        "group_label": group_label,
                        "due_date": due_date,
                        "amount": remainder_amount,
                    })

    events: Dict[datetime, float] = {}

    payments_by_group: Dict[Optional[str], List[Dict[str, Any]]] = {}
    if payments:
        for payment in payments:
            label = payment.get("group_label")
            payments_by_group.setdefault(label, []).append({
                "date": _coerce_date(payment.get("date") or payment.get("payment_date")),
                "amount": parse_amount(payment.get("amount")),
            })

    obligations_by_group: Dict[Optional[str], List[Dict[str, Any]]] = {}
    for obligation in obligations:
        label = obligation.get("group_label")
        obligations_by_group.setdefault(label, []).append(obligation)

    for label, group_obligations in obligations_by_group.items():
        group_obligations.sort(key=lambda o: o.get("due_date") or datetime.min)
        group_payments = payments_by_group.get(label, [])
        group_payments.sort(key=lambda p: p.get("date") or datetime.min)

        for obligation in group_obligations:
            amount = float(obligation.get("amount") or 0)
            if amount <= 0:
                continue
            due_date = obligation.get("due_date")
            if not due_date:
                continue

            if group_payments:
                for payment in group_payments:
                    pay_amount = payment.get("amount") or 0.0
                    payment_date = payment.get("date")
                    if pay_amount <= 0 or not payment_date:
                        continue
                    if payment_date < due_date and amount > 0:
                        applied = min(amount, pay_amount)
                        amount -= applied
                        payment["amount"] = pay_amount - applied

            if amount <= 0:
                continue
            interest_start = due_date + timedelta(days=1)
            events[interest_start] = events.get(interest_start, 0.0) + amount

        if group_payments:
            for payment in group_payments:
                pay_amount = payment.get("amount") or 0.0
                payment_date = payment.get("date")
                if pay_amount <= 0 or not payment_date:
                    continue
                event_date = payment_date + timedelta(days=1)
                events[event_date] = events.get(event_date, 0.0) - pay_amount

    if payments_by_group.get(None):
        for payment in payments_by_group.get(None, []):
            amount = payment.get("amount") or 0.0
            payment_date = payment.get("date")
            if amount <= 0 or not payment_date:
                continue
            event_date = payment_date + timedelta(days=1)
            events[event_date] = events.get(event_date, 0.0) - amount

    if not events:
        return {"total_interest": 0.0, "detailed_calc": []}

    key_rates = get_key_rates_from_395gk()
    event_dates = sorted(events.keys())
    today = datetime.today()
    total_interest = 0.0
    detailed_calc: List[Dict[str, Any]] = []
    current_sum = 0.0

    for idx, event_date in enumerate(event_dates):
        increase = events[event_date]
        current_sum += increase
        if current_sum < 0:
            current_sum = 0.0
        period_start = event_date
        period_end = (
            event_dates[idx + 1] - timedelta(days=1)
            if idx + 1 < len(event_dates) else today
        )
        if period_end < period_start:
            continue
        rate_periods = split_period_by_key_rate(period_start, period_end, key_rates)
        for rp_index, (start, end, rate) in enumerate(rate_periods):
            days = (end - start).days + 1
            year_days = 366 if (
                start.year % 4 == 0 and
                start.year % 100 != 0 or
                start.year % 400 == 0
            ) else 365
            interest = current_sum * days * rate / 100 / year_days
            total_interest += interest
            detailed_calc.append({
                "sum": current_sum,
                "date_from": start,
                "date_to": end,
                "days": days,
                "rate": rate,
                "year_days": year_days,
                "interest": interest,
                "increase_sum": increase if rp_index == 0 else 0.0,
                "increase_date": event_date if rp_index == 0 else None,
            })

    return {
        "total_interest": total_interest,
        "detailed_calc": detailed_calc,
    }


def build_interest_note_for_groups(
    groups: List[Dict[str, Any]],
    default_payment_days: int
) -> str:
    earliest_due: Optional[datetime] = None
    for group in groups:
        received_str = group.get("docs_received_date")
        if not received_str:
            continue
        group_days = group.get("payment_days")
        if group_days is None:
            group_days = default_payment_days
        try:
            group_days = int(group_days)
        except (TypeError, ValueError):
            group_days = 0
        if group_days <= 0:
            continue
        received_date = parse_date_str(received_str)
        if not received_date:
            continue
        calendar = load_work_calendar(received_date.year)
        due_date = add_working_days(received_date, group_days, calendar)
        if earliest_due is None or due_date < earliest_due:
            earliest_due = due_date
    if not earliest_due:
        return ""
    interest_start = earliest_due + timedelta(days=1)
    if interest_start <= datetime.today():
        return ""
    return (
        "Срок оплаты истекает "
        f"{earliest_due.strftime('%d.%m.%Y')}, "
        "проценты по ст. 395 ГК РФ начнут начисляться "
        f"с {interest_start.strftime('%d.%m.%Y')}."
    )


def replace_placeholders_robust(doc, replacements):
    """
    Заменяет плейсхолдеры в документе, применяя жирное начертание
    только для указанных полей и строк.
    """
    # Переменная для будущего использования жирных плейсхолдеров
    # bold_placeholders = [
    #     '{plaintiff_name}', '{defendant_name}',
    #     '{total_claim}', '{duty}'
    # ]
    # Ключевые фразы для жирного
    bold_lines = [
        'Арбитражный суд по месту нахождения ответчика',
        'Истец:',
        'Ответчик:',
        'Цена иска:',
        'Государственная пошлина:'
    ]
    multiline_placeholders = {
        'documents_list',
        'defendant_block',
        'plaintiff_block',
        'intro_paragraph',
        'legal_fees_block',
        'requirements_summary',
        'shipping_info',
        'payment_terms',
        'plaintiff_birth_info',
    }

    is_plaintiff_ip = (
        'ИП' in str(replacements.get('{plaintiff_name}', '')) or
        'Индивидуальный предприниматель' in str(
            replacements.get('{plaintiff_name}', ''))
    )
    is_defendant_ip = (
        'ИП' in str(replacements.get('{defendant_name}', '')) or
        'Индивидуальный предприниматель' in str(
            replacements.get('{defendant_name}', ''))
    )

    def is_missing(value: object) -> bool:
        if value is None:
            return True
        text = str(value).strip()
        return not text or text == 'Не указано'

    def split_list_values(value: object) -> list:
        if value is None:
            return []
        if isinstance(value, (list, tuple)):
            items = [str(item).strip() for item in value if str(item).strip()]
            return items
        text = str(value).strip()
        if not text or text == 'Не указано':
            return []
        return [part.strip() for part in re.split(r'[;,]', text) if part.strip()]

    def count_list_items(value: object) -> int:
        return len(split_list_values(value))

    def normalize_track_item(item: str) -> str:
        cleaned = re.sub(r'^[№\s]+', '', item)
        return cleaned.strip()

    def format_track_phrase(value: object) -> str:
        items = [normalize_track_item(item) for item in split_list_values(value)]
        if not items:
            return ''
        label = 'трек номерами' if len(items) > 1 else 'трек номером'
        return f"с {label} № {', '.join(items)}"

    def format_received_verb(track_value: object, date_value: object) -> str:
        count = max(
            count_list_items(track_value),
            count_list_items(date_value)
        )
        return 'получены' if count > 1 else 'получен'

    for paragraph in doc.paragraphs:
        full_text = ''.join(run.text for run in paragraph.runs)
        original_alignment = paragraph.alignment
        text_changed = False
        skip_replacements = False

        # Удаляем строку про оригиналы, если нет трек-номера или даты получения
        if 'Оригиналы документов' in full_text:
            track_value = replacements.get(
                '{docs_track_number}',
                replacements.get('{{docs_track_number}}', '')
            )
            date_value = replacements.get(
                '{docs_received_date}',
                replacements.get('{{docs_received_date}}', '')
            )
            if is_missing(track_value) or is_missing(date_value):
                p = paragraph._element
                p.getparent().remove(p)
                continue
            track_items = [
                normalize_track_item(item)
                for item in split_list_values(track_value)
            ]
            date_items = split_list_values(date_value)
            if track_items and date_items and len(track_items) == len(date_items):
                if len(track_items) == 1:
                    full_text = (
                        "Оригиналы документов по перевозкам отправлялись "
                        "почтовым отправлением "
                        f"с трек номером № {track_items[0]} "
                        f"получен {date_items[0]}."
                    )
                else:
                    pairs = [
                        f"№ {track} (получен {date})"
                        for track, date in zip(track_items, date_items)
                    ]
                    full_text = (
                        "Оригиналы документов по перевозкам отправлялись "
                        "почтовым отправлением "
                        f"с трек номерами {', '.join(pairs)}."
                    )
            else:
                full_text = (
                    "Оригиналы документов по перевозкам отправлялись почтовым "
                    f"отправлением {format_track_phrase(track_value)} "
                    f"{format_received_verb(track_value, date_value)} "
                    f"{str(date_value).strip()}."
                )
            text_changed = True
            skip_replacements = True

        # Удаляем строки с КПП для ИП
        if (
            is_plaintiff_ip
            and 'КПП' in full_text
            and ('{plaintiff_kpp}' in full_text or '{{plaintiff_kpp}}' in full_text)
        ):
            before_kpp = full_text
            lines = full_text.split('\n')
            filtered_lines = []
            for line in lines:
                if not (
                    'КПП' in line
                    and ('{plaintiff_kpp}' in line or '{{plaintiff_kpp}}' in line)
                ):
                    filtered_lines.append(line)
            full_text = '\n'.join(filtered_lines)
            if full_text != before_kpp:
                text_changed = True

        if (
            is_defendant_ip
            and 'КПП' in full_text
            and ('{defendant_kpp}' in full_text or '{{defendant_kpp}}' in full_text)
        ):
            before_kpp = full_text
            lines = full_text.split('\n')
            filtered_lines = []
            for line in lines:
                if not (
                    'КПП' in line
                    and ('{defendant_kpp}' in line or '{{defendant_kpp}}' in line)
                ):
                    filtered_lines.append(line)
            full_text = '\n'.join(filtered_lines)
            if full_text != before_kpp:
                text_changed = True

        if not skip_replacements:
            replaced_any = False
            for key in sorted(replacements.keys(), key=len, reverse=True):
                if key in full_text:
                    replaced_any = True
                    value = replacements[key]
                    placeholder_name = key.strip('{}')
                    if placeholder_name in multiline_placeholders:
                        clean_value = str(value).replace('\r', '').strip()
                    else:
                        clean_value = str(value).replace(
                            '\n', ' ').replace('\r', ' ').strip()
                    full_text = full_text.replace(key, clean_value)
            if replaced_any:
                text_changed = True

        if "итого задолженност" in full_text.lower():
            debt_amount = replacements.get(
                "{debt_amount}",
                replacements.get("{{debt_amount}}", "")
            )
            debt_kopeks = replacements.get(
                "{debt_kopeks}",
                replacements.get("{{debt_kopeks}}", "")
            )
            if debt_amount:
                updated = re.sub(
                    r'итого задолженность:\s*[\d\s,\.]+',
                    f"Итого задолженность: {debt_amount} ",
                    full_text,
                    flags=re.IGNORECASE
                )
                if updated != full_text:
                    full_text = updated
                    text_changed = True
            if debt_kopeks:
                updated = re.sub(
                    r'рублей\s+\d{2}\s+копеек',
                    f"рублей {debt_kopeks} копеек",
                    full_text,
                    flags=re.IGNORECASE
                )
                if updated != full_text:
                    full_text = updated
                    text_changed = True
        if full_text.startswith("В Арбитражный суд "):
            full_text = full_text[2:]
            text_changed = True
        updated = replace_ruble_words(full_text)
        if updated != full_text:
            full_text = updated
            text_changed = True

        signature_value = replacements.get(
            "{signatory}",
            replacements.get("{{signatory}}", "")
        )
        signature_value = str(signature_value or "").replace('\n', ' ').strip()
        if signature_value == "Не указано":
            signature_value = ""
        if "________________/" in full_text and full_text.strip().endswith("/"):
            if signature_value:
                updated = re.sub(
                    r'_{5,}/.*?/',
                    f"________________/{signature_value}/",
                    full_text
                )
            else:
                updated = re.sub(
                    r'_{5,}/.*?/',
                    "________________/__________________/",
                    full_text
                )
            if updated != full_text:
                full_text = updated
                text_changed = True

        if text_changed:
            if not full_text.strip():
                p = paragraph._element
                p.getparent().remove(p)
                continue
            paragraph.clear()
            if original_alignment is not None:
                paragraph.alignment = original_alignment
            # Проверяем, должна ли строка быть жирной
            is_bold = False
            for bold_line in bold_lines:
                if full_text.strip().startswith(bold_line):
                    is_bold = True

            # Специальная обработка для строк с названием суда
            if '{court_name}' in full_text or 'Арбитражный суд' in full_text:
                # Разбиваем текст на части и применяем жирное начертание к названию суда
                parts = full_text.split('Арбитражный суд')
                if len(parts) > 1:
                    # Первая часть (обычно "В")
                    if parts[0].strip():
                        prefix = parts[0]
                        if not prefix.endswith(' '):
                            prefix += ' '
                        run = paragraph.add_run(prefix)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        run.bold = False

                    # Название суда (жирное)
                    court_part = 'Арбитражный суд' + parts[1]
                    run = paragraph.add_run(court_part)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.bold = True
                else:
                    run = paragraph.add_run(full_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.bold = True
            # Специально для "Истец:" и "Ответчик:" — только первая строка жирная
            elif full_text.strip().startswith('Истец:') or full_text.strip().startswith('Ответчик:'):
                lines = full_text.split('\n')
                for i, line in enumerate(lines):
                    run = paragraph.add_run(line)
                    run.bold = (i == 0)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    if i < len(lines) - 1:
                        paragraph.add_run('\n')
            else:
                run = paragraph.add_run(full_text)
                run.bold = is_bold
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)


def fix_number_spacing(text: str) -> str:
    """
    Добавляет пробел после '№', если его нет.
    """
    return re.sub(r'№(\d)', r'№ \1', text)


ATTACHMENTS_EXTRA_TOP = [
    "Претензия – копия",
    "Чек и опись об отправке требования – копия",
]
ATTACHMENTS_EXTRA_TAIL = [
    "Квитанция об уплате государственной пошлины",
    "Документы, подтверждающие отправку искового заявления Ответчику - копия",
    "Доверенность на представителя – копия",
]
F107_EXCLUDED_ATTACHMENTS = [
    "Квитанция об уплате государственной пошлины",
    "Документы, подтверждающие отправку искового заявления Ответчику - копия",
    "Доверенность на представителя – копия",
]
F107_MAX_ITEMS = 14


def normalize_attachment_text(value: str) -> str:
    cleaned = str(value).lower().replace('ё', 'е')
    cleaned = re.sub(r'[\s\.,;:–—-]+', ' ', cleaned)
    return cleaned.strip()


def build_isk_attachments_list(attachments) -> List[str]:
    if isinstance(attachments, str):
        attachments = [attachments]
    attachments = attachments or []

    base_attachments = []
    for att in attachments:
        if not att or str(att).strip() == "Не указано":
            continue
        att_clean = fix_number_spacing(normalize_document_item(att))
        if att_clean:
            base_attachments.append(att_clean)

    final_attachments = []
    seen = set()

    def add_unique(item: str) -> None:
        key = normalize_attachment_text(item)
        if not key or key in seen:
            return
        seen.add(key)
        final_attachments.append(item)

    for item in ATTACHMENTS_EXTRA_TOP:
        add_unique(item)
    for item in base_attachments:
        add_unique(item)
    for item in ATTACHMENTS_EXTRA_TAIL:
        add_unique(item)

    return final_attachments


def resolve_defendant_display_name(
    name_short: Optional[str],
    name_full: Optional[str]
) -> str:
    for value in (name_short, name_full):
        if value and str(value).strip() and str(value).strip() != 'Не указано':
            return str(value).strip()
    return "Ответчик"


def build_f107_items(
    attachments,
    defendant_name: str
) -> List[str]:
    items = build_isk_attachments_list(attachments)
    excluded = {
        normalize_attachment_text(item)
        for item in F107_EXCLUDED_ATTACHMENTS
    }
    filtered_items = [
        item for item in items
        if normalize_attachment_text(item) not in excluded
    ]

    final_items = []
    seen = set()

    def add_unique(item: str) -> None:
        normalized = normalize_attachment_text(item)
        if not normalized or normalized in seen:
            return
        seen.add(normalized)
        final_items.append(item)

    claim_item = normalize_document_item(
        f"Исковое заявление к {defendant_name}"
    )
    add_unique(claim_item)
    for item in filtered_items:
        add_unique(item)

    if len(final_items) > F107_MAX_ITEMS:
        logging.warning(
            "Слишком много приложений для Ф107: %s, используется %s",
            len(final_items),
            F107_MAX_ITEMS
        )
        final_items = final_items[:F107_MAX_ITEMS]

    return final_items


def replace_attachments_with_paragraphs(
    doc,
    attachments,
    use_claim_extras: bool = True
):
    """
    Заменяет плейсхолдер {attachments} списком приложений с нумерацией и добавляет заголовок 'Приложения:'.
    Ожидается, что в шаблоне только {attachments} на отдельной строке.
    """
    import logging
    idx = None
    add_header = True
    parent = None
    header_idx = None
    placeholder_found = False
    placeholders = ['{attachments}', '{{attachments}}']
    for i, paragraph in enumerate(doc.paragraphs):
        if any(ph in paragraph.text for ph in placeholders):
            idx = i
            parent = paragraph._element.getparent()
            placeholder_found = True
            break

    if idx is None:
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip() == "Приложения:":
                idx = i
                parent = paragraph._element.getparent()
                header_idx = i
                break

    inserted_header = None
    if idx is not None:
        if header_idx is not None:
            # Удаляем существующий заголовок и старый список приложений
            p = doc.paragraphs[header_idx]._element
            parent.remove(p)
            while idx < len(doc.paragraphs):
                paragraph = doc.paragraphs[idx]
                text = paragraph.text.strip()
                if not text:
                    p = paragraph._element
                    parent.remove(p)
                    continue
                if (
                    "{{plaintiff_name_short}}" in text
                    or "{plaintiff_name_short}" in text
                    or text.startswith("Дата:")
                    or re.match(r"^_+", text)
                ):
                    break
                p = paragraph._element
                parent.remove(p)
            add_header = True

        if add_header:
            # Удаляем параграф с плейсхолдером
            if placeholder_found:
                p = doc.paragraphs[idx]._element
                parent.remove(p)

            # Добавляем заголовок "Приложения:"
            new_par = doc.add_paragraph()
            run = new_par.add_run("Приложения:")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            new_par.alignment = WD_ALIGN_PARAGRAPH.LEFT
            parent.insert(idx, new_par._element)
            inserted_header = new_par._element
            idx += 1
        else:
            # Удаляем старый список приложений из шаблона
            while idx < len(doc.paragraphs):
                paragraph = doc.paragraphs[idx]
                text = paragraph.text.strip()
                if not text:
                    p = paragraph._element
                    parent.remove(p)
                    continue
                if (
                    "{{plaintiff_name_short}}" in text
                    or "{plaintiff_name_short}" in text
                    or text.startswith("Дата:")
                    or re.match(r"^_+", text)
                ):
                    break
                p = paragraph._element
                parent.remove(p)

        if use_claim_extras:
            final_attachments = build_isk_attachments_list(attachments)
        else:
            if isinstance(attachments, str):
                attachments = [attachments]
            final_attachments = [
                fix_number_spacing(normalize_document_item(att))
                for att in (attachments or [])
                if att and str(att).strip() != "Не указано"
            ]

        # Динамические приложения с нумерацией
        attachment_number = 1
        for att in final_attachments:
            new_par = doc.add_paragraph()
            run = new_par.add_run(f"{attachment_number}. {att};")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = False
            new_par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            parent.insert(idx, new_par._element)
            idx += 1
            attachment_number += 1
        if add_header and inserted_header is not None:
            for paragraph in list(doc.paragraphs):
                if paragraph._element is inserted_header:
                    continue
                if paragraph.text.strip() == "Приложения:":
                    p = paragraph._element
                    p.getparent().remove(p)
    else:
        if attachments and attachments != ["Не указано"]:
            logging.warning(
                "Плейсхолдер {attachments} или блок 'Приложения:' не найден"
            )


def replace_documents_list_with_paragraphs(
    doc,
    structured_items: List[Tuple[int, str]]
) -> bool:
    """
    Заменяет {documents_list} на список параграфов с отступами.
    """
    placeholders = ['{documents_list}', '{{documents_list}}']
    idx = None
    for i, paragraph in enumerate(doc.paragraphs):
        if any(ph in paragraph.text for ph in placeholders):
            idx = i
            break

    if idx is None:
        return False

    p = doc.paragraphs[idx]._element
    parent = p.getparent()
    parent.remove(p)

    first_application = True
    for level, text in structured_items:
        line = text.strip()
        if not line:
            continue
        if level > 0 and not line.startswith(('-', '–', '—', '•')):
            line = f"• {line}"
        if level == 0 and not first_application:
            spacer = doc.add_paragraph()
            spacer_run = spacer.add_run("")
            spacer_run.font.name = 'Times New Roman'
            spacer_run.font.size = Pt(12)
            parent.insert(idx, spacer._element)
            idx += 1
        new_par = doc.add_paragraph()
        run = new_par.add_run(line)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        if level == 0:
            run.bold = True
        new_par.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if level > 0:
            new_par.paragraph_format.left_indent = Pt(18)
            new_par.paragraph_format.first_line_indent = Pt(-9)
        parent.insert(idx, new_par._element)
        idx += 1
        if level == 0:
            first_application = False

    return True


def iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested in cell.tables:
                yield from iter_table_paragraphs(nested)


def iter_document_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        yield from iter_table_paragraphs(table)


def replace_placeholders_simple(doc, replacements: Dict[str, str]) -> None:
    for paragraph in iter_document_paragraphs(doc):
        if not paragraph.runs:
            continue
        full_text = ''.join(run.text for run in paragraph.runs)
        if not full_text:
            continue
        updated = full_text
        for key, value in replacements.items():
            if key in updated:
                updated = updated.replace(key, value)
        if updated != full_text:
            paragraph.runs[0].text = updated
            for run in paragraph.runs[1:]:
                run.text = ''


def create_f107_document(
    items: List[str],
    sender_name: str,
    sender_company: str,
    output_path: Optional[str] = None
) -> str:
    template_dir = os.path.dirname(__file__)
    candidate_templates = [
        os.path.join(template_dir, 'templates', 'F107.docx'),
        os.path.join(template_dir, 'F107.docx'),
    ]
    template_path = next(
        (path for path in candidate_templates if os.path.exists(path)),
        None
    )
    if template_path is None:
        raise FileNotFoundError("Шаблон Ф107 не найден.")

    doc = Document(template_path)
    replacements: Dict[str, str] = {}
    total_quantity = 0
    total_value = 0

    for index in range(F107_MAX_ITEMS):
        suffix = '' if index == 0 else str(index)
        if index < len(items):
            item_text = normalize_document_item(items[index])
            replacements[f"${{predmet{suffix}}}"] = item_text
            replacements[f"${{kolich_predm{suffix}}}"] = "1"
            replacements[f"${{sum_predm{suffix}}}"] = "0"
            total_quantity += 1
        else:
            replacements[f"${{predmet{suffix}}}"] = ""
            replacements[f"${{kolich_predm{suffix}}}"] = ""
            replacements[f"${{sum_predm{suffix}}}"] = ""

    replacements["${sum_predmetov}"] = str(total_quantity) if total_quantity else ""
    replacements["${sum_kolich}"] = str(total_value) if total_quantity else ""
    replacements["${namef107}"] = ""
    replacements["${company}"] = sender_company or ""

    replace_placeholders_simple(doc, replacements)

    if output_path is None:
        output_name = f"Опись_вложения_F107_{uuid.uuid4().hex}.docx"
        output_path = os.path.join(os.path.dirname(__file__), output_name)
    doc.save(output_path)
    return output_path


def format_poa_date(value: Optional[datetime] = None) -> str:
    target = value or datetime.today()
    months = [
        'января', 'февраля', 'марта', 'апреля', 'мая', 'июня',
        'июля', 'августа', 'сентября', 'октября', 'ноября', 'декабря'
    ]
    month_name = months[target.month - 1]
    return f"«{target:%d}» {month_name} {target.year} г."


def create_power_of_attorney_document(
    replacements: Dict[str, str],
    output_path: Optional[str] = None
) -> str:
    template_dir = os.path.dirname(__file__)
    candidate_templates = [
        os.path.join(template_dir, 'templates', 'ДОВЕРЕННОСТЬ.docx'),
        os.path.join(template_dir, 'ДОВЕРЕННОСТЬ.docx'),
    ]
    template_path = next(
        (path for path in candidate_templates if os.path.exists(path)),
        None
    )
    if template_path is None:
        raise FileNotFoundError("Шаблон доверенности не найден.")

    doc = Document(template_path)
    replace_placeholders_simple(doc, replacements)

    if output_path is None:
        output_name = f"Доверенность_{uuid.uuid4().hex}.docx"
        output_path = os.path.join(os.path.dirname(__file__), output_name)
    doc.save(output_path)
    return output_path


def number_attachments_section(doc):
    """
    Добавляет нумерацию к списку приложений в шаблоне,
    если он указан как отдельные строки без номеров.
    """
    start_idx = None
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "Приложения:":
            start_idx = i + 1
            break
    if start_idx is None:
        return

    number = 1
    for i in range(start_idx, len(doc.paragraphs)):
        paragraph = doc.paragraphs[i]
        text = paragraph.text.strip()
        if not text:
            break
        if (
            "{{plaintiff_name_short}}" in text
            or "{plaintiff_name_short}" in text
            or re.match(r"^_+", text)
        ):
            break
        if re.match(r"^\d+[.)]\s+", text):
            number += 1
            continue

        cleaned = re.sub(r"^[-–—]\s*", "", text).strip()
        original_alignment = paragraph.alignment
        paragraph.clear()
        if original_alignment is not None:
            paragraph.alignment = original_alignment
        run = paragraph.add_run(f"{number}. {cleaned}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = False
        number += 1


def format_organization_name_short(full_name: str) -> str:
    """
    Форматирует полное название организации/ИП в краткий формат.

    Args:
        full_name: Полное название организации или ИП

    Returns:
        Краткий формат названия в кавычках
    """
    if not full_name or full_name == 'Не указано':
        return 'Не указано'
    full_name = normalize_company_name(full_name)

    # Для ИП
    if 'Индивидуальный предприниматель' in full_name:
        # Извлекаем ФИО после "Индивидуальный предприниматель"
        fio = full_name.replace('Индивидуальный предприниматель', '').strip()
        # Форматируем ФИО в формат "Фамилия И.О."
        parts = fio.split()
        if len(parts) >= 2:
            surname = parts[0]
            # Берем первые 2 инициала
            initials = '.'.join([part[0] for part in parts[1:3]]) + '.'
            return f'ИП «{surname} {initials}»'
        else:
            return f'ИП «{fio}»'

    # Для ООО
    elif 'Общество с ограниченной ответственностью' in full_name:
        # Извлекаем название в кавычках
        match = re.search(r'«(.+?)»', full_name)
        if match:
            return f'ООО «{match.group(1)}»'
        else:
            # Если нет кавычек, берем все после ООО
            name = full_name.replace(
                'Общество с ограниченной ответственностью',
                '').strip()
            return f'ООО «{name}»'

    # Для ЗАО
    elif 'Закрытое акционерное общество' in full_name:
        match = re.search(r'«(.+?)»', full_name)
        if match:
            return f'ЗАО «{match.group(1)}»'
        else:
            name = full_name.replace(
                'Закрытое акционерное общество',
                '').strip()
            return f'ЗАО «{name}»'

    # Для ПАО
    elif 'Публичное акционерное общество' in full_name:
        match = re.search(r'«(.+?)»', full_name)
        if match:
            return f'ПАО «{match.group(1)}»'
        else:
            name = full_name.replace(
                'Публичное акционерное общество', '').strip()
            return f'ПАО «{name}»'

    # Для ОАО
    elif 'Открытое акционерное общество' in full_name:
        match = re.search(r'«(.+?)»', full_name)
        if match:
            return f'ОАО «{match.group(1)}»'
        else:
            name = full_name.replace(
                'Открытое акционерное общество', '').strip()
            return f'ОАО «{name}»'

    # Для АО
    elif 'Акционерное общество' in full_name:
        match = re.search(r'«(.+?)»', full_name)
        if match:
            return f'АО «{match.group(1)}»'
        else:
            name = full_name.replace('Акционерное общество', '').strip()
            return f'АО «{name}»'

    # Если уже в кратком формате (содержит аббревиатуру)
    elif any(abbr in full_name for abbr in ['ООО', 'ИП', 'ЗАО', 'ПАО', 'ОАО', 'АО']):
        # Проверяем, есть ли уже кавычки
        if '«' in full_name and '»' in full_name:
            return full_name
        else:
            # Добавляем кавычки если их нет
            match = re.search(r'(ООО|ИП|ЗАО|ПАО|ОАО|АО)\s+(.+)', full_name)
            if match:
                return f'{match.group(1)} «{match.group(2).strip()}»'
            else:
                return full_name

    # Для остальных случаев возвращаем как есть
    return full_name


def create_isk_document(
    data: dict,
    interest_data: dict,
    duty_data: dict,
    replacements: dict,
    documents_list_structured: Optional[List[Tuple[int, str]]] = None,
    output_path: Optional[str] = None,
    proofread_protected_values: Optional[List[str]] = None
) -> str:
    template_dir = os.path.dirname(__file__)
    candidate_templates = [
        os.path.join(template_dir, 'templates', 'template_isk.docx'),
        os.path.join(template_dir, 'template_isk.docx'),
    ]
    fallback_templates = [
        os.path.join(template_dir, 'templates', 'template.docx'),
        os.path.join(template_dir, 'template.docx'),
    ]
    template_path = next(
        (path for path in candidate_templates if os.path.exists(path)),
        None
    )
    if template_path is None:
        template_path = next(
            (path for path in fallback_templates if os.path.exists(path)),
            None
        )
        logging.warning(
            "Шаблон template_isk.docx не найден, используется template.docx"
        )
    if template_path is None:
        raise FileNotFoundError(
            "Шаблон искового заявления не найден."
        )
    doc = Document(template_path)
    replacements = replacements.copy()
    if documents_list_structured:
        inserted = replace_documents_list_with_paragraphs(
            doc,
            documents_list_structured
        )
        if inserted:
            replacements.pop('{documents_list}', None)
            replacements.pop('{{documents_list}}', None)
    replace_placeholders_robust(doc, expand_placeholder_map(replacements))
    attachment_placeholders = ['{attachments}', '{{attachments}}']
    has_attachment_placeholder = any(
        ph in paragraph.text
        for paragraph in doc.paragraphs
        for ph in attachment_placeholders
    )
    has_attachments_header = any(
        paragraph.text.strip() == "Приложения:"
        for paragraph in doc.paragraphs
    )
    if has_attachment_placeholder or has_attachments_header:
        replace_attachments_with_paragraphs(
            doc,
            data.get('attachments', [])
        )
    number_attachments_section(doc)
    inserted_interest = False
    if interest_data.get('table_rows'):
        inserted_interest = insert_interest_table_from_rows(
            doc,
            interest_data.get('table_rows')
        )
    if not inserted_interest:
        insert_interest_table(
            doc,
            interest_data.get('detailed_calc', []),
            interest_data.get('total_interest')
        )
    proofread_docx_document(doc, protected_values=proofread_protected_values)
    enforce_times_new_roman(doc)
    if output_path is None:
        output_name = f"Исковое_заявление_{uuid.uuid4().hex}.docx"
        output_path = os.path.join(os.path.dirname(__file__), output_name)
    doc.save(output_path)
    return output_path


def remove_legal_fees_section(doc) -> None:
    """
    Удаляет блок про юридические услуги, если он пустой.
    """
    header_idx = None
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith("5. Договор об оказании юридических услуг"):
            header_idx = i
            break
    if header_idx is None:
        return
    indices = [header_idx]
    if header_idx + 1 < len(doc.paragraphs):
        next_text = doc.paragraphs[header_idx + 1].text.strip()
        if not next_text or "юридических услуг" in next_text.lower():
            indices.append(header_idx + 1)
    for idx in sorted(indices, reverse=True):
        p = doc.paragraphs[idx]._element
        p.getparent().remove(p)


def insert_awareness_block(doc, awareness_text: str):
    """
    Вставляет блок "осознанности" (частичные оплаты, гарантийные письма и т.д.)
    в документ после раздела "Итого задолженность".

    Args:
        doc: Документ Word
        awareness_text: Текст для вставки
    """
    if not awareness_text or not awareness_text.strip():
        return

    # Ищем абзац с "Итого задолженность"
    target_idx = None
    for i, para in enumerate(doc.paragraphs):
        text_lower = para.text.lower()
        if "итого задолженност" in text_lower or "итого долг" in text_lower:
            target_idx = i
            break

    if target_idx is None:
        # Если не нашли "Итого", ищем после "{debt_amount}"
        for i, para in enumerate(doc.paragraphs):
            if "{debt_amount}" in para.text or "рублей 00 копеек" in para.text.lower():
                target_idx = i
                break

    if target_idx is None:
        logger.warning("Не найдено место для вставки блока осознанности")
        return

    # Вставляем новый абзац после найденного
    target_para = doc.paragraphs[target_idx]
    new_para = doc.add_paragraph()

    # Перемещаем новый абзац на нужную позицию
    target_para._element.addnext(new_para._element)

    # Форматируем текст
    run = new_para.add_run("\n" + awareness_text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)


def adjust_pretension_quality_section(
    doc,
    documents_count: Optional[int]
) -> None:
    if documents_count != 1:
        return
    plural_text = (
        "Все перевозки выполнены надлежащим образом и приняты без претензий. "
        "Письменные претензии в установленный срок не заявлялись."
    )
    singular_text = (
        "Перевозка выполнена надлежащим образом и принята без претензий. "
        "Письменные претензии в установленный срок не заявлялись."
    )
    for paragraph in doc.paragraphs:
        if plural_text in paragraph.text:
            paragraph.text = paragraph.text.replace(plural_text, singular_text)
            return
        if "Все перевозки выполнены" in paragraph.text:
            paragraph.text = singular_text
            return


def adjust_pretension_interest_section(doc, interest_data: Dict[str, Any]) -> None:
    has_interest = False
    if interest_data.get("detailed_calc"):
        has_interest = True
    if float(interest_data.get("total_interest") or 0) > 0:
        has_interest = True
    if has_interest:
        return

    start_idx = None
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if re.search(r'Расч[её]т процентов по ст\.\s*395', text, re.IGNORECASE):
            start_idx = idx
            break
    if start_idx is None:
        return

    end_idx = len(doc.paragraphs)
    for idx in range(start_idx + 1, len(doc.paragraphs)):
        text = doc.paragraphs[idx].text.strip()
        if re.match(r'^\d+\.\s+', text) and not re.match(r'^\d+\.\d', text):
            end_idx = idx
            break

    for idx in range(end_idx - 1, start_idx - 1, -1):
        p = doc.paragraphs[idx]._element
        p.getparent().remove(p)

    section_keywords = (
        "Нормативное обоснование",
        "Договор об оказании юридических услуг",
        "Требования и срок исполнения",
    )
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not any(keyword in text for keyword in section_keywords):
            continue
        match = re.match(r'^(\d+)\.\s+(.*)$', text)
        if not match:
            continue
        number = int(match.group(1))
        if number >= 4:
            paragraph.text = f"{number - 1}. {match.group(2)}"


def enforce_times_new_roman(doc) -> None:
    def apply_runs(paragraphs):
        for paragraph in paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    apply_runs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                apply_runs(cell.paragraphs)


def _paragraph_has_uniform_runs(paragraph) -> bool:
    runs = paragraph.runs or []
    if not runs:
        return True
    def signature(run):
        return (
            run.bold,
            run.italic,
            run.underline,
            run.font.name,
            run.font.size,
        )
    first = signature(runs[0])
    return all(signature(run) == first for run in runs[1:])


def proofread_docx_document(doc, protected_values: Optional[List[str]] = None) -> None:
    enabled_raw = os.getenv("LLM_PROOFREAD_WHOLE_DOC")
    if not enabled_raw or enabled_raw.lower() not in ("1", "true", "yes", "on"):
        return

    cache: Dict[str, str] = {}

    def process_paragraph(paragraph) -> None:
        text = paragraph.text.strip()
        if not text:
            return
        if "{" in text and "}" in text:
            return
        if re.match(r"^\d+[.)]\s+", text):
            return
        if text.startswith(("•", "-", "–", "—")):
            return
        if text.lower() in {"приложения:", "приложения"}:
            return
        if not re.search(r"[A-Za-zА-Яа-яЁё]", text):
            return
        if not _paragraph_has_uniform_runs(paragraph):
            return
        cached = cache.get(text)
        if cached is None:
            cached = maybe_proofread_text(
                text,
                protected_values=protected_values
            )
            cache[text] = cached
        if cached and cached != text:
            paragraph.text = cached

    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            process_paragraph(paragraph)
        for paragraph in section.footer.paragraphs:
            process_paragraph(paragraph)


def create_pretension_document(
    data: dict,
    interest_data: dict,
    replacements: dict,
    documents_list_structured: Optional[List[Tuple[int, str]]] = None,
    attachments: Optional[List[str]] = None,
    output_path: Optional[str] = None,
    proofread_protected_values: Optional[List[str]] = None
) -> str:
    template_dir = os.path.dirname(__file__)
    candidate_templates = [
        os.path.join(template_dir, "templates", "template_pretension.docx"),
        os.path.join(template_dir, "template_pretension.docx"),
        os.path.join(template_dir, "templates", "ПРЕТЕНЗИЯ.docx"),
        os.path.join(template_dir, "ПРЕТЕНЗИЯ.docx"),
    ]
    template_path = next(
        (path for path in candidate_templates if os.path.exists(path)),
        None
    )
    if template_path is None:
        raise FileNotFoundError("Шаблон претензии не найден.")

    doc = Document(template_path)
    replacements = replacements.copy()
    if documents_list_structured:
        inserted = replace_documents_list_with_paragraphs(
            doc,
            documents_list_structured
        )
        if inserted:
            replacements.pop("{documents_list}", None)
            replacements.pop("{{documents_list}}", None)

    replace_placeholders_robust(doc, expand_placeholder_map(replacements))

    # Вставляем блок "осознанности" (частичные оплаты, гарантийные письма)
    awareness_block = replacements.get("{awareness_block}", "")
    if awareness_block and awareness_block.strip():
        insert_awareness_block(doc, awareness_block)

    legal_block = replacements.get("{legal_fees_block}", "")
    if not str(legal_block).strip():
        remove_legal_fees_section(doc)

    replace_attachments_with_paragraphs(
        doc,
        attachments or [],
        use_claim_extras=False
    )
    number_attachments_section(doc)
    documents_count = None
    groups = data.get("pretension_groups") or data.get("document_groups")
    if isinstance(groups, list) and groups:
        documents_count = len(groups)
    adjust_pretension_quality_section(doc, documents_count)
    insert_pretension_interest_table(
        doc,
        interest_data.get("detailed_calc", []),
        interest_data.get("total_interest")
    )
    proofread_docx_document(doc, protected_values=proofread_protected_values)
    enforce_times_new_roman(doc)

    if output_path is None:
        output_name = f"Претензия_{uuid.uuid4().hex}.docx"
        output_path = os.path.join(os.path.dirname(__file__), output_name)
    doc.save(output_path)
    return output_path


async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """
    Обработчик команды /start.

    Args:
        update: Объект обновления Telegram
        context: Контекст бота
    """
    if update.effective_user:
        logging.info(
            f"Received /start command from user {update.effective_user.id}"
        )
    if update.message:
        keyboard = [
            [
                InlineKeyboardButton(
                    "🧾 Исковое заявление",
                    callback_data="flow_claim"
                ),
                InlineKeyboardButton(
                    "📄 Претензия",
                    callback_data="flow_pretension"
                ),
            ],
            [
                InlineKeyboardButton(
                    "📋 Иск (внешняя претензия)",
                    callback_data="flow_external_claim"
                ),
            ]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        await update.message.reply_text(
            "Что нужно составить?",
            reply_markup=reply_markup
        )
    return ASK_FLOW


async def flow_chosen(update, context):
    query = update.callback_query
    await query.answer()
    choice = query.data

    if choice == "flow_claim":
        context.user_data.clear()
        context.user_data["flow"] = "claim"
        await query.edit_message_text(
            "Отправь .docx файл с досудебным требованием — "
            "я верну исковое заявление в формате Word."
        )
        return ASK_DOCUMENT

    if choice == "flow_pretension":
        context.user_data.clear()
        context.user_data["flow"] = "pretension"
        await query.edit_message_text(
            "Отправь PDF-файл с документами по перевозке. "
            "Я подготовлю претензию в формате Word."
        )
        return ASK_DOCUMENT

    if choice == "flow_external_claim":
        context.user_data.clear()
        context.user_data["flow"] = "external_claim"
        context.user_data["external_claim_files"] = []
        context.user_data["external_claim_stage"] = "claim"  # claim, docs, legal
        await query.edit_message_text(
            "📋 Режим: Исковое заявление на основе внешней претензии.\n\n"
            "Сначала отправьте PDF-файл с претензией от клиента.\n"
            "(Это основной документ с требованиями к должнику)"
        )
        return ASK_EXTERNAL_CLAIM_DOCUMENT

    await query.edit_message_text(
        "Не удалось определить выбор. Попробуй снова /start."
    )
    return ConversationHandler.END


async def ask_claim_status(update, context):
    keyboard = [
        [
            InlineKeyboardButton("✅ Да", callback_data='claim_received'),
            InlineKeyboardButton("❌ Нет", callback_data='claim_not_received'),
        ]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text(
        "Ответчик получил требование?",
        reply_markup=reply_markup
    )
    return ASK_CLAIM_STATUS


async def claim_status_chosen(update, context):
    query = update.callback_query
    await query.answer()
    context.user_data['claim_status'] = query.data

    await query.edit_message_text("Введите дату отправления (ДД.ММ.ГГГГ):")
    return ASK_SEND_DATE


async def ask_send_date(update, context):
    context.user_data['claim_date'] = update.message.text.strip()
    if context.user_data.get('claim_status') == 'claim_received':
        await update.message.reply_text("Введите трек-номер отправления:")
        return ASK_TRACK
    else:
        await update.message.reply_text("Введите трек-номер отправления:")
        return ASK_TRACK


async def ask_track(update, context):
    raw_track = update.message.text.strip()
    track_number = normalize_tracking_number(raw_track)

    if context.user_data.get('use_tracking_api'):
        if not is_valid_tracking_number(track_number):
            await update.message.reply_text(
                "Трек-номер некорректен. Введите номер из 10-20 цифр "
                "или формат S10 (например RA123456789RU)."
            )
            return ASK_TRACK
        try:
            records = fetch_russian_post_operations(track_number)
        except RussianPostTrackingError as exc:
            await update.message.reply_text(
                f"Не удалось получить данные по трек-номеру. {exc} "
                "Проверьте номер и попробуйте снова."
            )
            return ASK_TRACK

        send_date, receive_date = extract_tracking_dates(records)
        if not send_date:
            await update.message.reply_text(
                "Не удалось определить дату отправления по треку. "
                "Проверьте номер и попробуйте снова."
            )
            return ASK_TRACK

        context.user_data['claim_number'] = track_number
        context.user_data['claim_date'] = send_date
        context.user_data['postal_receive_date'] = receive_date or ''
        context.user_data['claim_status'] = (
            'claim_received' if receive_date else 'claim_not_received'
        )

        if receive_date:
            await update.message.reply_text(
                f"Найдены даты: отправление {send_date}, получение {receive_date}."
            )
        else:
            await update.message.reply_text(
                f"Найдена дата отправления {send_date}. "
                "Вручение адресату не найдено."
            )

        await finish_claim(update, context)
        return ConversationHandler.END

    context.user_data['claim_number'] = raw_track
    if context.user_data.get('claim_status') == 'claim_received':
        await update.message.reply_text("Введите дату получения (ДД.ММ.ГГГГ):")
        return ASK_RECEIVE_DATE
    else:
        await finish_claim(update, context)
        return ConversationHandler.END


async def ask_receive_date(update, context):
    context.user_data['postal_receive_date'] = update.message.text.strip()
    await finish_claim(update, context)
    return ConversationHandler.END


async def ask_birth_date(update, context):
    if not update.message:
        return ConversationHandler.END
    value = update.message.text.strip()
    if value.lower() == 'пропустить':
        context.user_data['skip_birth_info'] = True
        await finish_claim(update, context)
        return ConversationHandler.END
    context.user_data['plaintiff_birth_date'] = value
    await update.message.reply_text(
        "Укажите место рождения истца или «пропустить»:"
    )
    return ASK_BIRTH_PLACE


async def ask_birth_place(update, context):
    if not update.message:
        return ConversationHandler.END
    value = update.message.text.strip()
    if value.lower() == 'пропустить':
        context.user_data['plaintiff_birth_place'] = ''
    else:
        context.user_data['plaintiff_birth_place'] = value
    await finish_claim(update, context)
    return ConversationHandler.END


async def finish_claim(update, context):
    file_path = context.user_data.get('file_path')
    logging.info(
        "Trying to process file_path from user_data: %s",
        file_path
    )
    if not file_path or not os.path.exists(file_path):
        await update.message.reply_text(
            f'Ошибка: файл {file_path} не найден на диске.'
        )
        logging.error(
            "File for processing not found: %s",
            file_path
        )
        return

    # Извлекаем текст из документа для нового парсера
    doc = Document(file_path)
    text = "\n".join(p.text for p in doc.paragraphs)

    # Используем sliding window парсер
    claim_data = parse_documents_with_sliding_window(text)
    claim_data = apply_llm_fallback(text, claim_data)

    claim_data['claim_number'] = context.user_data.get('claim_number', '')
    claim_data['claim_date'] = context.user_data.get('claim_date', '')
    key_rates = get_key_rates_from_395gk()
    try:
        interest_data = calculate_full_395(
            file_path, key_rates=key_rates
        )
    except Exception as exc:
        logging.error(
            "Ошибка расчета процентов: %s",
            exc,
            exc_info=True
        )
        interest_data = {
            'total_interest': 0.0,
            'detailed_calc': [],
            'error': str(exc)
        }
    if interest_data.get('error') and update.message:
        await update.message.reply_text(
            "⚠️ Не удалось найти таблицу расчета процентов. "
            "Продолжаю без процентов."
        )

    # Получаем сумму долга из нового парсера
    debt_amount = parse_amount(claim_data.get('debt', '0'))
    debt_decimal = parse_amount_decimal(claim_data.get('debt', '0'))
    debt_rubles, debt_kopeks = split_rubles_kopeks(debt_decimal)
    debt_full = format_money_ru(debt_amount, 2)
    total_interest = parse_amount(interest_data.get('total_interest', 0.0))
    total_claim = debt_amount + total_interest

    duty_data = calculate_duty(total_claim)
    if 'error' in duty_data:
        await update.message.reply_text(str(duty_data['error']))
        return
    for key in ['claim_status', 'claim_date', 'claim_number', 'postal_receive_date']:
        if key not in context.user_data:
            context.user_data[key] = ''

    # Используем данные из нового парсера для истца и ответчика
    plaintiff_name = normalize_str(claim_data.get('plaintiff_name'))
    defendant_name = normalize_str(claim_data.get('defendant_name'))
    contract_parties = claim_data.get('contract_parties', '')
    contract_parties_short = claim_data.get('contract_parties_short', '')

    # Проверяем, является ли истец ИП
    is_plaintiff_ip = 'ИП' in plaintiff_name or 'Индивидуальный предприниматель' in plaintiff_name
    is_defendant_ip = 'ИП' in defendant_name or 'Индивидуальный предприниматель' in defendant_name

    birth_info_value = normalize_str(
        claim_data.get('plaintiff_birth_info'),
        default=''
    )
    birth_date = normalize_str(
        context.user_data.get('plaintiff_birth_date'),
        default=''
    )
    birth_place = normalize_str(
        context.user_data.get('plaintiff_birth_place'),
        default=''
    )
    if not birth_info_value and birth_date:
        birth_info_value = f"Дата рождения {birth_date}"
        if birth_place:
            birth_info_value += f"\nМесто рождения {birth_place}"

    if (
        is_plaintiff_ip
        and not birth_info_value
        and not context.user_data.get('skip_birth_info')
    ):
        target = None
        if getattr(update, "message", None):
            target = update.message
        elif getattr(update, "callback_query", None):
            target = update.callback_query.message
        if target:
            await target.reply_text(
                "Укажите дату рождения истца (ДД.ММ.ГГГГ) или «пропустить»:"
            )
            return ASK_BIRTH_DATE

    # Форматируем имена для использования в тексте (короткие названия)
    plaintiff_name_short = format_organization_name_short(plaintiff_name)
    defendant_name_short = format_organization_name_short(defendant_name)
    plaintiff_ogrn_type = get_ogrn_label(
        plaintiff_name,
        claim_data.get('plaintiff_inn', '')
    )

    # Получаем информацию о подсудности из контекста
    jurisdiction_info = context.user_data.get('jurisdiction_info')
    if jurisdiction_info:
        court_name, court_address = resolve_court_from_dadata(
            jurisdiction_info.court_name,
            jurisdiction_info.court_address
        )
    else:
        # Fallback на старую логику
        court_name, court_address = get_court_by_address(
            claim_data.get('defendant_address', 'Не указано')
        )

    legal_fees_value = parse_amount(claim_data.get('legal_fees', '0'))
    docs_track_number = join_list_values(
        claim_data.get('postal_numbers', [])
    ) or context.user_data.get('claim_number', '')
    docs_received_date = join_list_values(
        claim_data.get('postal_dates', [])
    ) or context.user_data.get('postal_receive_date', '')
    legal_contract_date = add_prefix_if_missing(
        normalize_str(claim_data.get('legal_contract_date')),
        'от '
    )
    documents_list_structured = extract_documents_list_structure(text)
    documents_list = build_documents_list(claim_data)
    plaintiff_birth_info = birth_info_value if is_plaintiff_ip else ''
    signatory_value = normalize_str(
        claim_data.get('signatory'),
        default=''
    ).replace('\n', ' ').strip()
    signature_block_value = normalize_str(
        claim_data.get('signature_block'),
        default=''
    )
    plaintiff_address_value = normalize_str(
        claim_data.get('plaintiff_address')
    ).replace('\n', ' ').strip()
    defendant_address_value = normalize_str(
        claim_data.get('defendant_address')
    ).replace('\n', ' ').strip()
    plaintiff_address_value = maybe_proofread_text(
        plaintiff_address_value,
        protected_values=[plaintiff_name, plaintiff_name_short]
    )
    defendant_address_value = maybe_proofread_text(
        defendant_address_value,
        protected_values=[defendant_name, defendant_name_short]
    )
    claim_paragraph_value = generate_claim_paragraph(context.user_data)
    claim_paragraph_value = maybe_proofread_text(
        claim_paragraph_value,
        protected_values=[plaintiff_name_short, defendant_name_short]
    )
    payment_terms_value = generate_payment_terms(claim_data)
    payment_terms_value = maybe_proofread_text(payment_terms_value)

    replacements = {
        '{claim_paragraph}': claim_paragraph_value,
        '{postal_block}': format_document_list(
            claim_data.get('postal_block', 'Не указано')
        ),
        '{postal_numbers_all}': (
            ', '.join(claim_data.get('postal_numbers', []))
            or 'Не указано'
        ),
        '{postal_dates_all}': (
            ', '.join(claim_data.get('postal_dates', []))
            or 'Не указано'
        ),
        '{court_name}': court_name,
        '{court_address}': court_address,
        '{plaintiff_name}': plaintiff_name,
        '{plaintiff_name_short}': plaintiff_name_short,
        '{plaintiff_name_formatted}': plaintiff_name_short,
        '{plaintiff_inn}': normalize_str(claim_data.get('plaintiff_inn')),
        '{plaintiff_kpp}': '' if is_plaintiff_ip else normalize_str(claim_data.get('plaintiff_kpp')),
        '{plaintiff_ogrn}': normalize_str(claim_data.get('plaintiff_ogrn')),
        '{plaintiff_address}': plaintiff_address_value,
        '{defendant_name}': defendant_name,
        '{defendant_name_short}': defendant_name_short,
        '{defendant_inn}': normalize_str(claim_data.get('defendant_inn')),
        '{defendant_kpp}': '' if is_defendant_ip else normalize_str(claim_data.get('defendant_kpp')),
        '{defendant_ogrn}': normalize_str(claim_data.get('defendant_ogrn')),
        '{defendant_address}': defendant_address_value,
        '{contract_parties}': contract_parties,
        '{contract_parties_short}': contract_parties_short,
        '{total_claim}': format_money_ru(total_claim, 2),
        '{claim_total}': format_money_ru(total_claim, 2),
        '{duty}': f"{duty_data['duty']:,.0f}".replace(',', ' '),
        '{debt}': debt_full,
        '{debt_amount}': debt_rubles,
        '{debt_kopeks}': debt_kopeks,
        '{payment_terms}': payment_terms_value,
        '{contracts}': format_document_list(claim_data.get('contracts', '')),
        '{contract_applications}': format_document_list(claim_data.get('contract_applications', '')),
        '{cargo_docs}': format_document_list(claim_data.get('cargo_docs', '')),
        '{invoice_blocks}': format_document_list(claim_data.get('invoice_blocks', '')),
        '{upd_blocks}': format_document_list(claim_data.get('upd_blocks', '')),
        '{invoices}': claim_data.get('invoice_blocks', 'Не указано'),
        '{upds}': claim_data.get('upd_blocks', 'Не указано'),
        '{claim_date}': normalize_str(context.user_data.get('claim_date', '')),
        '{claim_number}': normalize_str(context.user_data.get('claim_number', '')),
        '{claim_track_number}': normalize_str(
            context.user_data.get('claim_number', '')
        ),
        '{docs_track_number}': normalize_str(docs_track_number),
        '{docs_received_date}': normalize_str(docs_received_date),
        '{documents_list}': documents_list,
        '{total_interest}': format_money_ru(total_interest, 2),
        '{legal_fees}': format_money(legal_fees_value, 0),
        '{legal_fee}': format_money(legal_fees_value, 0),
        '{legal_contract_number}': normalize_str(
            claim_data.get('legal_contract_number')
        ),
        '{legal_contract_date}': legal_contract_date,
        '{legal_payment_number}': normalize_str(
            claim_data.get('legal_payment_number')
        ),
        '{legal_payment_date}': normalize_str(
            claim_data.get('legal_payment_date')
        ),
        '{total_expenses}': (
            f"{float(str(duty_data['duty'])) + legal_fees_value:,.0f}"
            .replace(',', ' ')
        ),
        '{calculation_date}': datetime.today().strftime('%d.%m.%Y г.'),
        '{signatory}': signatory_value,
        '{signature_block}': signature_block_value,
        '{postal_numbers}': normalize_str(
            context.user_data.get('claim_number', '')
        ),
        '{postal_receive_date}': normalize_str(
            context.user_data.get('postal_receive_date', '')
        ),
        '{payment_days}': claim_data.get('payment_days', 'Не указано'),
        '{plaintiff_ogrn_type}': plaintiff_ogrn_type,
        '{plaintiff_birth_info}': plaintiff_birth_info,
    }
    result_docx = create_isk_document(
        claim_data,
        interest_data,
        duty_data,
        replacements,
        documents_list_structured=documents_list_structured,
        proofread_protected_values=[
            plaintiff_name,
            defendant_name,
            plaintiff_name_short,
            defendant_name_short,
        ],
    )
    f107_path = None
    poa_path = None
    try:
        defendant_display_name = resolve_defendant_display_name(
            defendant_name_short,
            defendant_name
        )
        f107_items = build_f107_items(
            claim_data.get('attachments', []),
            defendant_display_name
        )
        sender_name = normalize_str(
            claim_data.get('signatory'),
            default=''
        )
        if sender_name == 'Не указано':
            sender_name = ''
        if not sender_name:
            sender_name = (
                plaintiff_name_short
                if plaintiff_name_short != 'Не указано'
                else ''
            )
        sender_company = (
            plaintiff_name
            if plaintiff_name != 'Не указано'
            else plaintiff_name_short
        )
        f107_path = create_f107_document(
            f107_items,
            sender_name,
            sender_company
        )
    except FileNotFoundError as exc:
        logging.warning("Не удалось сформировать Ф107: %s", exc)
    except Exception as exc:
        logging.error("Ошибка формирования Ф107: %s", exc, exc_info=True)
    try:
        poa_replacements = {
            '{poa_date}': format_poa_date(),
            '{plaintiff_name}': normalize_str(plaintiff_name),
            '{plaintiff_inn}': normalize_str(claim_data.get('plaintiff_inn')),
            '{plaintiff_ogrn}': normalize_str(claim_data.get('plaintiff_ogrn')),
            '{plaintiff_address}': normalize_str(
                claim_data.get('plaintiff_address')
            ).replace('\n', ' ').strip(),
        }
        poa_path = create_power_of_attorney_document(poa_replacements)
    except FileNotFoundError as exc:
        logging.warning("Не удалось сформировать доверенность: %s", exc)
    except Exception as exc:
        logging.error(
            "Ошибка формирования доверенности: %s",
            exc,
            exc_info=True
        )
    with open(result_docx, 'rb') as f:
        await update.message.reply_document(
            InputFile(f, filename="Исковое_заявление.docx"),
            caption="Исковое заявление по ст. 395 ГК РФ"
        )
    if f107_path:
        with open(f107_path, 'rb') as f:
            await update.message.reply_document(
                InputFile(f, filename="Опись_вложения_F107.docx"),
                caption="Опись вложения (форма Ф107)"
            )
    if poa_path:
        with open(poa_path, 'rb') as f:
            await update.message.reply_document(
                InputFile(f, filename="Доверенность.docx"),
                caption="Доверенность на представителя"
            )
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
        if os.path.exists(result_docx):
            os.remove(result_docx)
        if f107_path and os.path.exists(f107_path):
            os.remove(f107_path)
        if poa_path and os.path.exists(poa_path):
            os.remove(poa_path)
    except Exception as e:
        logging.warning(
            f"Не удалось удалить файл {file_path}: {e}"
        )


async def ask_jurisdiction(update, context):
    """
    Спрашивает пользователя о подсудности спора.
    """
    from jurisdiction import JurisdictionDetector, format_jurisdiction_for_user

    file_path = context.user_data.get('file_path')
    if not file_path or not os.path.exists(file_path):
        await update.message.reply_text('Ошибка: файл не найден.')
        return ConversationHandler.END

    # Извлекаем текст из документа
    doc = Document(file_path)
    text = "\n".join(p.text for p in doc.paragraphs)

    # Парсим базовые данные для определения адресов
    claim_data = parse_documents_with_sliding_window(text)
    claim_data = apply_llm_fallback(text, claim_data)
    defendant_address = claim_data.get('defendant_address', 'Не указано')

    # Определяем подсудность
    detector = JurisdictionDetector()
    jurisdiction_info = detector.detect_jurisdiction(
        text=text,
        defendant_address=defendant_address
    )

    # Сохраняем в контекст
    context.user_data['jurisdiction_info'] = jurisdiction_info
    context.user_data['claim_data'] = claim_data

    # Формируем сообщение для пользователя
    info_text = format_jurisdiction_for_user(jurisdiction_info)

    # Кнопки для выбора
    keyboard = []

    if jurisdiction_info.confidence > 0.7:
        # Высокая уверенность - предлагаем подтвердить
        keyboard.append([
            InlineKeyboardButton("✅ Верно", callback_data='jurisdiction_confirm')
        ])

    keyboard.extend([
        [InlineKeyboardButton("📝 Указать другой суд", callback_data='jurisdiction_custom')],
        [InlineKeyboardButton("❓ По месту ответчика (по умолчанию)", callback_data='jurisdiction_default')]
    ])

    reply_markup = InlineKeyboardMarkup(keyboard)

    message = (
        "🏛 *Определение подсудности*\n\n"
        f"{info_text}\n\n"
        "Подсудность определена верно?"
    )

    await update.message.reply_text(
        message,
        reply_markup=reply_markup,
        parse_mode='Markdown'
    )

    return ASK_JURISDICTION


async def jurisdiction_chosen(update, context):
    """
    Обрабатывает выбор пользователя по подсудности.
    """
    query = update.callback_query
    await query.answer()

    choice = query.data

    if choice == 'jurisdiction_confirm':
        # Подтверждено - переходим к вопросам о претензии
        await query.edit_message_text(
            "✅ Подсудность подтверждена.\n\nТеперь ответьте на вопросы о претензии."
        )
        return await ask_claim_status_after_jurisdiction(query, context)

    elif choice == 'jurisdiction_default':
        # Используем подсудность по умолчанию
        from jurisdiction import JurisdictionDetector

        claim_data = context.user_data.get('claim_data', {})
        defendant_address = claim_data.get('defendant_address', 'Не указано')

        detector = JurisdictionDetector()
        jurisdiction_info = detector._get_default_jurisdiction(defendant_address)
        context.user_data['jurisdiction_info'] = jurisdiction_info

        await query.edit_message_text(
            f"✅ Используется подсудность по месту ответчика.\n\n"
            f"Суд: {jurisdiction_info.court_name}\n\n"
            "Теперь ответьте на вопросы о претензии."
        )
        return await ask_claim_status_after_jurisdiction(query, context)

    elif choice == 'jurisdiction_custom':
        # Запрашиваем ручной ввод суда
        await query.edit_message_text(
            "Введите название суда (например: Арбитражный суд Московской области):"
        )
        return ASK_CUSTOM_COURT


async def handle_custom_court(update, context):
    """
    Обрабатывает ручной ввод названия суда.
    """
    from jurisdiction import JurisdictionDetector, JurisdictionInfo, JurisdictionType

    custom_court = update.message.text.strip()
    detector = JurisdictionDetector()

    # Сначала пробуем DaData справочник судов
    suggestion = fetch_dadata_court_suggest(custom_court, court_type="AS")
    if suggestion:
        parsed = parse_dadata_court(suggestion)
        court_name = parsed.get("name") or custom_court
        court_address = parsed.get("address") or "Уточните адрес суда"
        jurisdiction_info = JurisdictionInfo(
            type=JurisdictionType.CUSTOM,
            court_name=court_name,
            court_address=court_address,
            confidence=1.0
        )
        context.user_data['jurisdiction_info'] = jurisdiction_info

        await update.message.reply_text(
            f"✅ Суд установлен по DaData:\n{court_name}\n\n"
            "Теперь ответьте на вопросы о претензии."
        )
        return await ask_claim_status_after_jurisdiction(update, context)

    # Пытаемся найти суд в базе
    region = custom_court.replace('Арбитражный суд', '').strip()
    court_info_dict = detector._find_court_by_region(region)

    if court_info_dict:
        jurisdiction_info = JurisdictionInfo(
            type=JurisdictionType.CUSTOM,
            court_name=court_info_dict['name'],
            court_address=court_info_dict['address'],
            confidence=1.0
        )
        context.user_data['jurisdiction_info'] = jurisdiction_info

        await update.message.reply_text(
            f"✅ Суд установлен:\n{court_info_dict['name']}\n\n"
            "Теперь ответьте на вопросы о претензии."
        )
    else:
        # Не нашли в базе - сохраняем как есть
        jurisdiction_info = JurisdictionInfo(
            type=JurisdictionType.CUSTOM,
            court_name=custom_court,
            court_address="Уточните адрес суда",
            confidence=0.5
        )
        context.user_data['jurisdiction_info'] = jurisdiction_info

        await update.message.reply_text(
            f"⚠️ Суд не найден в базе. Использую введенное название:\n{custom_court}\n\n"
            "Не забудьте проверить адрес суда в готовом документе!\n\n"
            "Теперь ответьте на вопросы о претензии."
        )

    return await ask_claim_status_after_jurisdiction(update, context)


async def ask_claim_status_after_jurisdiction(update_or_query, context):
    """
    Переход к вопросам о претензии после определения подсудности.
    """
    config = get_russian_post_config()
    if config.get("enabled"):
        context.user_data['use_tracking_api'] = True
        if hasattr(update_or_query, 'message') and update_or_query.message:
            await update_or_query.message.reply_text(
                "Введите трек-номер отправления претензии. "
                "Я сам определю дату отправки и получения."
            )
        else:
            await update_or_query.message.reply_text(
                "Введите трек-номер отправления претензии. "
                "Я сам определю дату отправки и получения."
            )
        return ASK_TRACK

    context.user_data['use_tracking_api'] = False
    keyboard = [
        [
            InlineKeyboardButton("✅ Да", callback_data='claim_received'),
            InlineKeyboardButton("❌ Нет", callback_data='claim_not_received'),
        ]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)

    # Определяем, откуда пришел вызов
    if hasattr(update_or_query, 'message') and update_or_query.message:
        # Это обычный update с message
        await update_or_query.message.reply_text(
            "Ответчик получил требование?",
            reply_markup=reply_markup
        )
    else:
        # Это callback query
        await update_or_query.message.reply_text(
            "Ответчик получил требование?",
            reply_markup=reply_markup
        )

    return ASK_CLAIM_STATUS


async def handle_docx_entry(update, context):
    try:
        if update.effective_user:
            logging.info(
                "Received document from user %s",
                update.effective_user.id
            )
        if not update.message:
            return
        doc = update.message.document
        if not doc or not doc.file_name:
            logging.warning(
                "Invalid document"
            )
            await update.message.reply_text(
                'Пожалуйста, отправь файл Word (.docx).'
            )
            return
        if not doc.file_name.lower().endswith('.docx'):
            logging.warning(
                "Invalid file format"
            )
            await update.message.reply_text(
                'Пожалуйста, отправь файл Word (.docx).'
            )
            return
        os.makedirs('uploads', exist_ok=True)
        unique_name = f"{uuid.uuid4()}_{doc.file_name}"
        file_path = os.path.join('uploads', unique_name)
        telegram_file = await doc.get_file()
        await telegram_file.download_to_drive(file_path)
        logging.info(
            "Downloaded file: %s",
            file_path
        )
        context.user_data['file_path'] = file_path
        logging.info(
            "Saved file_path in user_data: %s",
            file_path
        )
        # ИЗМЕНЕНО: теперь сначала спрашиваем о подсудности
        return await ask_jurisdiction(update, context)
    except Exception as e:
        import traceback
        tb = traceback.format_exc()
        logging.error(f"Ошибка: {e}\n{tb}")
        if update.message:
            await update.message.reply_text(
                f'Ошибка обработки: {e}. Проверьте формат файла.'
            )


def is_missing_value(value: Any) -> bool:
    if value is None:
        return True
    if isinstance(value, list):
        return len(value) == 0
    text = str(value).strip()
    return not text or text == "Не указано"


def get_pretension_missing_fields(data: Dict[str, Any]) -> List[str]:
    missing = []
    shipping_method = normalize_shipping_method(data.get("shipping_method"))
    shipments = data.get("shipments") or []
    api_enabled = get_russian_post_config().get("enabled")
    skip_postal_lookup = bool(data.get("skip_postal_lookup"))
    groups = data.get("pretension_groups") or []

    for key in PRETENSION_FIELD_ORDER:
        value = data.get(key)
        if key == "debt":
            if parse_amount(value, 0) <= 0:
                missing.append(key)
            continue
        if key == "payment_days":
            if groups:
                group_days = []
                for group in groups:
                    raw_days = group.get("payment_days")
                    if raw_days is None:
                        group_days.append(0)
                        continue
                    try:
                        group_days.append(int(raw_days))
                    except (TypeError, ValueError):
                        group_days.append(0)
                if group_days and all(days > 0 for days in group_days):
                    continue
            try:
                days = int(re.sub(r"[^\d]", "", str(value)))
            except ValueError:
                days = 0
            if days <= 0:
                missing.append(key)
            continue
        if key == "shipping_method":
            if shipments:
                continue
            if shipping_method not in ("почта", "сдэк"):
                missing.append(key)
            continue
        if key == "docs_track_number":
            if shipments:
                continue
            if shipping_method == "почта" and api_enabled and not skip_postal_lookup:
                dates_raw = str(data.get("docs_received_date") or "")
                has_dates = bool(re.search(r"\d{2}\.\d{2}\.\d{4}", dates_raw))
                if not has_dates and is_missing_value(value):
                    missing.append(key)
            continue
        if key == "docs_received_date":
            if shipments:
                continue
            if shipping_method == "почта" and api_enabled and not skip_postal_lookup:
                continue
            dates_raw = str(value or "")
            if not re.search(r"\d{2}\.\d{2}\.\d{4}", dates_raw):
                missing.append(key)
            continue
        if is_missing_value(value):
            missing.append(key)
    return missing


async def ask_next_pretension_field(update, context):
    missing = context.user_data.get("pretension_missing_fields", [])
    if not missing:
        return await finish_pretension(update, context)

    key = missing[0]
    prompt = PRETENSION_FIELD_DEFS.get(key, {}).get("prompt")
    if not prompt:
        missing.pop(0)
        context.user_data["pretension_missing_fields"] = missing
        return await ask_next_pretension_field(update, context)

    target = None
    if getattr(update, "message", None):
        target = update.message
    elif getattr(update, "callback_query", None):
        target = update.callback_query.message
    if target:
        await target.reply_text(prompt)
    return ASK_PRETENSION_FIELD


async def handle_pretension_document(update, context):
    if update.message and update.message.document:
        doc = update.message.document
        if not doc.file_name or not doc.file_name.lower().endswith('.pdf'):
            await update.message.reply_text('Пожалуйста, отправь PDF-файл.')
            return ASK_DOCUMENT

        os.makedirs('uploads', exist_ok=True)
        unique_name = f"{uuid.uuid4()}_{doc.file_name}"
        file_path = os.path.join('uploads', unique_name)
        telegram_file = await doc.get_file()
        await telegram_file.download_to_drive(file_path)

        files = context.user_data.get("pretension_files", [])
        files.append({"path": file_path, "name": doc.file_name})
        context.user_data["pretension_files"] = files

        await update.message.reply_text(
            "Файл получен. Если есть еще PDF — отправьте. "
            "Если всё готово, напишите «готово»."
        )
        return ASK_DOCUMENT

    if not update.message or not update.message.text:
        await update.message.reply_text('Пожалуйста, отправь PDF-файл.')
        return ASK_DOCUMENT

    if update.message.text.strip().lower() not in {"готово", "done", "finish", "все"}:
        await update.message.reply_text(
            "Отправьте PDF-файлы или напишите «готово», когда закончите."
        )
        return ASK_DOCUMENT

    files = context.user_data.get("pretension_files", [])
    if not files:
        await update.message.reply_text(
            "Сначала отправьте один или несколько PDF-файлов."
        )
        return ASK_DOCUMENT

    combined_texts = []
    all_pages: List[str] = []
    low_pages_info = []
    for entry in files:
        try:
            pages, low_text_pages = extract_pdf_pages(entry["path"])
        except Exception as exc:
            await update.message.reply_text(
                f"Не удалось прочитать PDF {entry['name']}: {exc}"
            )
            return ASK_DOCUMENT
        processed_low_pages: List[int] = []
        if low_text_pages:
            ocr_pages = apply_vision_ocr_to_pages(
                entry["path"],
                pages,
                low_text_pages
            )
            if ocr_pages:
                processed_low_pages = ocr_pages
                await update.message.reply_text(
                    "✅ Ollama Vision OCR применён к страницам: "
                    + ", ".join(str(page) for page in ocr_pages)
                )

        # Дополнительный OCR по ключевым документам (заявки, накладные, счета, УПД, почтовые квитанции)
        config = get_vision_config()
        max_pages = int(config.get("max_pages") or 0)
        targeted_pages = collect_targeted_ocr_pages(pages, processed_low_pages)
        if targeted_pages:
            if max_pages > 0:
                remaining = max_pages - len(set(processed_low_pages))
                if remaining <= 0:
                    targeted_pages = []
                else:
                    targeted_pages = targeted_pages[:remaining]
            if targeted_pages:
                ocr_pages = apply_vision_ocr_to_pages(
                    entry["path"],
                    pages,
                    targeted_pages
                )
                if ocr_pages:
                    await update.message.reply_text(
                        "✅ OCR для ключевых документов применён к страницам: "
                        + ", ".join(str(page) for page in ocr_pages)
                    )

        # Vision-экстракция для нетипичных/уникальных документов
        scan_limit_raw = os.getenv("VISION_DOC_SCAN_PAGES", "2")
        try:
            scan_limit = int(scan_limit_raw)
        except ValueError:
            scan_limit = 2
        if scan_limit > 0:
            vision_pages = collect_vision_doc_pages(
                pages,
                processed_low_pages,
                limit=scan_limit
            )
            if vision_pages:
                processed = apply_vision_document_extraction(
                    entry["path"],
                    pages,
                    vision_pages
                )
                if processed:
                    await update.message.reply_text(
                        "✅ Vision-анализ документов применён к страницам: "
                        + ", ".join(str(page) for page in processed)
                    )
        all_pages.extend(pages)
        page_blocks = []
        for idx, page in enumerate(pages, start=1):
            page_blocks.append(f"[Страница {idx}]\n{page}")
        combined_texts.append(
            f"=== {entry['name']} ===\n" + "\n\n".join(page_blocks)
        )
        if low_text_pages:
            low_pages_info.append((entry["path"], entry["name"], low_text_pages))

    if low_pages_info:
        for file_path, name, pages in low_pages_info:
            pages_list = ", ".join(str(page) for page in pages)
            await update.message.reply_text(
                "⚠️ Текст распознан плохо в файле "
                f"{name} на страницах: {pages_list}. "
                "Возможно, часть данных придется ввести вручную."
            )
            image_paths = render_pdf_pages(file_path, pages)
            if image_paths:
                await update.message.reply_text(
                    "Показываю страницы с плохим распознаванием (первые несколько):"
                )
                for path in image_paths:
                    with open(path, "rb") as handle:
                        await update.message.reply_photo(InputFile(handle))
                for path in image_paths:
                    try:
                        os.remove(path)
                    except OSError:
                        pass

    combined_text = "\n\n".join(combined_texts)
    claim_data = parse_documents_with_sliding_window(combined_text)
    claim_data = apply_llm_fallback(combined_text, claim_data)
    claim_data["document_groups"] = build_document_groups(
        combined_text,
        claim_data
    )
    claim_data["source_files"] = [entry.get("name") for entry in files if entry.get("name")]

    applications = extract_applications_from_pages(all_pages)
    invoices = extract_invoices_from_pages(all_pages)
    upd_docs = extract_upd_from_pages(all_pages)
    cargo_docs = extract_cargo_docs_from_pages(all_pages)

    # Vision LLM обогащение данных из cargo_docs при наличии low_pages_info
    if low_pages_info:
        cargo_docs = enrich_cargo_docs_with_vision(
            cargo_docs, files, low_pages_info
        )

    shipments = extract_cdek_shipments_from_pages(all_pages)
    shipments.extend(extract_postal_shipments_from_pages(all_pages))

    if not shipments:
        numbers = claim_data.get("postal_numbers") or []
        dates = claim_data.get("postal_dates") or []
        if numbers:
            shipping_method = normalize_shipping_method(
                claim_data.get("shipping_method")
            ) or "почта"
            claim_data["shipping_method"] = shipping_method
            source = normalize_shipping_source(shipping_method) or "post"
            for idx, raw_number in enumerate(numbers):
                track_number = normalize_tracking_number(str(raw_number))
                if not track_number:
                    continue
                date_str = dates[idx] if idx < len(dates) else ""
                shipments.append({
                    "track_number": track_number,
                    "received_date": parse_date_str(date_str),
                    "received_date_str": date_str,
                    "source": source,
                })

    if not claim_data.get("shipping_method"):
        for page in all_pages:
            lower = page.lower()
            if "сдэк" in lower or "cdek" in lower:
                claim_data["shipping_method"] = "сдэк"
                break
            if any(
                token in lower
                for token in (
                    "почта россии",
                    "отчет об отслеживании",
                    "квитанц",
                    "рпо",
                    "идентификатор",
                    "трек",
                )
            ):
                claim_data["shipping_method"] = "почта"
                break

    if shipments:
        config = get_russian_post_config()
        if config.get("enabled") and not claim_data.get("skip_postal_lookup"):
            for shipment in shipments:
                if normalize_shipping_source(shipment.get("source")) != "post":
                    continue
                track_number = shipment.get("track_number") or ""
                if not is_valid_tracking_number(track_number):
                    continue
                try:
                    records = fetch_russian_post_operations(track_number)
                    shipment["api_records"] = len(records)
                    send_date, receive_date = extract_tracking_dates(records)
                    if receive_date:
                        shipment["received_date"] = parse_date_str(receive_date) or receive_date
                        shipment["received_date_str"] = receive_date
                    if send_date:
                        shipment["send_date"] = send_date
                except Exception as exc:
                    logger.warning(
                        f"Не удалось получить данные по треку {track_number}: {exc}"
                    )
    cargo_assignment_preview = assign_cargo_to_applications(applications, cargo_docs)
    matching_warnings = get_matching_warnings(cargo_assignment_preview)
    if matching_warnings:
        await update.message.reply_text(
            "⚠️ Есть сомнительные сопоставления документов с заявками. "
            "Проверьте их перед отправкой претензии:"
        )
        for warning in matching_warnings:
            await update.message.reply_text(warning)

    payment_terms_by_application = extract_application_payment_terms(
        all_pages,
        applications
    )
    groups = build_pretension_groups(
        applications,
        invoices,
        cargo_docs,
        upd_docs=upd_docs,
        payment_terms_by_application=payment_terms_by_application
    )
    assign_shipments_to_groups(groups, shipments)

    missing_doc_warnings = get_missing_document_warnings(groups)
    for warning in missing_doc_warnings:
        await update.message.reply_text(warning)

    payment_terms = None
    payment_days = None
    if payment_terms_by_application:
        terms_values = [
            normalize_payment_terms(item.get("terms") or "")
            for item in payment_terms_by_application.values()
            if item.get("terms")
        ]
        days_values = [
            item.get("days")
            for item in payment_terms_by_application.values()
            if item.get("days")
        ]
        unique_terms = {value for value in terms_values if value}
        unique_days = {value for value in days_values if value}
        if len(unique_terms) == 1:
            payment_terms = unique_terms.pop()
        if len(unique_days) == 1:
            payment_days = unique_days.pop()
    elif applications:
        app_pages = [page for page in all_pages if "заявк" in page.lower()]
        payment_terms, payment_days = extract_payment_terms_from_text(
            "\n".join(app_pages)
        )
    if not payment_terms:
        payment_terms, payment_days = extract_payment_terms_from_text(combined_text)

    if payment_terms:
        claim_data["payment_terms"] = payment_terms
    if payment_days:
        claim_data["payment_days"] = str(payment_days)

    parties = extract_parties_from_pages(all_pages)
    if parties:
        apply_extracted_parties(claim_data, parties)

    legal_docs = extract_legal_docs_from_pages(all_pages)
    if legal_docs:
        for key, value in legal_docs.items():
            if value and (is_missing_value(claim_data.get(key)) or key == "legal_fees"):
                claim_data[key] = value

    total_debt = sum(group.get("amount") or 0.0 for group in groups)
    if total_debt > 0:
        claim_data["debt"] = format_money(total_debt, 2)

    # Анализ "осознанности" документов: частичные оплаты, гарантийные письма и т.д.
    from decimal import Decimal
    original_debt_decimal = Decimal(str(total_debt)) if total_debt > 0 else None
    awareness_result = analyze_documents_for_special_cases(
        all_pages,
        original_debt=original_debt_decimal,
        use_llm=True
    )

    # Выводим предупреждения пользователю
    if awareness_result.warnings:
        warnings_text = "\n".join(awareness_result.warnings)
        await update.message.reply_text(
            f"📊 Результаты анализа документов:\n\n{warnings_text}"
        )

    # Корректируем данные иска при наличии особых случаев
    if (awareness_result.has_partial_payments or
        awareness_result.has_guarantee_letters or
        awareness_result.has_debt_acknowledgment):
        claim_data = adjust_claim_data(claim_data, awareness_result)
        # Если есть частичные оплаты - информируем о скорректированной сумме
        if awareness_result.has_partial_payments and awareness_result.adjusted_debt is not None:
            adjusted_str = format_money(float(awareness_result.adjusted_debt), 2)
            original_str = format_money(total_debt, 2)
            await update.message.reply_text(
                f"💰 Сумма долга скорректирована с учётом частичных оплат:\n"
                f"Исходная сумма: {original_str} руб.\n"
                f"Частичные оплаты: {format_money(float(awareness_result.total_partial_payments), 2)} руб.\n"
                f"Итого к взысканию: {adjusted_str} руб."
            )

    # Обработка актов сверки: строгая привязка оплат к заявкам
    reconciliation_entries, reconciliation_sales = extract_reconciliation_entries(
        all_pages
    )
    reconciliation_payments = [
        entry for entry in reconciliation_entries
        if entry.get("entry_type") == "payment"
    ]
    if reconciliation_payments:
        allocated, unassigned = match_reconciliation_payments_to_groups(
            groups,
            reconciliation_payments,
            sales=reconciliation_sales
        )
        if allocated:
            reconciled: List[Dict[str, Any]] = []
            allocated_total = 0.0
            for payment in allocated:
                date_value = payment.get("date")
                date_str = (
                    date_value.strftime("%d.%m.%Y")
                    if hasattr(date_value, "strftime") else str(date_value or "")
                )
                amount_value = float(payment.get("amount") or 0)
                allocated_total += amount_value
                reconciled.append({
                    "amount": str(amount_value),
                    "date": date_str,
                    "payment_number": payment.get("payment_number"),
                    "group_label": payment.get("group_label"),
                })
            if allocated_total > 0:
                adjusted_debt = max(total_debt - allocated_total, 0.0)
                claim_data["debt"] = format_money(adjusted_debt, 2)
            claim_data["partial_payments_info"] = reconciled
            claim_data["partial_payments_total"] = str(allocated_total)
        if not allocated:
            claim_data["partial_payments_info"] = []
            claim_data["partial_payments_total"] = "0"
            claim_data["debt"] = format_money(total_debt, 2)
        if unassigned:
            samples = []
            for payment in unassigned[:5]:
                date_value = payment.get("date")
                date_str = (
                    date_value.strftime("%d.%m.%Y")
                    if hasattr(date_value, "strftime") else str(date_value or "")
                )
                samples.append(
                    f"{date_str} — {format_money(payment.get('amount') or 0, 2)} руб."
                )
            details = "; ".join(samples)
            await update.message.reply_text(
                "⚠️ В акте сверки есть оплаты, которые не удалось строго "
                f"сопоставить с заявками. Примеры: {details}"
            )

    if shipments:
        for shipment in shipments:
            shipment["source"] = normalize_shipping_source(shipment.get("source"))
        claim_data["shipments"] = shipments
        if any(item.get("source") == "post" for item in shipments):
            claim_data["shipping_method"] = "почта"
        elif any(item.get("source") == "cdek" for item in shipments):
            claim_data["shipping_method"] = "сдэк"
        shipments_for_numbers = shipments
        if any(item.get("api_records", 0) > 0 for item in shipments):
            shipments_for_numbers = [
                item for item in shipments if item.get("api_records", 0) > 0
            ]
        claim_data["postal_numbers"] = [
            item["track_number"] for item in shipments_for_numbers if item.get("track_number")
        ]
        postal_dates = []
        for item in shipments:
            received = item.get("received_date")
            if isinstance(received, datetime):
                postal_dates.append(received.strftime("%d.%m.%Y"))
            elif isinstance(received, str):
                parsed = parse_date_str(received)
                if parsed:
                    postal_dates.append(parsed.strftime("%d.%m.%Y"))
        claim_data["postal_dates"] = postal_dates
        claim_data["docs_track_number"] = get_first_list_value(
            claim_data.get("postal_numbers", [])
        )
        claim_data["docs_received_date"] = get_first_list_value(
            claim_data.get("postal_dates", [])
        )

    claim_data["pretension_groups"] = groups
    context.user_data["pretension_data"] = claim_data

    missing = get_pretension_missing_fields(claim_data)
    context.user_data["pretension_missing_fields"] = missing
    if missing:
        return await ask_next_pretension_field(update, context)
    return await finish_pretension(update, context)


async def handle_pretension_field(update, context):
    data = context.user_data.get("pretension_data", {})
    missing = context.user_data.get("pretension_missing_fields", [])
    if not missing:
        return await finish_pretension(update, context)

    key = missing[0]
    raw = update.message.text.strip() if update.message else ""
    field_def = PRETENSION_FIELD_DEFS.get(key, {})
    required = field_def.get("required", True)

    if raw.lower() == "пропустить" and not required:
        if key == "docs_track_number":
            data["skip_postal_lookup"] = True
        data[key] = ""
        context.user_data["pretension_data"] = data
        context.user_data["pretension_missing_fields"] = get_pretension_missing_fields(data)
        return await ask_next_pretension_field(update, context)

    if key == "debt":
        amount = parse_amount(raw, 0)
        if amount <= 0:
            await update.message.reply_text(
                "Введите сумму задолженности в рублях (например: 210000)."
            )
            return ASK_PRETENSION_FIELD
        data[key] = format_money(amount, 2)
    elif key == "payment_days":
        digits = re.sub(r"[^\d]", "", raw)
        if not digits:
            await update.message.reply_text(
                "Введите срок оплаты числом (например: 15)."
            )
            return ASK_PRETENSION_FIELD
        data[key] = digits
        groups = data.get("pretension_groups") or []
        if groups:
            for group in groups:
                if not group.get("payment_days"):
                    group["payment_days"] = int(digits)
    elif key == "shipping_method":
        method = normalize_shipping_method(raw)
        if method not in ("почта", "сдэк"):
            await update.message.reply_text(
                "Укажите «почта» или «сдэк»."
            )
            return ASK_PRETENSION_FIELD
        data["shipping_method"] = method
    elif key == "docs_received_date":
        dates = re.findall(r"\d{2}\.\d{2}\.\d{4}", raw)
        if not dates:
            await update.message.reply_text(
                "Введите дату в формате ДД.ММ.ГГГГ."
            )
            return ASK_PRETENSION_FIELD
        data["postal_dates"] = dates
        data[key] = dates[0]
        if dates and not data.get("shipments"):
            source = normalize_shipping_source(data.get("shipping_method"))
            data["shipments"] = [
                {
                    "track_number": None,
                    "received_date": parse_date_str(date_str),
                    "source": source,
                }
                for date_str in dates
            ]
    elif key == "docs_track_number":
        parts = re.split(r"[;\n,]+", raw)
        track_numbers = []
        pairs = []
        for part in parts:
            track_match = re.search(r"[A-Za-z0-9]{8,}", part)
            date_match = re.search(r"\d{2}\.\d{2}\.\d{4}", part)
            if track_match and date_match:
                pairs.append((track_match.group(0), date_match.group(0)))
            elif track_match:
                track_numbers.append(track_match.group(0))
        if pairs:
            data["postal_numbers"] = [normalize_tracking_number(t) for t, _ in pairs]
            data["postal_dates"] = [d for _, d in pairs]
            data[key] = data["postal_numbers"][0] if data["postal_numbers"] else ""
            data["docs_received_date"] = data["postal_dates"][0] if data["postal_dates"] else ""
            data.pop("skip_postal_lookup", None)
            data["shipments"] = [
                {
                    "track_number": normalize_tracking_number(t),
                    "received_date": parse_date_str(d),
                    "source": normalize_shipping_source(data.get("shipping_method")),
                }
                for t, d in pairs
            ]
        else:
            track_numbers = [normalize_tracking_number(t) for t in track_numbers if t]
            if not track_numbers:
                await update.message.reply_text(
                    "Укажите трек-номер (например: 10152795200)."
                )
                return ASK_PRETENSION_FIELD
            data["postal_numbers"] = track_numbers
            data[key] = track_numbers[0]
            data.pop("skip_postal_lookup", None)
            if not data.get("shipping_method"):
                if all(t.isdigit() and len(t) in (13, 14) for t in track_numbers):
                    data["shipping_method"] = "почта"
            if normalize_shipping_method(data.get("shipping_method")) == "почта":
                config = get_russian_post_config()
                if config.get("enabled"):
                    postal_dates = []
                    for track in track_numbers:
                        try:
                            records = fetch_russian_post_operations(track)
                            _, receive_date = extract_tracking_dates(records)
                        except Exception:
                            receive_date = None
                        if receive_date:
                            postal_dates.append(receive_date)
                    if postal_dates:
                        data["postal_dates"] = postal_dates
                        data["docs_received_date"] = postal_dates[0]
                        data["shipments"] = [
                            {
                                "track_number": track,
                                "received_date": parse_date_str(date_str),
                                "source": "post",
                            }
                            for track, date_str in zip(track_numbers, postal_dates)
                        ]
                    else:
                        data.pop("skip_postal_lookup", None)
                        await update.message.reply_text(
                            "Не удалось получить даты по трек-номерам через API "
                            "Почты России. Продолжаю без дат; при необходимости "
                            "можете указать их вручную позже."
                        )
    else:
        data[key] = raw

    context.user_data["pretension_data"] = data
    missing = get_pretension_missing_fields(data)
    context.user_data["pretension_missing_fields"] = missing
    if missing:
        return await ask_next_pretension_field(update, context)
    return await finish_pretension(update, context)


async def finish_pretension(update, context):
    files = context.user_data.get("pretension_files", [])
    file_paths = [entry["path"] for entry in files if entry.get("path")]
    if not file_paths:
        file_path = context.user_data.get("file_path")
        if file_path:
            file_paths = [file_path]
    if not file_paths or not all(os.path.exists(path) for path in file_paths):
        await update.message.reply_text(
            "Ошибка: файл не найден на диске."
        )
        return ConversationHandler.END

    claim_data = context.user_data.get("pretension_data", {})
    groups = claim_data.get("pretension_groups", []) or []
    document_groups = claim_data.get("document_groups", []) or []
    shipments = claim_data.get("shipments") or []

    if not shipments and claim_data.get("postal_numbers"):
        numbers = claim_data.get("postal_numbers", []) or []
        dates = claim_data.get("postal_dates", []) or []
        shipping_method = normalize_shipping_method(
            claim_data.get("shipping_method")
        )
        for number, date_str in zip(numbers, dates):
            shipments.append({
                "track_number": number,
                "received_date": parse_date_str(date_str),
                "source": "cdek" if shipping_method == "сдэк" else "post",
            })

    if claim_data.get("docs_track_number") and not claim_data.get("postal_numbers"):
        claim_data["postal_numbers"] = [claim_data.get("docs_track_number")]
    if claim_data.get("docs_received_date") and not claim_data.get("postal_dates"):
        claim_data["postal_dates"] = [claim_data.get("docs_received_date")]

    if shipments:
        for shipment in shipments:
            shipment["source"] = normalize_shipping_source(shipment.get("source"))
        if groups and any(
            not group.get("docs_received_date") or not group.get("docs_track_number")
            for group in groups
        ):
            assign_shipments_to_groups(groups, shipments)

    raw_plaintiff_name = normalize_str(claim_data.get("plaintiff_name"))
    raw_defendant_name = normalize_str(claim_data.get("defendant_name"))
    plaintiff_name_short = format_organization_name_short(raw_plaintiff_name)
    defendant_name_short = format_organization_name_short(raw_defendant_name)
    plaintiff_name = plaintiff_name_short
    defendant_name = defendant_name_short
    is_plaintiff_ip = (
        "ИП" in raw_plaintiff_name
        or "Индивидуальный предприниматель" in raw_plaintiff_name
    )
    is_defendant_ip = (
        "ИП" in raw_defendant_name
        or "Индивидуальный предприниматель" in raw_defendant_name
    )

    debt_amount = parse_amount(claim_data.get("debt", "0"))
    debt_decimal = parse_amount_decimal(claim_data.get("debt", "0"))
    debt_rubles, debt_kopeks = split_rubles_kopeks(debt_decimal)
    payment_days_raw = claim_data.get("payment_days", "0")
    try:
        payment_days = int(re.sub(r"[^\d]", "", str(payment_days_raw)))
    except ValueError:
        payment_days = 0

    interest_data = {"total_interest": 0.0, "detailed_calc": []}
    partial_payments = claim_data.get("partial_payments_info") or []
    has_group_payment_days = False
    if groups:
        for group in groups:
            try:
                if int(group.get("payment_days") or 0) > 0:
                    has_group_payment_days = True
                    break
            except (TypeError, ValueError):
                continue

    has_group_terms = any(
        normalize_payment_terms(group.get("payment_terms") or "")
        for group in (groups or [])
    )

    if groups and (payment_days > 0 or has_group_payment_days or has_group_terms):
        interest_data = calculate_pretension_interest_schedule(
            groups,
            payment_days,
            payments=partial_payments
        )
    else:
        docs_received_date = parse_date_str(claim_data.get("docs_received_date", ""))
        if docs_received_date and payment_days > 0 and debt_amount > 0:
            calendar = load_work_calendar(docs_received_date.year)
            due_date = add_working_days(
                docs_received_date,
                payment_days,
                calendar
            )
            interest_start = due_date + timedelta(days=1)
            interest_data = calculate_pretension_interest(
                debt_amount,
                interest_start,
                payments=partial_payments
            )
    total_interest = parse_amount(interest_data.get("total_interest", 0))
    legal_fees_value = parse_amount(claim_data.get("legal_fees", "0"))

    payment_terms_text = normalize_payment_terms(
        claim_data.get("payment_terms", "")
    )
    if not payment_terms_text or payment_terms_text == "Не указано":
        if payment_days > 0:
            payment_terms_text = (
                "Оплата не позднее "
                f"{payment_days} рабочих дней с даты получения документов, "
                "подтверждающих перевозку"
            )
        else:
            payment_terms_text = "Не указано"
    payment_terms_text = build_payment_terms_summary(
        groups,
        payment_terms_text,
        payment_days
    )

    applications = [
        group.get("application")
        for group in (groups or document_groups)
        if group.get("application")
    ]
    cargo_docs = []
    for group in groups:
        cargo_docs.extend(group.get("cargo_docs", []))
    if not cargo_docs:
        cargo_docs = split_document_items(claim_data.get("cargo_docs"))
    intro_paragraph = build_intro_paragraph(
        plaintiff_name_short,
        applications,
        cargo_docs,
        include_docs=False
    )

    # LLM-коррекция склонений/регистра/опечаток (без изменения реквизитов)
    plaintiff_address_value = normalize_str(
        claim_data.get("plaintiff_address")
    ).replace("\n", " ").strip()
    defendant_address_value = normalize_str(
        claim_data.get("defendant_address")
    ).replace("\n", " ").strip()
    plaintiff_address_value = maybe_proofread_text(
        plaintiff_address_value,
        protected_values=[plaintiff_name, plaintiff_name_short]
    )
    defendant_address_value = maybe_proofread_text(
        defendant_address_value,
        protected_values=[defendant_name, defendant_name_short]
    )
    intro_paragraph = maybe_proofread_text(
        intro_paragraph,
        protected_values=[plaintiff_name_short]
    )
    payment_terms_text = maybe_proofread_text(payment_terms_text)

    plaintiff_ogrn_type = get_ogrn_label(
        plaintiff_name,
        claim_data.get("plaintiff_inn", "")
    )
    defendant_ogrn_type = get_ogrn_label(
        defendant_name,
        claim_data.get("defendant_inn", "")
    )

    defendant_block = build_party_block(
        "Кому",
        defendant_name_short,
        normalize_str(claim_data.get("defendant_inn")),
        normalize_str(claim_data.get("defendant_kpp")),
        normalize_str(claim_data.get("defendant_ogrn")),
        defendant_ogrn_type,
        defendant_address_value,
        defendant_address_value,
        is_defendant_ip
    )
    plaintiff_block = build_party_block(
        "От кого",
        plaintiff_name_short,
        normalize_str(claim_data.get("plaintiff_inn")),
        normalize_str(claim_data.get("plaintiff_kpp")),
        normalize_str(claim_data.get("plaintiff_ogrn")),
        plaintiff_ogrn_type,
        plaintiff_address_value,
        plaintiff_address_value,
        is_plaintiff_ip
    )

    if groups:
        documents_list_structured = build_documents_list_structured_for_groups(groups)
    else:
        documents_list_structured = build_documents_list_structured(document_groups)
    attachments = build_pretension_attachments(groups or document_groups, claim_data)
    shipping_summary = build_shipping_summary(
        shipments,
        documents_count=len(groups) if groups else None
    )

    replacements = {
        "{defendant_block}": defendant_block,
        "{plaintiff_block}": plaintiff_block,
        "{intro_paragraph}": intro_paragraph,
        "{documents_list}": build_documents_list(claim_data),
        "{debt_amount}": debt_rubles,
        "{debt_kopeks}": debt_kopeks,
        "{payment_terms}": payment_terms_text,
        "{legal_fees_block}": build_legal_fees_block(claim_data),
        "{requirements_summary}": build_requirements_summary(
            debt_amount,
            total_interest,
            legal_fees_value
        ),
        "{pretension_date}": format_russian_date(),
        "{shipping_info}": shipping_summary,
        "{docs_track_number}": normalize_str(
            claim_data.get("docs_track_number", ""),
            default=""
        ),
        "{docs_received_date}": normalize_str(
            claim_data.get("docs_received_date", ""),
            default=""
        ),
        "{plaintiff_name}": plaintiff_name,
        "{defendant_name}": defendant_name,
    }

    # Добавляем текст "осознанности" (частичные оплаты, гарантийные письма и т.д.)
    awareness_text = claim_data.get("awareness_text", "")
    if awareness_text:
        # Если есть скорректированная сумма - подставляем её в текст
        adjusted_debt_str = claim_data.get("debt", format_money(debt_amount, 2))
        awareness_text = awareness_text.replace("{adjusted_debt}", adjusted_debt_str)
        replacements["{awareness_block}"] = awareness_text
    else:
        replacements["{awareness_block}"] = ""

    result_docx = create_pretension_document(
        claim_data,
        interest_data,
        replacements,
        documents_list_structured=documents_list_structured,
        attachments=attachments,
        proofread_protected_values=[
            plaintiff_name,
            defendant_name,
            plaintiff_name_short,
            defendant_name_short,
        ],
    )

    with open(result_docx, "rb") as f:
        await update.message.reply_document(
            InputFile(f, filename="Претензия.docx"),
            caption="Претензия по документам перевозки"
        )

    try:
        for path in file_paths:
            if os.path.exists(path):
                os.remove(path)
        if os.path.exists(result_docx):
            os.remove(result_docx)
    except Exception as exc:
        logging.warning("Не удалось удалить временные файлы: %s", exc)

    return ConversationHandler.END


# ============ ОБРАБОТЧИКИ ВНЕШНИХ ПРЕТЕНЗИЙ ============

async def handle_external_claim_document(update, context):
    """Обработчик загрузки документов для внешних претензий."""
    stage = context.user_data.get("external_claim_stage", "claim")
    files = context.user_data.get("external_claim_files", [])

    # Обработка загрузки документа
    if update.message and update.message.document:
        doc = update.message.document
        if not doc.file_name or not doc.file_name.lower().endswith('.pdf'):
            await update.message.reply_text('Пожалуйста, отправьте PDF-файл.')
            return ASK_EXTERNAL_CLAIM_DOCUMENT

        os.makedirs('uploads', exist_ok=True)
        unique_name = f"{uuid.uuid4()}_{doc.file_name}"
        file_path = os.path.join('uploads', unique_name)
        telegram_file = await doc.get_file()
        await telegram_file.download_to_drive(file_path)

        files.append({
            "path": file_path,
            "name": doc.file_name,
            "type": stage  # claim, docs, legal
        })
        context.user_data["external_claim_files"] = files

        if stage == "claim":
            await update.message.reply_text(
                f"✅ Претензия загружена: {doc.file_name}\n\n"
                "Теперь отправьте PDF-файлы с пакетами документов (СП) по каждой перевозке.\n"
                "Когда все файлы загружены, напишите «готово»."
            )
            context.user_data["external_claim_stage"] = "docs"
            return ASK_EXTERNAL_CLAIM_DOCUMENT

        elif stage == "docs":
            await update.message.reply_text(
                f"✅ Документ загружен: {doc.file_name}\n"
                "Отправьте ещё документы или напишите «готово»."
            )
            return ASK_EXTERNAL_CLAIM_DOCUMENT

        elif stage == "legal":
            await update.message.reply_text(
                f"✅ Договор юр. услуг загружен: {doc.file_name}"
            )
            # Переходим к обработке
            return await process_external_claim(update, context)

    # Обработка текстовых команд
    if update.message and update.message.text:
        text = update.message.text.strip().lower()

        if text in {"готово", "done", "finish", "все"}:
            if stage == "claim":
                await update.message.reply_text(
                    "Сначала загрузите PDF-файл с претензией."
                )
                return ASK_EXTERNAL_CLAIM_DOCUMENT

            elif stage == "docs":
                # Проверяем, есть ли документы помимо претензии
                doc_files = [f for f in files if f.get("type") == "docs"]
                if not doc_files:
                    await update.message.reply_text(
                        "Загрузите хотя бы один пакет документов (СП)."
                    )
                    return ASK_EXTERNAL_CLAIM_DOCUMENT

                # Спрашиваем про юр. услуги
                keyboard = [
                    [
                        InlineKeyboardButton("✅ Да", callback_data='ext_legal_yes'),
                        InlineKeyboardButton("❌ Нет", callback_data='ext_legal_no'),
                    ]
                ]
                reply_markup = InlineKeyboardMarkup(keyboard)
                await update.message.reply_text(
                    "Есть договор на юридические услуги?",
                    reply_markup=reply_markup
                )
                return ASK_EXTERNAL_CLAIM_DOCUMENT

        elif text in {"пропустить", "skip"}:
            if stage == "legal":
                return await process_external_claim(update, context)

    await update.message.reply_text(
        "Отправьте PDF-файл или напишите «готово»."
    )
    return ASK_EXTERNAL_CLAIM_DOCUMENT


async def handle_external_legal_choice(update, context):
    """Обработчик выбора: есть ли договор юр. услуг."""
    query = update.callback_query
    await query.answer()
    choice = query.data

    if choice == "ext_legal_yes":
        context.user_data["external_claim_stage"] = "legal"
        await query.edit_message_text(
            "Отправьте PDF-файл с договором на юридические услуги."
        )
        return ASK_EXTERNAL_CLAIM_DOCUMENT

    elif choice == "ext_legal_no":
        await query.edit_message_text("Обрабатываю документы...")
        return await process_external_claim(update, context)

    return ASK_EXTERNAL_CLAIM_DOCUMENT


async def process_external_claim(update, context):
    """Обработка загруженных документов внешней претензии."""
    files = context.user_data.get("external_claim_files", [])

    # Разделяем файлы по типам
    claim_file = None
    doc_files = []
    legal_file = None

    for f in files:
        if f.get("type") == "claim":
            claim_file = f["path"]
        elif f.get("type") == "docs":
            doc_files.append(f["path"])
        elif f.get("type") == "legal":
            legal_file = f["path"]

    if not claim_file:
        message = update.message or update.callback_query.message
        await message.reply_text("Ошибка: файл претензии не найден.")
        return ConversationHandler.END

    message = update.message or update.callback_query.message
    await message.reply_text("⏳ Анализирую документы с помощью LLM...")

    try:
        # Парсим претензию
        claim_data = parse_external_claim(claim_file)

        # Парсим пакеты документов
        if doc_files:
            doc_packages = parse_document_packages(doc_files)

            # Связываем документы
            claim_data = link_documents_full(claim_data, doc_packages)

        # Парсим договор юр. услуг (если есть)
        # TODO: добавить парсер договора юр. услуг

        # Формируем данные для искового
        lawsuit_data = convert_external_to_lawsuit(claim_data)

        # Выводим результат
        apps = claim_data.applications or []
        await message.reply_text(
            f"✅ Документы обработаны!\n\n"
            f"📝 Истец: {claim_data.plaintiff.name if claim_data.plaintiff else 'Не определён'}\n"
            f"📝 Ответчик: {claim_data.defendant.name if claim_data.defendant else 'Не определён'}\n"
            f"💰 Общий долг: {claim_data.total_debt:,.2f} руб.\n"
            f"📋 Заявок: {len(apps)}\n"
        )

        # Показываем предупреждения
        if claim_data.warnings:
            warnings_text = "\n".join(f"⚠️ {w}" for w in claim_data.warnings[:5])
            await message.reply_text(f"Предупреждения:\n{warnings_text}")

        # Сохраняем данные
        context.user_data["external_claim_data"] = lawsuit_data
        context.user_data["external_claim_raw"] = claim_data

        # Проверяем недостающие поля
        missing = get_external_claim_missing_fields(lawsuit_data)
        context.user_data["external_claim_missing"] = missing

        if missing:
            return await ask_next_external_claim_field(update, context)

        return await finish_external_claim(update, context)

    except Exception as exc:
        logging.exception("Error processing external claim")
        await message.reply_text(
            f"❌ Ошибка при обработке документов: {exc}"
        )
        return ConversationHandler.END


def convert_external_to_lawsuit(claim_data: ExternalClaimData) -> dict:
    """Конвертирует ExternalClaimData в формат для искового заявления."""
    from decimal import Decimal
    result = {
        "plaintiff_name": claim_data.plaintiff.name if claim_data.plaintiff else "",
        "plaintiff_inn": claim_data.plaintiff.inn if claim_data.plaintiff else "",
        "plaintiff_kpp": claim_data.plaintiff.kpp if claim_data.plaintiff else "",
        "plaintiff_ogrn": claim_data.plaintiff.ogrn if claim_data.plaintiff else "",
        "plaintiff_address": claim_data.plaintiff.address if claim_data.plaintiff else "",

        "defendant_name": claim_data.defendant.name if claim_data.defendant else "",
        "defendant_inn": claim_data.defendant.inn if claim_data.defendant else "",
        "defendant_kpp": claim_data.defendant.kpp if claim_data.defendant else "",
        "defendant_ogrn": claim_data.defendant.ogrn if claim_data.defendant else "",
        "defendant_address": claim_data.defendant.address if claim_data.defendant else "",

        "base_contract_number": claim_data.base_contract.number if claim_data.base_contract else "",
        "base_contract_date": claim_data.base_contract.date if claim_data.base_contract else "",

        "debt": str(claim_data.total_debt) if claim_data.total_debt else "0",
        "claim_date": claim_data.claim_date or "",

        "pretension_groups": [],
        "source_files": claim_data.source_files or [],
        "warnings": claim_data.warnings or [],
    }

    # Конвертируем заявки в группы
    for app in (claim_data.applications or []):
        group = {
            "application": f"Заявка № {app.number} от {app.date}",
            "application_number": app.number,
            "application_date": app.date,
            "amount": float(app.amount) if app.amount else 0.0,
            "amount_without_vat": float(app.amount / Decimal("1.2")) if app.amount else 0.0,
            "route": app.route or "",
            "vehicle_plate": app.vehicle_plate or "",
            "trailer_plate": app.trailer_plate or "",
            "driver_name": app.driver_name or "",
            "driver_inn": "",
            "load_date": app.load_date or "",
            "payment_days": app.payment_days or 0,
            "payment_terms": app.payment_terms or "",
            "waybill_number": app.linked_waybill.number if app.linked_waybill else "",
            "waybill_date": app.linked_waybill.date if app.linked_waybill else "",
            "docs_track_number": app.linked_shipment.track_number if app.linked_shipment else "",
            "docs_received_date": app.linked_shipment.received_date if app.linked_shipment else "",
            "cargo_docs": [],
        }
        if app.linked_waybill:
            group["cargo_docs"].append(
                f"ТН № {app.linked_waybill.number} от {app.linked_waybill.date or ''}"
            )
        result["pretension_groups"].append(group)

    return result


# Поля для внешних претензий
EXTERNAL_CLAIM_FIELD_ORDER = [
    "plaintiff_name",
    "plaintiff_inn",
    "defendant_name",
    "defendant_inn",
]

EXTERNAL_CLAIM_FIELD_DEFS = {
    "plaintiff_name": {
        "prompt": "Укажите полное наименование истца:",
        "required": True,
    },
    "plaintiff_inn": {
        "prompt": "Укажите ИНН истца:",
        "required": True,
    },
    "defendant_name": {
        "prompt": "Укажите полное наименование ответчика:",
        "required": True,
    },
    "defendant_inn": {
        "prompt": "Укажите ИНН ответчика:",
        "required": True,
    },
}


def get_external_claim_missing_fields(data: dict) -> list:
    """Возвращает список отсутствующих обязательных полей."""
    missing = []
    for key in EXTERNAL_CLAIM_FIELD_ORDER:
        field_def = EXTERNAL_CLAIM_FIELD_DEFS.get(key, {})
        if not field_def.get("required"):
            continue
        value = data.get(key)
        if not value or value == "Не указано":
            missing.append(key)
    return missing


async def ask_next_external_claim_field(update, context):
    """Запрашивает следующее отсутствующее поле."""
    missing = context.user_data.get("external_claim_missing", [])
    if not missing:
        return await finish_external_claim(update, context)

    key = missing[0]
    field_def = EXTERNAL_CLAIM_FIELD_DEFS.get(key, {})
    prompt = field_def.get("prompt", f"Введите {key}:")

    message = update.message or update.callback_query.message
    await message.reply_text(prompt)
    return ASK_EXTERNAL_CLAIM_FIELD


async def handle_external_claim_field(update, context):
    """Обработчик ввода поля для внешней претензии."""
    data = context.user_data.get("external_claim_data", {})
    missing = context.user_data.get("external_claim_missing", [])

    if not missing:
        return await finish_external_claim(update, context)

    key = missing[0]
    raw = update.message.text.strip() if update.message else ""

    # Валидация ИНН
    if key in ("plaintiff_inn", "defendant_inn"):
        digits = re.sub(r"[^\d]", "", raw)
        if len(digits) not in (10, 12):
            await update.message.reply_text(
                "ИНН должен содержать 10 (для юр. лиц) или 12 (для ИП) цифр."
            )
            return ASK_EXTERNAL_CLAIM_FIELD
        data[key] = digits
    else:
        data[key] = raw

    context.user_data["external_claim_data"] = data
    missing = get_external_claim_missing_fields(data)
    context.user_data["external_claim_missing"] = missing

    if missing:
        return await ask_next_external_claim_field(update, context)
    return await finish_external_claim(update, context)


async def finish_external_claim(update, context):
    """Генерирует исковое заявление на основе внешней претензии."""
    from decimal import Decimal

    data = context.user_data.get("external_claim_data", {})
    raw_data = context.user_data.get("external_claim_raw")

    message = update.message or update.callback_query.message
    await message.reply_text("⏳ Рассчитываю проценты и генерирую исковое заявление...")

    try:
        groups = data.get("pretension_groups", [])

        # Расчёт процентов для каждой перевозки
        total_interest = Decimal("0")
        for group in groups:
            received_date_str = group.get("docs_received_date")
            payment_days = group.get("payment_days") or 20
            amount = Decimal(str(group.get("amount", 0)))

            if received_date_str and amount > 0:
                received_date = parse_date_str(received_date_str)
                if received_date:
                    # Рассчитываем дату начала просрочки
                    calendar = load_work_calendar(received_date.year)
                    due_date = add_working_days(received_date, payment_days, calendar)
                    interest_start = due_date + timedelta(days=1)

                    # Рассчитываем проценты
                    interest_result = calculate_pretension_interest(
                        float(amount),
                        interest_start
                    )
                    group["interest"] = interest_result.get("total_interest", 0)
                    group["interest_details"] = interest_result.get("detailed_calc", [])
                    group["due_date"] = due_date.strftime("%d.%m.%Y")
                    group["interest_start"] = interest_start.strftime("%d.%m.%Y")
                    total_interest += Decimal(str(group["interest"]))

        data["total_interest"] = float(total_interest)
        data["pretension_groups"] = groups

        # Расчёт общей суммы иска
        debt = Decimal(str(data.get("debt", 0)))
        claim_total = debt + total_interest
        data["claim_total"] = float(claim_total)

        # Расчёт госпошлины
        duty = calculate_duty(float(claim_total))
        data["duty"] = duty

        # Выводим итоги
        await message.reply_text(
            f"📊 Расчёт выполнен:\n\n"
            f"💰 Основной долг: {format_money(float(debt), 0)} руб.\n"
            f"📈 Проценты (ст. 395 ГК РФ): {format_money(float(total_interest), 2)} руб.\n"
            f"📋 Итого сумма иска: {format_money(float(claim_total), 2)} руб.\n"
            f"⚖️ Госпошлина: {format_money(duty, 0)} руб."
        )

        # Показываем детали по перевозкам
        details = []
        for i, group in enumerate(groups, 1):
            app_num = group.get("application_number", f"#{i}")
            amount = group.get("amount", 0)
            interest = group.get("interest", 0)
            received = group.get("docs_received_date", "?")
            details.append(
                f"{i}. {app_num}: {format_money(amount, 0)} руб. + {format_money(interest, 2)} руб. (получено {received})"
            )

        if details:
            await message.reply_text(
                "📋 Детали по перевозкам:\n" + "\n".join(details)
            )

        # TODO: Генерация DOCX искового заявления
        await message.reply_text(
            "✅ Данные подготовлены.\n"
            "Генерация DOCX искового заявления будет добавлена в следующей версии."
        )

        # Очистка временных файлов
        files = context.user_data.get("external_claim_files", [])
        for f in files:
            try:
                if os.path.exists(f.get("path", "")):
                    os.remove(f["path"])
            except Exception:
                pass

        return ConversationHandler.END

    except Exception as exc:
        logging.exception("Error finishing external claim")
        await message.reply_text(f"❌ Ошибка: {exc}")
        return ConversationHandler.END


async def handle_document(update, context):
    flow = context.user_data.get("flow")
    if flow == "claim":
        return await handle_docx_entry(update, context)
    if flow == "pretension":
        return await handle_pretension_document(update, context)
    if flow == "external_claim":
        return await handle_external_claim_document(update, context)
    if update.message:
        await update.message.reply_text(
            "Сначала выбери тип документа через /start."
        )
    return ConversationHandler.END


conv_handler = ConversationHandler(
    entry_points=[CommandHandler("start", start)],
    states={
        ASK_FLOW: [CallbackQueryHandler(flow_chosen)],
        ASK_DOCUMENT: [
            MessageHandler(filters.Document.ALL, handle_document),
            MessageHandler(filters.TEXT & ~filters.COMMAND, handle_document),
        ],
        ASK_JURISDICTION: [CallbackQueryHandler(jurisdiction_chosen)],
        ASK_CUSTOM_COURT: [
            MessageHandler(filters.TEXT & ~filters.COMMAND, handle_custom_court)
        ],
        ASK_CLAIM_STATUS: [CallbackQueryHandler(claim_status_chosen)],
        ASK_TRACK: [
            MessageHandler(filters.TEXT & ~filters.COMMAND, ask_track)
        ],
        ASK_RECEIVE_DATE: [
            MessageHandler(filters.TEXT & ~filters.COMMAND, ask_receive_date)
        ],
        ASK_SEND_DATE: [
            MessageHandler(filters.TEXT & ~filters.COMMAND, ask_send_date)
        ],
        ASK_BIRTH_DATE: [
            MessageHandler(filters.TEXT & ~filters.COMMAND, ask_birth_date)
        ],
        ASK_BIRTH_PLACE: [
            MessageHandler(filters.TEXT & ~filters.COMMAND, ask_birth_place)
        ],
        ASK_PRETENSION_FIELD: [
            MessageHandler(
                filters.TEXT & ~filters.COMMAND, handle_pretension_field
            )
        ],
        ASK_EXTERNAL_CLAIM_DOCUMENT: [
            MessageHandler(
                filters.Document.ALL, handle_external_claim_document
            ),
            MessageHandler(
                filters.TEXT & ~filters.COMMAND, handle_external_claim_document
            ),
            CallbackQueryHandler(handle_external_legal_choice),
        ],
        ASK_EXTERNAL_CLAIM_FIELD: [
            MessageHandler(
                filters.TEXT & ~filters.COMMAND, handle_external_claim_field
            )
        ],
    },
    fallbacks=[],
    per_message=False,
)


def generate_postal_block(postal_numbers, postal_dates):
    if not postal_numbers or not postal_dates or len(postal_numbers) != len(postal_dates):
        return "Не указано"
    if len(postal_numbers) == 1:
        return (
            f"Почтовым уведомлением № {postal_numbers[0]} об отправке и получении "
            f"{postal_dates[0]} оригиналов документов Заказчиком."
        )
    else:
        pairs = [
            f"№ {num} от {date}" for num, date in zip(postal_numbers, postal_dates)
        ]
        return (
            f"Почтовыми уведомлениями {', '.join(pairs)} "
            f"об отправке и получении оригиналов документов Заказчиком."
        )


def clean_uploads_folder():
    """Удаляет все файлы из папки uploads при запуске бота."""
    uploads_dir = os.path.join(os.path.dirname(__file__), 'uploads')
    if not os.path.exists(uploads_dir):
        return
    for filename in os.listdir(uploads_dir):
        file_path = os.path.join(uploads_dir, filename)
        try:
            if os.path.isfile(file_path):
                os.remove(file_path)
                logging.info(f"Удален файл из uploads: {file_path}")
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
                logging.info(f"Удалена папка из uploads: {file_path}")
        except Exception as e:
            logging.warning(f"Не удалось удалить {file_path}: {e}")


def main() -> None:
    """Запускает Telegram бота."""
    logging.info("Starting bot...")
    clean_uploads_folder()  # Очищаем uploads при запуске
    if not TOKEN:
        logging.error(
            "TOKEN is not set. Please provide a valid Telegram bot token."
        )
        raise ValueError(
            "TOKEN is not set. Please provide a valid Telegram bot token.")
    app = Application.builder().token(TOKEN).build()
    logging.info("Bot initialized")
    app.add_handler(conv_handler)
    logging.info("Handlers added")
    app.run_polling()
    logging.info("Bot is polling")


if __name__ == '__main__':
    main()
