# -*- coding: utf-8 -*-
"""
Модуль для парсинга данных из .docx файлов (исковые требования, периоды).
"""

import logging
import re
from datetime import datetime
from typing import Any, Dict, List, Tuple

from docx import Document


def parse_periods_from_docx(
    docx_path: str,
) -> Tuple[List[Dict[str, Any]], float]:
    """
    Парсит таблицу из .docx и возвращает периоды для расчета процентов.
    """
    doc = Document(docx_path)
    if not doc.tables:
        raise ValueError("В файле .docx отсутствует таблица")
    table = doc.tables[0]

    if len(table.rows) == 0:
        raise ValueError("Таблица пуста")

    header_cells = table.rows[0].cells
    if len(header_cells) < 7:
        raise ValueError(
            f"Таблица должна содержать минимум 7 столбцов, найдено: "
            f"{len(header_cells)}"
        )

    logging.info(
        f"Найдена таблица с {len(header_cells)} столбцами и "
        f"{len(table.rows)} строками"
    )

    periods = []
    current_sum = 0.0

    for row in table.rows[1:]:
        cells = row.cells
        if len(cells) < 7:
            logging.warning(
                f"Некорректная строка таблицы: {len(cells)} столбцов"
            )
            continue
        try:
            amount_text = cells[0].text.strip()
            logging.debug(f"Оригинальная сумма: '{amount_text}'")
            if not amount_text:
                logging.debug("Сумма пустая, пропускаем строку")
                continue
            cleaned_for_check = re.sub(r'[^\d.,+\-]', '', amount_text)
            if not cleaned_for_check:
                logging.debug(f"Строка не содержит чисел: '{amount_text}'")
                continue
            if re.search(r'[А-Яа-яA-Za-z]', amount_text):
                if any(word in amount_text.lower()
                       for word in ["сумма", "задолж", "процент"]):
                    logging.debug(
                        f"Строка пропущена по ключевому слову: "
                        f"'{amount_text}'"
                    )
                    continue
            amount_text = re.sub(r'[^\d.,+\-]', '', amount_text)
            logging.debug(f"После очистки: '{amount_text}'")
            if not amount_text:
                logging.debug("Сумма пустая после очистки, пропускаем строку")
                continue
            amount_text = amount_text.rstrip('.').rstrip(',').rstrip()
            logging.debug(f"После rstrip: '{amount_text}'")
            if not amount_text:
                logging.debug("Сумма пустая после rstrip, пропускаем строку")
                continue
            if not re.match(r'^[\+\-]?\d+([.,]\d+)?$', amount_text):
                logging.debug(f"Сумма не является числом: '{amount_text}'")
                continue
            try:
                if amount_text.startswith('+'):
                    amount = float(amount_text[1:].replace(',', '.'))
                    current_sum += amount
                    logging.debug(
                        f"Добавлена сумма: {amount}, текущая: {current_sum}"
                    )
                else:
                    current_sum = float(amount_text.replace(',', '.'))
                    logging.debug(f"Установлена сумма: {current_sum}")
            except ValueError as e:
                logging.debug(
                    f"Ошибка конвертации суммы '{amount_text}': {e}"
                )
                continue
            date_from_text = cells[1].text.strip()
            date_to_text = cells[2].text.strip()
            if not date_from_text or not date_to_text:
                continue
            if date_to_text.lower() == "новая задолженность":
                continue
            if (
                not re.match(r'\d{2}\.\d{2}\.\d{4}', date_from_text)
                or not re.match(r'\d{2}\.\d{2}\.\d{4}', date_to_text)
            ):
                continue
            date_from = datetime.strptime(date_from_text, '%d.%m.%Y')
            date_to = datetime.strptime(date_to_text, '%d.%m.%Y')
            days_text = cells[3].text.strip()
            if not days_text or not days_text.isdigit():
                continue
            days = int(days_text)
            rate_text = cells[4].text.strip()
            if not rate_text:
                continue
            rate = float(rate_text.replace(',', '.'))
            interest_text = cells[6].text.strip()
            if not interest_text:
                continue
            interest_text = re.sub(r'[^\d.,]', '', interest_text)
            interest_text = interest_text.rstrip('.').rstrip(',').rstrip()
            if not interest_text:
                continue
            if not re.match(r'^\d+([.,]\d+)?$', interest_text):
                logging.debug(
                    f"Проценты не являются числом: '{interest_text}'"
                )
                continue
            try:
                interest = float(interest_text.replace(',', '.'))
            except ValueError as e:
                logging.debug(
                    f"Ошибка конвертации процентов '{interest_text}': {e}"
                )
                continue
            periods.append({
                'sum': current_sum,
                'date_from': date_from,
                'date_to': date_to,
                'days': days,
                'rate': rate,
                'interest': interest,
                'formula': cells[5].text.strip() if len(cells) > 5 else ""
            })
        except Exception as e:
            logging.warning(
                f"Ошибка парсинга строки: {e}"
            )
            logging.warning(
                f"Содержимое ячеек: {[cell.text.strip() for cell in cells]}"
            )
            continue
    if not periods:
        raise ValueError("Не удалось распарсить ни одной строки из таблицы")
    return periods, current_sum


def is_valid_date(date_str: str) -> bool:
    """
    Проверяет, является ли строка корректной датой.

    Args:
        date_str: Строка с датой

    Returns:
        True если дата корректна, False иначе
    """
    try:
        d = datetime.strptime(date_str.replace('/', '.'), '%d.%m.%Y')
        return 2000 <= d.year <= 2099
    except Exception:
        return False


def parse_contract_applications(text: str) -> str:
    """
    Парсит договоры-заявки на перевозку груза используя правильную логику группировки.
    """
    from simple_parser import find_documents_with_numbers_grouped

    # Используем правильный парсер из simple_parser.py
    documents = find_documents_with_numbers_grouped(text)

    # Возвращаем contract_applications из результатов
    return documents.get('contract_applications', 'Не указано')


def parse_cargo_documents(text: str) -> str:
    """
    Парсит грузосопроводительные документы.
    Возвращает только реальные документы и 'Комплектом сопроводительных документов на груз'.
    """
    found_documents = []

    # 1. Комплект сопроводительных документов - только эту фразу
    komplekt_match = re.search(
        r'Комплектом сопроводительных документов на груз', text, re.IGNORECASE)
    if komplekt_match:
        found_documents.append(
            "Комплектом сопроводительных документов на груз")

    # 2. Транспортная накладная - название + номер + дата
    transport_matches = re.finditer(
        r'([Тт]ранспортн[а-яё]* накладн[а-яё]*)\s*'
        r'(?:№\s*(\d+)\s*от\s*(\d{2}\.\d{2}\.\d{4})|'
        r'от\s*(\d{2}\.\d{2}\.\d{4}))',
        text, re.IGNORECASE
    )
    for match in transport_matches:
        name = match.group(1)
        if match.group(2) and match.group(3):  # есть номер и дата
            found_documents.append(
                f"{name} № {match.group(2)} от {match.group(3)} г.")
        elif match.group(4):  # есть только дата
            found_documents.append(
                f"{name} от {match.group(4)} г.")

    # 3. Товарная накладная - название + номер + дата
    tovar_matches = re.finditer(
        r'([Тт]оварн[а-яё]* накладн[а-яё]*)\s*'
        r'(?:№\s*(\d+)\s*от\s*(\d{2}\.\d{2}\.\d{4})|'
        r'от\s*(\d{2}\.\d{2}\.\d{4}))',
        text, re.IGNORECASE
    )
    for match in tovar_matches:
        name = match.group(1)
        if match.group(2) and match.group(3):  # есть номер и дата
            found_documents.append(
                f"{name} № {match.group(2)} от {match.group(3)} г.")
        elif match.group(4):  # есть только дата
            found_documents.append(
                f"{name} от {match.group(4)} г.")

    # 4. Товарно-транспортная накладная - название + номер + дата
    tt_matches = re.finditer(
        r'([Тт]оварно-транспортн[а-яё]* накладн[а-яё]*)\s*'
        r'(?:№\s*(\d+)\s*от\s*(\d{2}\.\d{2}\.\d{4})|'
        r'от\s*(\d{2}\.\d{2}\.\d{4}))',
        text, re.IGNORECASE
    )
    for match in tt_matches:
        name = match.group(1)
        if match.group(2) and match.group(3):  # есть номер и дата
            found_documents.append(
                f"{name} № {match.group(2)} от {match.group(3)} г.")
        elif match.group(4):  # есть только дата
            found_documents.append(
                f"{name} от {match.group(4)} г.")

    # 5. Счет-фактура - название + номер + дата
    sf_matches = re.finditer(
        r'([Сс]чет-фактур[а-яё]*)\s*'
        r'(?:№\s*(\d+)\s*от\s*(\d{2}\.\d{2}\.\d{4})|'
        r'от\s*(\d{2}\.\d{2}\.\d{4}))',
        text, re.IGNORECASE
    )
    for match in sf_matches:
        name = match.group(1)
        if match.group(2) and match.group(3):  # есть номер и дата
            found_documents.append(
                f"{name} № {match.group(2)} от {match.group(3)} г.")
        elif match.group(4):  # есть только дата
            found_documents.append(
                f"{name} от {match.group(4)} г.")

    # Исключаем любые фразы вида 'Грузосопроводительными документами', 'Грузосопроводительный документ' и т.п.
    found_documents = [d for d in found_documents if not re.match(
        r'Грузосопроводительн', d, re.IGNORECASE)]

    if found_documents:
        return '; '.join(found_documents)
    else:
        return "Не указано"


def parse_attachments(text: str) -> list:
    """
    Парсит раздел 'Приложения' и возвращает список приложений.
    Ищет все строки после заголовка до первой пустой строки или нового раздела.
    Разделяет по точке с запятой.
    """
    attachments = []
    match = re.search(
        r'Приложени[еяи]{1,2}\s*:?\s*([\s\S]+)', text, re.IGNORECASE)
    if not match:
        logging.info('DEBUG: Не найден раздел Приложения')
        return attachments
    after_header = match.group(1)
    logging.info('DEBUG: Текст после Приложения: %s', after_header[:300])
    # Берём только до первой пустой строки или до "Генеральный директор"
    stop_match = re.search(r'\n\s*\n|Генеральный директор', after_header)
    if stop_match:
        after_header = after_header[:stop_match.start()]
    # Разделяем по точке с запятой
    for item in after_header.split(';'):
        item = item.strip()
        if item:
            attachments.append(item)
    logging.info('DEBUG: Найденные приложения: %s', attachments)
    return attachments


def parse_claim_data(docx_path: str) -> Dict[str, Any]:
    """
    Парсит данные из .docx файла с досудебным требованием.

    Args:
        docx_path: Путь к .docx файлу

    Returns:
        Словарь с данными истца, ответчика и другими параметрами
    """
    doc = Document(docx_path)
    text = "\n".join(p.text for p in doc.paragraphs)

    # Извлекаем блоки истца и ответчика
    defendant_block = None
    plaintiff_block = None

    # Ищем блок ответчика (начинается с "Обществу с ограниченной ответственностью")
    defendant_pattern = (
        r'(Обществу с ограниченной ответственностью.*?)'
        r'(?=\n\n|\nот|\nТРЕБОВАНИЕ|\nПРЕТЕНЗИЯ)'
    )
    defendant_match = re.search(defendant_pattern, text, re.DOTALL)
    if defendant_match:
        defendant_block = defendant_match.group(1).strip()

    # Ищем блок истца (начинается с "от Индивидуального предпринимателя")
    plaintiff_pattern = (
        r'(от Индивидуального предпринимателя.*?)'
        r'(?=\n\n|\nТРЕБОВАНИЕ|\nПРЕТЕНЗИЯ)'
    )
    plaintiff_match = re.search(plaintiff_pattern, text, re.DOTALL)
    if plaintiff_match:
        plaintiff_block = plaintiff_match.group(1).strip()

    def extract_requisites(block: str) -> Dict[str, str]:
        """
        Извлекает реквизиты из блока текста.

        Args:
            block: Текст с реквизитами

        Returns:
            Словарь с реквизитами
        """
        # Простой парсер реквизитов - берем как есть
        lines = block.split('\n')
        name = ''
        inn = ''
        kpp = ''
        ogrn = ''
        ogrnip = ''
        address = ''

        for line in lines:
            line = line.strip()
            if not line:
                continue
            if 'ИНН' in line:
                inn_match = re.search(r'ИНН\s*(\d+)', line)
                inn = inn_match.group(1) if inn_match else ''
            elif 'КПП' in line:
                kpp_match = re.search(r'КПП\s*(\d+)', line)
                kpp = kpp_match.group(1) if kpp_match else ''
            elif 'ОГРНИП' in line:
                ogrnip_match = re.search(r'ОГРНИП\s*(\d+)', line)
                ogrnip = ogrnip_match.group(1) if ogrnip_match else ''
            elif 'ОГРН' in line:
                ogrn_match = re.search(r'ОГРН\s*(\d+)', line)
                ogrn = ogrn_match.group(1) if ogrn_match else ''
            elif inn and (kpp or ogrnip) and not address:
                # Адрес - первая строка после всех реквизитов
                address = line
            elif not name and ('ООО' in line or 'ИП' in line or 'Обществ' in line or 'Индивидуального предпринимателя' in line):
                # Имя - первая строка с названием организации
                name = line
                # Ограничиваем до кавычек, если есть
                if '«' in name and '»' in name:
                    start = name.find('«')
                    end = name.find('»')
                    if start != -1 and end != -1:
                        org_type = name[:start].strip()
                        org_name = name[start+1:end].strip()
                        name = f"{org_type} «{org_name}»"

        # Для ИП используем ОГРНИП вместо ОГРН, и не показываем КПП
        if 'ИП' in name or 'Индивидуальный предприниматель' in name:
            return {
                'name': name if name else 'Не указано',
                'inn': inn if inn else 'Не указано',
                'kpp': '',  # У ИП нет КПП
                'ogrn': ogrnip if ogrnip else 'Не указано',  # Используем ОГРНИП
                'address': address if address else 'Не указано',
            }
        else:
            return {
                'name': name if name else 'Не указано',
                'inn': inn if inn else 'Не указано',
                'kpp': kpp if kpp else 'Не указано',
                'ogrn': ogrn if ogrn else 'Не указано',
                'address': address if address else 'Не указано',
            }

    # Если нашли оба блока — используем их, иначе fallback на старый способ
    if defendant_block and plaintiff_block:
        defendant = extract_requisites(defendant_block)
        plaintiff = extract_requisites(plaintiff_block)
    else:
        # --- Fallback: старый способ ---
        plaintiff_match = re.search(
            r"((?:Обществ[оа] с ограниченной ответственностью|"
            r"Индивидуальный предприниматель|"
            r"Закрытое акционерное общество|"
            r"Публичное акционерное общество|"
            r"Открытое акционерное общество|"
            r"Акционерное общество|АО|ООО|ИП|ЗАО|ПАО)[^«]*«.+?»)\s+ИНН\s+"
            r"(\d+)\s+КПП\s+(\d+)?\s+ОГРН\s+(\d+)?\s+(.+?)\s*\n",
            text, re.DOTALL
        )
        plaintiff = {
            'name': (
                plaintiff_match.group(1).strip()
                if plaintiff_match else "Не указано"
            ),
            'inn': (
                plaintiff_match.group(2).strip()
                if plaintiff_match else "Не указано"
            ),
            'kpp': (
                plaintiff_match.group(3).strip()
                if plaintiff_match else "Не указано"
            ),
            'ogrn': (
                plaintiff_match.group(4).strip()
                if plaintiff_match else "Не указано"
            ),
            'address': (
                plaintiff_match.group(5).strip()
                if plaintiff_match else "Не указано"
            ),
        }
        # Сокращаем длинные формы до аббревиатуры
        plaintiff['name'] = re.sub(
            r"Обществ[оа] с ограниченной ответственностью", "ООО", plaintiff['name'])
        plaintiff['name'] = re.sub(
            r"Индивидуальный предприниматель", "ИП", plaintiff['name'])
        plaintiff['name'] = re.sub(
            r"Закрытое акционерное общество", "ЗАО", plaintiff['name'])
        plaintiff['name'] = re.sub(
            r"Публичное акционерное общество", "ПАО", plaintiff['name'])
        plaintiff['name'] = re.sub(
            r"Открытое акционерное общество", "ОАО", plaintiff['name'])
        plaintiff['name'] = re.sub(
            r"Акционерное общество", "АО", plaintiff['name'])

        defendant_match = re.search(
            r"(Обществ[оа] с ограниченной ответственностью|"
            r"Индивидуальному предпринимателю|"
            r"Закрытому акционерному обществу|"
            r"Публичному акционерному обществу|"
            r"Открытому акционерному обществу|"
            r"Акционерному обществу|АО|ООО|ИП|ЗАО|ПАО)\s+«(.+?)»\s+"
            r"ИНН\s+(\d+)\s+КПП\s+(\d+)\s+ОГРН\s+(\d+)\s+(.+?)\s*\n",
            text, re.DOTALL
        )
        defendant = {
            'name': (
                f"{defendant_match.group(1)} «{defendant_match.group(2)}»"
                if defendant_match else "Не указано"
            ),
            'inn': (
                defendant_match.group(3).strip()
                if defendant_match else "Не указано"
            ),
            'kpp': (
                defendant_match.group(4).strip()
                if defendant_match else "Не указано"
            ),
            'ogrn': (
                defendant_match.group(5).strip()
                if defendant_match else "Не указано"
            ),
            'address': (
                defendant_match.group(6).strip()
                if defendant_match else "Не указано"
            ),
        }

    # Извлекаем сумму задолженности
    debt_match = re.search(
        r'Стоимость услуг по договор[^0-9]*составила\s*([0-9\s,]+)\s*рубл',
        text, re.IGNORECASE
    )
    debt = 0.0
    if debt_match:
        debt_str = debt_match.group(1).replace(' ', '').replace(',', '.')
        try:
            debt = float(debt_str)
        except ValueError:
            pass

    # Извлекаем юридические услуги
    legal_fees_match = re.search(
        r'юридические услуги\s*([0-9\s,]+)\s*рубл',
        text, re.IGNORECASE
    )
    legal_fees = 0.0
    if legal_fees_match:
        legal_fees_str = legal_fees_match.group(
            1).replace(' ', '').replace(',', '.')
        try:
            legal_fees = float(legal_fees_str)
        except ValueError:
            pass

    # Извлекаем подписанта
    signatory_match = re.search(
        r'_________________\s*/([^/]+)/',
        text
    )
    signatory = signatory_match.group(
        1).strip() if signatory_match else "Не указано"

    # Извлекаем приложения
    attachments = parse_attachments(text)

    # Извлекаем счета и УПД
    invoices = []
    upds = []

    # Счета на оплату
    invoice_matches = re.finditer(
        r'Счет[а-яё]* на оплату\s*№\s*(\d+)\s*от\s*(\d{2}\.\d{2}\.\d{4})',
        text, re.IGNORECASE
    )
    for match in invoice_matches:
        invoices.append(f"№ {match.group(1)} от {match.group(2)}")

    # УПД
    upd_matches = re.finditer(
        r'УПД\s*№\s*(\d+)\s*от\s*(\d{2}\.\d{2}\.\d{4})',
        text, re.IGNORECASE
    )
    for match in upd_matches:
        upds.append(f"№ {match.group(1)} от {match.group(2)}")

    # Извлекаем блоки документов
    from sliding_window_parser import parse_documents_with_sliding_window
    document_blocks = parse_documents_with_sliding_window(text)

    return {
        'plaintiff': plaintiff,
        'defendant': defendant,
        'debt': debt,
        'legal_fees': legal_fees,
        'signatory': signatory,
        'attachments': attachments,
        'invoices': invoices,
        'upds': upds,
        'contract_applications': parse_contract_applications(text),
        'cargo_docs': parse_cargo_documents(text),
        'document_blocks': document_blocks
    }


def parse_invoice_blocks(text: str) -> str:
    """
    Парсит целые фразы, связанные со счетами на оплату (все формы).
    Возвращает строку для подстановки в шаблон.
    """
    patterns = [
        # Счет на оплату (ед. и мн. число, с номерами и датами)
        (r'(?:счет[а-яё]* на оплату|счета на оплату|счет[а-яё]*|счета)[^;\n\.]*?'
         r'(?:№\s*\d+[^;\n\.]*?от\s*\d{2}\.\d{2}\.\d{4}[^;\n\.]*?г?\.?|'
         r'от\s*\d{2}\.\d{2}\.\d{4}[^;\n\.]*?г?\.?|[^;\n\.]*?г?\.?)[^;\n\.]*'),
    ]
    found_blocks = []
    for pattern in patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            block = match.group(0).strip()
            block = re.sub(r'\s+', ' ', block)
            if block not in found_blocks:
                found_blocks.append(block)
    if found_blocks:
        return '; '.join(found_blocks)
    else:
        return "Не указано"


def parse_upd_blocks(text: str) -> str:
    """
    Парсит целые фразы, связанные с УПД (универсальный передаточный документ).
    Возвращает строку для подстановки в шаблон.
    """
    patterns = [
        # УПД и универсальный передаточный документ (ед. и мн. число)
        (r'(?:УПД|универсальн[а-яё ]*передаточн[а-яё ]*документ[а-яё]*)[^;\n\.]*?'
         r'(?:№\s*\d+[^;\n\.]*?от\s*\d{2}\.\d{2}\.\d{4}[^;\n\.]*?г?\.?|'
         r'от\s*\d{2}\.\d{2}\.\d{4}[^;\n\.]*?г?\.?|[^;\n\.]*?г?\.?)[^;\n\.]*'),
    ]
    found_blocks = []
    for pattern in patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            block = match.group(0).strip()
            block = re.sub(r'\s+', ' ', block)
            if block not in found_blocks:
                found_blocks.append(block)
    if found_blocks:
        return '; '.join(found_blocks)
    else:
        return "Не указано"


def extract_document_blocks(text: str) -> Dict[str, str]:
    """
    Основная функция извлечения блоков документов.
    Использует парсер из sliding_window_parser.py.

    Args:
        text: Текст для парсинга

    Returns:
        Словарь с извлеченными блоками документов
    """
    from sliding_window_parser import parse_documents_with_sliding_window
    return parse_documents_with_sliding_window(text)
