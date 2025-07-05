#!/usr/bin/env python3
"""
Тестовый скрипт для проверки парсинга почтовых уведомлений.
"""

import re

from docx import Document

from cal import calculate_duty
from calc_395 import calculate_full_395, get_key_rates_from_395gk
from main import get_court_by_address
from parsing import parse_cargo_documents, parse_claim_data


def test_postal_parsing(docx_path: str):
    """Тестирует различные регулярные выражения для парсинга почтовых уведомлений."""
    doc = Document(docx_path)
    text = "\n".join(p.text for p in doc.paragraphs)

    print("=== ТЕКСТ ДОКУМЕНТА ===")
    print(text)
    print("\n" + "="*50 + "\n")

    # Текущее регулярное выражение из parsing.py
    print("=== ТЕКУЩЕЕ РЕГУЛЯРНОЕ ВЫРАЖЕНИЕ ИЗ PARSING.PY ===")
    postal_patterns = [
        # Стандартный формат: "почтовым уведомлением № 123 от 01.01.2024"
        r"почтов[а-яё ]*уведомлени[еия][^№\d]*(?:№\s*)?(\d+)[^\d]*(\d{2}\.\d{2}\.\d{4})",
        # Формат: "№ 123 от 01.01.2024 - почтовое уведомление"
        r"(?:№\s*)?(\d+)[^\d]*(\d{2}\.\d{2}\.\d{4})[^.]*почтов[а-яё ]*уведомлени[еия]",
        # Формат: "уведомлением № 123 от 01.01.2024"
        r"уведомлени[еия][^№\d]*(?:№\s*)?(\d+)[^\d]*(\d{2}\.\d{2}\.\d{4})",
    ]

    all_matches = []
    for i, pattern in enumerate(postal_patterns, 1):
        matches = re.findall(pattern, text, re.IGNORECASE)
        print(f"Паттерн {i}: найдено {len(matches)} совпадений")
        for j, (num, date) in enumerate(matches, 1):
            print(f"  {j}. Номер: {num}, Дата: {date}")
        all_matches.extend(matches)

    # Убираем дубликаты, сохраняя порядок
    seen = set()
    unique_matches = []
    for num, date in all_matches:
        if (num, date) not in seen:
            seen.add((num, date))
            unique_matches.append((num, date))

    print(f"\nУникальных совпадений: {len(unique_matches)}")
    for i, (num, date) in enumerate(unique_matches, 1):
        print(f"{i}. № {num} от {date}")

    print("\n" + "="*50 + "\n")

    # Тестируем функцию generate_postal_block
    print("=== ТЕСТ ФУНКЦИИ GENERATE_POSTAL_BLOCK ===")
    postal_numbers = [num for num, _ in unique_matches]
    postal_dates = [date for _, date in unique_matches]

    if len(postal_numbers) == 1:
        result = (
            f"Почтовым уведомлением № {postal_numbers[0]} об отправке и получении "
            f"{postal_dates[0]} оригиналов документов Заказчиком."
        )
    else:
        pairs = [
            f"№ {num} от {date}" for num, date in zip(postal_numbers, postal_dates)
        ]
        result = (
            f"Почтовыми уведомлениями {', '.join(pairs)} "
            f"об отправке и получении оригиналов документов Заказчиком."
        )

    print(f"Результат: {result}")


def find_postal_placeholders(docx_path: str):
    doc = Document(docx_path)
    pattern = re.compile(
        r'\{postal(_block|_numbers|_receive_date)[^}]*\}', re.IGNORECASE)
    print(f"=== Поиск плейсхолдеров в {docx_path} ===")
    for i, p in enumerate(doc.paragraphs, 1):
        matches = pattern.findall(p.text)
        if matches:
            print(f"Строка {i}: {p.text}")
    print("=== Конец поиска ===")


def test_postal_parsing_real_doc(docx_path: str):
    print(f"=== Тест парсинга почтовых уведомлений в {docx_path} ===")
    claim_data = parse_claim_data(docx_path)
    postal_numbers = claim_data.get('postal_numbers', [])
    postal_dates = claim_data.get('postal_dates', [])
    print(f"Найдено уведомлений: {len(postal_numbers)}")
    for i, (num, date) in enumerate(zip(postal_numbers, postal_dates), 1):
        print(f"{i}. № {num} от {date}")
    print("=== Конец теста ===")


def debug_cargo_parsing():
    """Отладка парсинга грузосопроводительных документов"""
    try:
        # Парсим документ
        claim_data = parse_claim_data("test_document.docx")

        print("=== ОТЛАДКА ПАРСИНГА ГРУЗОСОПРОВОДИТЕЛЬНЫХ ДОКУМЕНТОВ ===")
        print(
            f"cargo_documents: {claim_data.get('cargo_documents', 'НЕ НАЙДЕНО')}")
        print()

        # Читаем весь текст документа
        doc = Document("test_document.docx")
        full_text = "\n".join(p.text for p in doc.paragraphs)

        print("=== ПОЛНЫЙ ТЕКСТ ДОКУМЕНТА ===")
        print(full_text)
        print()

        # Ищем вручную фразы с грузосопроводительными документами
        print("=== ПОИСК ФРАЗ С ГРУЗОСОПРОВОДИТЕЛЬНЫМИ ДОКУМЕНТАМИ ===")
        lines = full_text.split('\n')
        for i, line in enumerate(lines):
            if any(keyword in line.lower() for keyword in [
                'комплектом сопроводительных документов',
                'транспортная накладная',
                'товарно-транспортная накладная',
                'товарная накладная',
                'счет-фактура'
            ]):
                print(f"Строка {i+1}: {line}")

        print()

        # Тестируем функцию parse_cargo_documents отдельно
        print("=== ТЕСТ ФУНКЦИИ parse_cargo_documents ===")
        result = parse_cargo_documents(full_text)
        print(f"Результат parse_cargo_documents: {result}")

    except Exception as e:
        print(f"Ошибка: {e}")
        import traceback
        traceback.print_exc()


def debug_cargo_parsing_detailed():
    """Детальная отладка парсинга грузосопроводительных документов"""
    try:
        # Читаем весь текст документа
        doc = Document("test_document_with_cargo.docx")
        full_text = "\n".join(p.text for p in doc.paragraphs)

        print("=== ДЕТАЛЬНАЯ ОТЛАДКА ПАРСИНГА ГРУЗОСОПРОВОДИТЕЛЬНЫХ ДОКУМЕНТОВ ===")
        print("Полный текст документа:")
        print(full_text)
        print("\n" + "="*80)

        # Тестируем каждый паттерн отдельно
        from parsing import parse_cargo_documents

        # Паттерны из функции parse_cargo_documents
        patterns = [
            # Комплект сопроводительных документов
            r'Комплектом сопроводительных документов на груз[^;\n]*',
            # Транспортная накладная
            (r'[Тт]ранспортн[а-яё]* накладн[а-яё]*[^;\n]*?'
             r'(?:№\s*\d+[^;\n]*?от\s*\d{2}\.\d{2}\.\d{4}[^;\n]*?г?\.?|'
             r'от\s*\d{2}\.\d{2}\.\d{4}[^;\n]*?г?\.?|[^;\n]*?г?\.?)[^;\n]*'),
            # Товарно-транспортная накладная
            (r'[Тт]оварно-транспортн[а-яё]* накладн[а-яё]*[^;\n]*?'
             r'(?:№\s*\d+[^;\n]*?от\s*\d{2}\.\d{2}\.\d{4}[^;\n]*?г?\.?|'
             r'от\s*\d{2}\.\d{2}\.\d{4}[^;\n]*?г?\.?|[^;\n]*?г?\.?)[^;\n]*'),
            # Товарная накладная
            (r'[Тт]оварн[а-яё]* накладн[а-яё]*[^;\n]*?'
             r'(?:№\s*\d+[^;\n]*?от\s*\d{2}\.\d{2}\.\d{4}[^;\n]*?г?\.?|'
             r'от\s*\d{2}\.\d{2}\.\d{4}[^;\n]*?г?\.?|[^;\n]*?г?\.?)[^;\n]*'),
            # Счет-фактура
            (r'[Сс]чет-фактур[а-яё]*[^;\n]*?'
             r'(?:№\s*\d+[^;\n]*?от\s*\d{2}\.\d{2}\.\d{4}[^;\n]*?г?\.?|'
             r'от\s*\d{2}\.\d{2}\.\d{4}[^;\n]*?г?\.?|[^;\n]*?г?\.?)[^;\n]*'),
        ]

        print("Тестирование каждого паттерна:")
        for i, pattern in enumerate(patterns):
            print(f"\nПаттерн {i+1}: {pattern}")
            matches = list(re.finditer(pattern, full_text, re.IGNORECASE))
            print(f"Найдено совпадений: {len(matches)}")
            for j, match in enumerate(matches):
                print(f"  {j+1}. '{match.group(0)}'")

        print("\n" + "="*80)
        print("Результат функции parse_cargo_documents:")
        result = parse_cargo_documents(full_text)
        print(f"'{result}'")

    except Exception as e:
        print(f"Ошибка: {e}")
        import traceback
        traceback.print_exc()


def debug_placeholders():
    file_path = 'test_document.docx'
    claim_data = parse_claim_data(file_path)
    key_rates = get_key_rates_from_395gk()
    interest_data = calculate_full_395(file_path, key_rates=key_rates)
    total_claim = claim_data['debt'] + interest_data['total_interest']
    duty_data = calculate_duty(total_claim)
    court_name, court_address = get_court_by_address(
        claim_data['defendant']['address'])
    replacements = {
        '{court_name}': court_name,
        '{court_address}': court_address,
        '{plaintiff_name}': claim_data['plaintiff']['name'],
        '{plaintiff_inn}': claim_data['plaintiff']['inn'],
        '{plaintiff_kpp}': claim_data['plaintiff']['kpp'],
        '{plaintiff_ogrn}': claim_data['plaintiff']['ogrn'],
        '{plaintiff_address}': claim_data['plaintiff']['address'],
        '{defendant_name}': claim_data['defendant']['name'],
        '{defendant_inn}': claim_data['defendant']['inn'],
        '{defendant_kpp}': claim_data['defendant']['kpp'],
        '{defendant_ogrn}': claim_data['defendant']['ogrn'],
        '{defendant_address}': claim_data['defendant']['address'],
        '{total_claim}': f"{total_claim:,.2f}".replace(',', ' '),
        '{duty}': f"{duty_data['duty']:,.0f}".replace(',', ' '),
        '{debt}': f"{claim_data['debt']:,.2f}".replace(',', ' '),
        '{contracts}': ", ".join(claim_data['contracts']) if claim_data['contracts'] else 'Не указано',
        '{contract_applications}': claim_data['contract_applications'],
        '{cargo_docs}': claim_data['cargo_documents'],
        '{invoice_blocks}': claim_data['invoice_blocks'],
        '{upd_blocks}': claim_data['upd_blocks'],
        '{invoices}': ", ".join(claim_data['invoices']) if claim_data['invoices'] else 'Не указано',
        '{upds}': ", ".join(claim_data['upds']) if claim_data['upds'] else 'Не указано',
        '{total_interest}': f"{interest_data['total_interest']:,.2f}".replace(',', ' '),
        '{legal_fees}': f"{claim_data['legal_fees']:,.2f}".replace(',', ' '),
        '{signatory}': claim_data['signatory'],
        '{claim_paragraph}': 'Тестовый параграф претензии',
        '{postal_block}': 'Тестовый почтовый блок',
        '{postal_numbers_all}': '123456',
        '{postal_dates_all}': '01.01.2025',
        '{claim_date}': '01.01.2025',
        '{claim_number}': '123456',
        '{total_expenses}': f"{duty_data['duty'] + claim_data['legal_fees']:,.0f}".replace(',', ' '),
        '{calculation_date}': '01.01.2025',
        '{postal_numbers}': '123456',
        '{postal_receive_date}': '01.01.2025'
    }
    print("=== КЛЮЧИ REPLACEMENTS ===")
    for k, v in replacements.items():
        print(f"{k}: {v}")
    print("\n=== ПЛЕЙСХОЛДЕРЫ В ШАБЛОНЕ ===")
    doc = Document('template.docx')
    found = set()
    for paragraph in doc.paragraphs:
        for word in paragraph.text.split():
            if word.startswith('{') and word.endswith('}'):
                found.add(word)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for word in cell.text.split():
                    if word.startswith('{') and word.endswith('}'):
                        found.add(word)
    for ph in sorted(found):
        print(ph)


if __name__ == "__main__":
    # Тестируем с документом с одним почтовым уведомлением
    test_postal_parsing("test_single_postal.docx")
    find_postal_placeholders("template.docx")
    # Укажите здесь путь к вашему реальному docx-файлу с несколькими уведомлениями
    test_postal_parsing_real_doc("test_document.docx")
    debug_cargo_parsing()
    debug_cargo_parsing_detailed()
    debug_placeholders()
