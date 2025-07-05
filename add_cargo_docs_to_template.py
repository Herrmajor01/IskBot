#!/usr/bin/env python3
"""
Скрипт для добавления плейсхолдера {cargo_docs} в шаблон
"""

from docx import Document


def add_cargo_docs_to_template():
    """Добавляет строку с грузосопроводительными документами в шаблон"""
    doc = Document('template.docx')

    # Находим параграф со счетами на оплату (строка 13)
    target_paragraph = None
    for i, paragraph in enumerate(doc.paragraphs):
        if '{invoices}' in paragraph.text:
            target_paragraph = paragraph
            target_index = i
            break

    if target_paragraph:
        print(f"Найден параграф со счетами на строке {target_index + 1}")

        # Создаем новый параграф для грузосопроводительных документов
        new_paragraph = doc.add_paragraph()
        new_paragraph.text = "· Грузосопроводительными документами {cargo_docs};"

        # Вставляем новый параграф после параграфа со счетами
        # Для этого нужно переместить все последующие параграфы
        paragraphs_to_move = []
        for i in range(target_index + 1, len(doc.paragraphs)):
            paragraphs_to_move.append(doc.paragraphs[i])

        # Удаляем старые параграфы
        for paragraph in paragraphs_to_move:
            p = paragraph._element
            p.getparent().remove(p)

        # Добавляем новый параграф
        new_p = doc.add_paragraph()
        new_p.text = "· Грузосопроводительными документами {cargo_docs};"

        # Добавляем обратно перемещенные параграфы
        for paragraph in paragraphs_to_move:
            doc.add_paragraph(paragraph.text)

        # Сохраняем изменения
        doc.save('template.docx')
        print("Плейсхолдер {cargo_docs} добавлен в шаблон")
    else:
        print("Не найден параграф со счетами на оплату")


if __name__ == "__main__":
    add_cargo_docs_to_template()
