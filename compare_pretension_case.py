import argparse
import json
import re
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document

from case_registry import CaseRegistry
from main import (
    add_working_days,
    apply_llm_fallback,
    apply_extracted_parties,
    assign_shipments_to_groups,
    build_document_groups,
    build_documents_list,
    build_documents_list_structured,
    build_documents_list_structured_for_groups,
    build_intro_paragraph,
    build_legal_fees_block,
    build_party_block,
    build_payment_terms_summary,
    build_pretension_attachments,
    build_pretension_groups,
    build_requirements_summary,
    build_shipping_summary,
    calculate_pretension_interest,
    calculate_pretension_interest_schedule,
    create_pretension_document,
    extract_application_payment_terms,
    extract_applications_from_pages,
    extract_cargo_docs_from_pages,
    extract_cdek_shipments_from_pages,
    extract_invoices_from_pages,
    extract_parties_from_pages,
    extract_payment_terms_from_text,
    extract_postal_shipments_from_pages,
    extract_upd_from_pages,
    extract_pdf_pages,
    apply_vision_ocr_to_pages,
    format_money,
    format_organization_name_short,
    get_ogrn_label,
    get_pretension_missing_fields,
    load_work_calendar,
    normalize_payment_terms,
    normalize_shipping_source,
    normalize_str,
    parse_amount,
    parse_date_str,
    parse_documents_with_sliding_window,
    split_document_items,
)


def extract_docx_text(path: Path) -> List[str]:
    doc = Document(path)
    lines: List[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                lines.append(row_text)
    return lines


def extract_manual_fields(lines: List[str]) -> Tuple[Dict[str, Any], List[float], List[Tuple[str, Optional[str]]]]:
    manual_fields: Dict[str, Any] = {}

    def extract_block(label: str) -> Dict[str, str]:
        for idx, line in enumerate(lines):
            if line.startswith(label):
                name = line.split(":", 1)[1].strip()
                block = {"name": name}
                for j in range(idx + 1, len(lines)):
                    current = lines[j]
                    if current.startswith("Кому:") or current.startswith("От кого:") or current.startswith("ПРЕТЕНЗИЯ"):
                        break
                    if current.startswith("ИНН"):
                        block["inn"] = re.sub(r"[^0-9]", "", current)
                    elif current.startswith("КПП"):
                        block["kpp"] = re.sub(r"[^0-9]", "", current)
                    elif current.startswith("ОГРН"):
                        block["ogrn"] = re.sub(r"[^0-9]", "", current)
                    else:
                        if "address" not in block:
                            block["address"] = current
                return block
        return {}

    manual_fields["defendant"] = extract_block("Кому:")
    manual_fields["plaintiff"] = extract_block("От кого:")

    manual_text = "\n".join(lines)
    debt_match = re.search(r"Итого задолженность:\s*([\d\s]+)", manual_text)
    if debt_match:
        manual_fields["debt"] = debt_match.group(1).strip()

    shipping_match = re.search(
        r"Документы по перевозкам были отправлены.*",
        manual_text,
        re.IGNORECASE
    )
    if shipping_match:
        shipping_line = shipping_match.group(0)
        lower = shipping_line.lower()
        if "сдэк" in lower:
            manual_fields["shipping_method"] = "сдэк"
        elif "почтов" in lower:
            manual_fields["shipping_method"] = "почта"
        manual_fields["postal_numbers"] = re.findall(r"\d{8,}", shipping_line)
        manual_fields["postal_dates"] = re.findall(r"\d{2}\.\d{2}\.\d{4}", shipping_line)

    contract_match = re.search(
        r"Договор оказания юридических услуг\s*№\s*([^\s]+)\s*от\s*(\d{2}\.\d{2}\.\d{4})",
        manual_text
    )
    if contract_match:
        manual_fields["legal_contract_number"] = contract_match.group(1)
        manual_fields["legal_contract_date"] = contract_match.group(2)

    payment_match = re.search(
        r"Плат[её]жным поручением\s*№\s*([^\s]+)\s*от\s*(\d{2}\.\d{2}\.\d{4})",
        manual_text
    )
    if payment_match:
        manual_fields["legal_payment_number"] = payment_match.group(1)
        manual_fields["legal_payment_date"] = payment_match.group(2)

    legal_fee_match = re.search(r"на сумму\s*([\d\s]+)\s*руб", manual_text)
    if legal_fee_match:
        manual_fields["legal_fees"] = legal_fee_match.group(1).strip()

    price_values = [
        parse_amount(value)
        for value in re.findall(r"Цена перевозки\s*([\d\s]+)", manual_text)
    ]
    shipment_lines = [
        line for line in lines
        if "отправлен" in line.lower() and "№" in line
    ]
    shipment_pairs: List[Tuple[str, Optional[str]]] = []
    for line in shipment_lines:
        track_match = re.search(r"№\s*(\d{8,})", line)
        date_match = re.search(r"(\d{2}\.\d{2}\.\d{4})", line)
        if track_match:
            shipment_pairs.append((track_match.group(1), date_match.group(1) if date_match else None))

    return manual_fields, price_values, shipment_pairs


def fill_claim_from_manual(
    claim_data: Dict[str, Any],
    groups: List[Dict[str, Any]],
    manual_fields: Dict[str, Any],
    price_values: List[float],
    shipment_pairs: List[Tuple[str, Optional[str]]]
) -> List[str]:
    filled_fields: List[str] = []

    def fill_if_missing(key: str, value: Optional[str]) -> None:
        if not value:
            return
        current = claim_data.get(key)
        if current is None or str(current).strip() in ("", "Не указано"):
            claim_data[key] = value
            filled_fields.append(key)

    plaintiff = manual_fields.get("plaintiff", {})
    defendant = manual_fields.get("defendant", {})
    fill_if_missing("plaintiff_name", plaintiff.get("name"))
    fill_if_missing("plaintiff_inn", plaintiff.get("inn"))
    fill_if_missing("plaintiff_kpp", plaintiff.get("kpp"))
    fill_if_missing("plaintiff_ogrn", plaintiff.get("ogrn"))
    fill_if_missing("plaintiff_address", plaintiff.get("address"))
    fill_if_missing("defendant_name", defendant.get("name"))
    fill_if_missing("defendant_inn", defendant.get("inn"))
    fill_if_missing("defendant_kpp", defendant.get("kpp"))
    fill_if_missing("defendant_ogrn", defendant.get("ogrn"))
    fill_if_missing("defendant_address", defendant.get("address"))

    fill_if_missing("debt", manual_fields.get("debt"))
    fill_if_missing("legal_contract_number", manual_fields.get("legal_contract_number"))
    fill_if_missing("legal_contract_date", manual_fields.get("legal_contract_date"))
    fill_if_missing("legal_payment_number", manual_fields.get("legal_payment_number"))
    fill_if_missing("legal_payment_date", manual_fields.get("legal_payment_date"))
    fill_if_missing("legal_fees", manual_fields.get("legal_fees"))

    if manual_fields.get("shipping_method") and not claim_data.get("shipping_method"):
        claim_data["shipping_method"] = manual_fields["shipping_method"]
        filled_fields.append("shipping_method")

    if manual_fields.get("postal_numbers") and not claim_data.get("postal_numbers"):
        claim_data["postal_numbers"] = manual_fields["postal_numbers"]
        filled_fields.append("postal_numbers")
    if manual_fields.get("postal_dates") and not claim_data.get("postal_dates"):
        claim_data["postal_dates"] = manual_fields["postal_dates"]
        filled_fields.append("postal_dates")
    if manual_fields.get("postal_numbers") and not claim_data.get("docs_track_number"):
        claim_data["docs_track_number"] = manual_fields["postal_numbers"][0]
        filled_fields.append("docs_track_number")
    if manual_fields.get("postal_dates") and not claim_data.get("docs_received_date"):
        claim_data["docs_received_date"] = manual_fields["postal_dates"][0]
        filled_fields.append("docs_received_date")

    if price_values and groups:
        for group, price in zip(groups, price_values):
            if (group.get("amount") or 0) <= 0 and price > 0:
                group["amount"] = price
                if group.get("invoice"):
                    group["invoice_amount"] = price
                elif group.get("upd"):
                    group["upd_amount"] = price

    if shipment_pairs and groups:
        source = normalize_shipping_source(claim_data.get("shipping_method"))
        for group, (track, date_str) in zip(groups, shipment_pairs):
            if track and not group.get("docs_track_number"):
                group["docs_track_number"] = track
            if date_str and not group.get("docs_received_date"):
                group["docs_received_date"] = date_str
            if source and not group.get("shipping_source"):
                group["shipping_source"] = source

    if not claim_data.get("shipments") and manual_fields.get("postal_numbers"):
        source = normalize_shipping_source(claim_data.get("shipping_method"))
        claim_data["shipments"] = [
            {
                "track_number": num,
                "received_date": parse_date_str(date_str) if date_str else None,
                "source": source,
            }
            for num, date_str in zip(
                manual_fields.get("postal_numbers", []),
                manual_fields.get("postal_dates", []),
            )
        ]

    return filled_fields


def build_claim_from_folder(folder: Path) -> Tuple[Dict[str, Any], List[Dict[str, Any]], List[str]]:
    pdfs = sorted(folder.glob("*.pdf"))
    combined_texts: List[str] = []
    all_pages: List[str] = []

    for pdf in pdfs:
        pages, low_text_pages = extract_pdf_pages(str(pdf))
        if low_text_pages:
            apply_vision_ocr_to_pages(str(pdf), pages, low_text_pages)
        all_pages.extend(pages)
        blocks = [f"[Страница {idx + 1}]\n{page}" for idx, page in enumerate(pages)]
        combined_texts.append(f"=== {pdf.name} ===\n" + "\n\n".join(blocks))

    combined_text = "\n\n".join(combined_texts)
    claim_data = parse_documents_with_sliding_window(combined_text)
    claim_data = apply_llm_fallback(combined_text, claim_data)
    claim_data["document_groups"] = build_document_groups(combined_text, claim_data)
    claim_data["source_files"] = [pdf.name for pdf in pdfs]

    applications = extract_applications_from_pages(all_pages)
    invoices = extract_invoices_from_pages(all_pages)
    upd_docs = extract_upd_from_pages(all_pages)
    cargo_docs = extract_cargo_docs_from_pages(all_pages)
    shipments = extract_cdek_shipments_from_pages(all_pages)
    shipments.extend(extract_postal_shipments_from_pages(all_pages))

    payment_terms_by_application = extract_application_payment_terms(all_pages, applications)
    groups = build_pretension_groups(
        applications,
        invoices,
        cargo_docs,
        upd_docs=upd_docs,
        payment_terms_by_application=payment_terms_by_application,
    )
    assign_shipments_to_groups(groups, shipments)

    payment_terms = None
    payment_days = None
    if payment_terms_by_application:
        terms_values = [
            normalize_payment_terms(item.get("terms") or "")
            for item in payment_terms_by_application.values()
            if item.get("terms")
        ]
        days_values = [item.get("days") for item in payment_terms_by_application.values() if item.get("days")]
        unique_terms = {value for value in terms_values if value}
        unique_days = {value for value in days_values if value}
        if len(unique_terms) == 1:
            payment_terms = unique_terms.pop()
        if len(unique_days) == 1:
            payment_days = unique_days.pop()
    elif applications:
        app_pages = [page for page in all_pages if "заявк" in page.lower()]
        payment_terms, payment_days = extract_payment_terms_from_text("\n".join(app_pages))

    if not payment_terms:
        payment_terms, payment_days = extract_payment_terms_from_text(combined_text)

    if payment_terms:
        claim_data["payment_terms"] = payment_terms
    if payment_days:
        claim_data["payment_days"] = str(payment_days)

    parties = extract_parties_from_pages(all_pages)
    if parties:
        apply_extracted_parties(claim_data, parties)

    total_debt = sum(group.get("amount") or 0.0 for group in groups)
    if total_debt > 0:
        claim_data["debt"] = format_money(total_debt, 0)

    if shipments:
        for shipment in shipments:
            shipment["source"] = normalize_shipping_source(shipment.get("source"))
        claim_data["shipments"] = shipments
        if any(item.get("source") == "post" for item in shipments):
            claim_data["shipping_method"] = "почта"
        elif any(item.get("source") == "cdek" for item in shipments):
            claim_data["shipping_method"] = "сдэк"
        claim_data["postal_numbers"] = [
            item["track_number"] for item in shipments if item.get("track_number")
        ]
        postal_dates = []
        for item in shipments:
            received = item.get("received_date")
            if isinstance(received, datetime):
                postal_dates.append(received.strftime("%d.%m.%Y"))
            elif isinstance(received, str):
                parsed = parse_date_str(received)
                if parsed:
                    postal_dates.append(parsed.strftime("%d.%m.%Y"))
        claim_data["postal_dates"] = postal_dates
        claim_data["docs_track_number"] = claim_data["postal_numbers"][0] if claim_data["postal_numbers"] else ""
        claim_data["docs_received_date"] = claim_data["postal_dates"][0] if claim_data["postal_dates"] else ""

    claim_data["pretension_groups"] = groups
    missing_fields = get_pretension_missing_fields(claim_data)
    return claim_data, groups, missing_fields


def generate_pretension_docx(
    claim_data: Dict[str, Any],
    groups: List[Dict[str, Any]],
    output_path: Path
) -> Path:
    plaintiff_name = normalize_str(claim_data.get("plaintiff_name"))
    defendant_name = normalize_str(claim_data.get("defendant_name"))
    plaintiff_name_short = format_organization_name_short(plaintiff_name)
    defendant_name_short = format_organization_name_short(defendant_name)
    is_plaintiff_ip = "ИП" in plaintiff_name or "Индивидуальный предприниматель" in plaintiff_name
    is_defendant_ip = "ИП" in defendant_name or "Индивидуальный предприниматель" in defendant_name

    debt_amount = parse_amount(claim_data.get("debt", "0"))
    payment_days_raw = claim_data.get("payment_days", "0")
    try:
        payment_days = int(re.sub(r"[^\d]", "", str(payment_days_raw)))
    except ValueError:
        payment_days = 0

    interest_data = {"total_interest": 0.0, "detailed_calc": []}
    has_group_payment_days = False
    for group in groups:
        try:
            if int(group.get("payment_days") or 0) > 0:
                has_group_payment_days = True
                break
        except (TypeError, ValueError):
            continue
    if groups and (payment_days > 0 or has_group_payment_days):
        interest_data = calculate_pretension_interest_schedule(groups, payment_days)
    else:
        docs_received_date = parse_date_str(claim_data.get("docs_received_date", ""))
        if docs_received_date and payment_days > 0 and debt_amount > 0:
            calendar = load_work_calendar(docs_received_date.year)
            due_date = add_working_days(docs_received_date, payment_days, calendar)
            interest_start = due_date + timedelta(days=1)
            interest_data = calculate_pretension_interest(debt_amount, interest_start)

    payment_terms_text = normalize_payment_terms(claim_data.get("payment_terms", ""))
    if not payment_terms_text or payment_terms_text == "Не указано":
        if payment_days > 0:
            payment_terms_text = (
                "Оплата не позднее "
                f"{payment_days} рабочих дней с даты получения документов, "
                "подтверждающих перевозку"
            )
        else:
            payment_terms_text = "Не указано"

    payment_terms_text = build_payment_terms_summary(
        groups,
        payment_terms_text,
        payment_days,
    )

    applications_list = [
        group.get("application")
        for group in groups
        if group.get("application")
    ]
    cargo_docs_list: List[str] = []
    for group in groups:
        cargo_docs_list.extend(group.get("cargo_docs", []))
    if not cargo_docs_list:
        cargo_docs_list = split_document_items(claim_data.get("cargo_docs"))

    intro_paragraph = build_intro_paragraph(
        plaintiff_name_short,
        applications_list,
        cargo_docs_list,
    )

    plaintiff_ogrn_type = get_ogrn_label(plaintiff_name, claim_data.get("plaintiff_inn", ""))
    defendant_ogrn_type = get_ogrn_label(defendant_name, claim_data.get("defendant_inn", ""))

    defendant_block = build_party_block(
        "Кому",
        defendant_name_short,
        normalize_str(claim_data.get("defendant_inn")),
        normalize_str(claim_data.get("defendant_kpp")),
        normalize_str(claim_data.get("defendant_ogrn")),
        defendant_ogrn_type,
        normalize_str(claim_data.get("defendant_address")),
        normalize_str(claim_data.get("defendant_address")),
        is_defendant_ip,
    )
    plaintiff_block = build_party_block(
        "От кого",
        plaintiff_name_short,
        normalize_str(claim_data.get("plaintiff_inn")),
        normalize_str(claim_data.get("plaintiff_kpp")),
        normalize_str(claim_data.get("plaintiff_ogrn")),
        plaintiff_ogrn_type,
        normalize_str(claim_data.get("plaintiff_address")),
        normalize_str(claim_data.get("plaintiff_address")),
        is_plaintiff_ip,
    )

    documents_list_structured = (
        build_documents_list_structured_for_groups(groups)
        if groups
        else build_documents_list_structured(claim_data.get("document_groups") or [])
    )

    attachments = build_pretension_attachments(groups or claim_data.get("document_groups") or [], claim_data)
    shipping_summary = build_shipping_summary(
        claim_data.get("shipments") or [],
        documents_count=len(groups) if groups else None
    )

    replacements = {
        "{defendant_block}": defendant_block,
        "{plaintiff_block}": plaintiff_block,
        "{intro_paragraph}": intro_paragraph,
        "{documents_list}": build_documents_list(claim_data),
        "{debt_amount}": format_money(debt_amount, 0),
        "{payment_terms}": payment_terms_text,
        "{legal_fees_block}": build_legal_fees_block(claim_data),
        "{requirements_summary}": build_requirements_summary(
            debt_amount,
            parse_amount(interest_data.get("total_interest", 0)),
            parse_amount(claim_data.get("legal_fees", 0))
        ),
        "{pretension_date}": "",
        "{shipping_info}": shipping_summary,
        "{docs_track_number}": normalize_str(claim_data.get("docs_track_number", ""), default=""),
        "{docs_received_date}": normalize_str(claim_data.get("docs_received_date", ""), default=""),
        "{plaintiff_name}": plaintiff_name,
        "{defendant_name}": defendant_name,
    }

    create_pretension_document(
        claim_data,
        interest_data,
        replacements,
        documents_list_structured=documents_list_structured,
        attachments=attachments,
        output_path=str(output_path),
    )
    return output_path


def build_snapshot(data: Dict[str, Any]) -> Dict[str, Any]:
    keys = [
        "plaintiff_name",
        "defendant_name",
        "plaintiff_inn",
        "defendant_inn",
        "plaintiff_address",
        "defendant_address",
        "debt",
        "payment_terms",
        "payment_days",
        "shipping_method",
        "docs_track_number",
        "docs_received_date",
        "legal_fees",
        "legal_contract_number",
        "legal_contract_date",
        "legal_payment_number",
        "legal_payment_date",
    ]
    return {key: data.get(key) for key in keys if key in data}


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Сравнивает претензию с эталоном и сохраняет кейс в SQLite."
    )
    parser.add_argument(
        "folder",
        nargs="?",
        default="Претензия и документы к ней",
        help="Папка с PDF и ПРЕТЕНЗИЯ.docx",
    )
    parser.add_argument("--db", help="Путь к SQLite реестру кейсов")
    parser.add_argument("--case-name", help="Название кейса")
    parser.add_argument("--output-dir", default="isk_outputs", help="Папка для выходных файлов")
    parser.add_argument("--no-fill", action="store_true", help="Не заполнять поля из ручной претензии")
    parser.add_argument("--notes", default="", help="Комментарий для кейса")
    args = parser.parse_args()

    folder = Path(args.folder)
    manual_docx = folder / "ПРЕТЕНЗИЯ.docx"
    if not manual_docx.exists():
        raise FileNotFoundError(f"Не найден файл: {manual_docx}")

    claim_data, groups, missing_fields = build_claim_from_folder(folder)
    manual_lines = extract_docx_text(manual_docx)
    manual_fields, price_values, shipment_pairs = extract_manual_fields(manual_lines)

    filled_fields: List[str] = []
    if not args.no_fill:
        filled_fields = fill_claim_from_manual(
            claim_data,
            groups,
            manual_fields,
            price_values,
            shipment_pairs,
        )

    claim_data["pretension_groups"] = groups

    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    output_docx = output_dir / "pretension_generated_compare.docx"
    generate_pretension_docx(claim_data, groups, output_docx)

    manual_text = extract_docx_text(manual_docx)
    generated_text = extract_docx_text(output_docx)

    diff_lines = list(
        __import__("difflib").unified_diff(
            manual_text,
            generated_text,
            lineterm="",
        )
    )
    diff_path = output_dir / "pretension_diff_compare.txt"
    diff_path.write_text("\n".join(diff_lines), encoding="utf-8")

    registry = CaseRegistry(args.db)
    summary = {
        "missing_fields": missing_fields,
        "filled_fields": filled_fields,
        "diff_lines": len(diff_lines),
    }
    case_id = registry.add_case(
        case_name=args.case_name or folder.name,
        folder_path=str(folder),
        manual_docx_path=str(manual_docx),
        generated_docx_path=str(output_docx),
        diff_path=str(diff_path),
        missing_fields={"missing": missing_fields},
        filled_fields={"filled": filled_fields},
        extracted_fields=build_snapshot(claim_data),
        manual_fields=manual_fields,
        summary=json.dumps(summary, ensure_ascii=False),
    )
    if args.notes:
        registry.add_observation(
            case_id=case_id,
            kind="note",
            note=args.notes,
        )
    registry.close()

    print(f"Generated: {output_docx}")
    print(f"Diff: {diff_path}")
    print(f"Missing fields: {missing_fields}")
    print(f"Filled fields: {filled_fields}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
